from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "usage-guide-editable-assets"
FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
DOC_FONT = "Arial Unicode MS"

OFFICIAL_DOC = OUT_DIR / "头号空间-官方运营使用教程_可编辑版.docx"
MERCHANT_DOC = OUT_DIR / "头号空间-商家门店使用教程_可编辑版.docx"
PLAYER_DOC = OUT_DIR / "头号空间-玩家体验使用教程_可编辑版.docx"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int):
    try:
        return ImageFont.truetype(FONT_PATH, size=size)
    except Exception:
        return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            width = draw.textbbox((0, 0), trial, font=font)[2]
            if width > max_width and current:
                lines.append(current)
                current = ch
            else:
                current = trial
        lines.append(current if current else "")
    return lines


def draw_box(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, title: str, desc: str, fill: str, outline: str) -> None:
    draw.rounded_rectangle([x, y, x + w, y + h], radius=20, fill=fill, outline=outline, width=3)
    title_font = load_font(28)
    desc_font = load_font(20)
    title_lines = wrap_text(draw, title, title_font, w - 28)
    desc_lines = wrap_text(draw, desc, desc_font, w - 28)
    ty = y + 18
    for line in title_lines:
        draw.text((x + 14, ty), line, font=title_font, fill="#0F172A")
        ty += 34
    ty += 6
    for line in desc_lines:
        draw.text((x + 14, ty), line, font=desc_font, fill="#334155")
        ty += 26


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#2563EB") -> None:
    draw.line([start, end], fill=color, width=5)
    if start[0] == end[0]:
        direction = 1 if end[1] > start[1] else -1
        points = [end, (end[0] - 9, end[1] - 16 * direction), (end[0] + 9, end[1] - 16 * direction)]
    else:
        direction = 1 if end[0] > start[0] else -1
        points = [end, (end[0] - 16 * direction, end[1] - 9), (end[0] - 16 * direction, end[1] + 9)]
    draw.polygon(points, fill=color)


def draw_title(draw: ImageDraw.ImageDraw, text: str, subtitle: str) -> None:
    title_font = load_font(44)
    sub_font = load_font(24)
    draw.text((70, 40), text, font=title_font, fill="#0F172A")
    draw.text((70, 98), subtitle, font=sub_font, fill="#475569")


def make_official_flow(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), "#F8FBFF")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "官方运营交付流程", "商家开通线与内容管理线并行推进，在内容分发处汇合")
    top_items = [
        ("商家线 1 新建商家", "开通后台账号、确认负责人", 80, 220),
        ("商家线 2 新建店铺", "补齐门店地址、状态和注册码", 400, 220),
        ("商家线 3 录入主机并生成 Token", "主机分配到门店后生成激活凭证", 720, 220),
        ("商家线 4 录入头显并绑定主机", "确保同店铺设备关系正确", 1040, 220),
    ]
    bottom_items = [
        ("内容线 1 录入 / 审核游戏", "把可运营内容放进平台游戏库", 170, 500),
        ("内容线 2 维护版本与运营参数", "明确平台、时长、消耗和状态", 520, 500),
        ("汇合点 内容分发", "把已准备好的游戏推送到目标门店", 900, 500),
        ("门店准备完成", "结合支付配置进入营业测试", 1250, 500),
    ]
    for title, desc, x, y in top_items:
        draw_box(draw, x, y, 250, 120, title, desc, "#FFFFFF", "#93C5FD")
    for title, desc, x, y in bottom_items:
        draw_box(draw, x, y, 250, 120, title, desc, "#FFFFFF", "#86EFAC" if "汇合点" in title or "门店准备完成" in title else "#FBBF24")
    for sx, ex in [(330, 400), (650, 720), (970, 1040)]:
        draw_arrow(draw, (sx, 280), (ex, 280))
    draw_arrow(draw, (1165, 340), (1165, 440))
    draw_arrow(draw, (295, 560), (520, 560))
    draw_arrow(draw, (770, 560), (900, 560))
    draw_arrow(draw, (1150, 560), (1250, 560))
    note_font = load_font(22)
    draw.rounded_rectangle([70, 760, 1530, 850], radius=18, fill="#E0F2FE", outline="#7DD3FC", width=2)
    note = "关键理解：录入游戏属于平台内容管理线，不是商家开通线的串行步骤。真正与门店交付汇合的节点是“内容分发到门店”。"
    for i, line in enumerate(wrap_text(draw, note, note_font, 1400)):
        draw.text((100, 790 + i * 28), line, font=note_font, fill="#0F172A")
    img.save(path)


def make_official_device_map(path: Path) -> None:
    img = Image.new("RGB", (1600, 780), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "官方交付示意图", "账号、设备、内容、支付四类信息要成套交付")
    draw_box(draw, 80, 220, 280, 150, "平台侧交付包", "商家后台账号\n店铺信息\n负责人联系方式", "#F8FAFC", "#94A3B8")
    draw_box(draw, 470, 220, 280, 150, "设备交付包", "主机编号\n头显清单\nToken\n建议设备密码", "#F8FAFC", "#94A3B8")
    draw_box(draw, 860, 220, 280, 150, "内容交付包", "已分发游戏清单\n主推内容\n是否需 USB 安装", "#F8FAFC", "#94A3B8")
    draw_box(draw, 1250, 220, 280, 150, "支付交付包", "拉卡拉状态\n测试单结果\n异常联系人", "#F8FAFC", "#94A3B8")
    draw_arrow(draw, (360, 295), (470, 295), "#64748B")
    draw_arrow(draw, (750, 295), (860, 295), "#64748B")
    draw_arrow(draw, (1140, 295), (1250, 295), "#64748B")
    draw.rounded_rectangle([140, 500, 1460, 650], radius=22, fill="#EFF6FF", outline="#60A5FA", width=3)
    big_font = load_font(28)
    lines = [
        "只有当这四类信息完整交接给商家后，门店才能顺利完成后台配置、点播系统 Token 绑定和首单测试。",
        "任何一个环节缺失，都会在现场表现为：看不到内容、装不上游戏、无法开局或不能收款。",
    ]
    for idx, line in enumerate(lines):
        draw.text((180, 540 + idx * 40), line, font=big_font, fill="#1E3A8A")
    img.save(path)


def make_merchant_prep_flow(path: Path) -> None:
    img = Image.new("RGB", (1600, 840), "#FCFFFE")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "门店开业准备流程", "门店管理员在营业前需要完成的最小闭环")
    items = [
        ("1 核对交付资料", "账号、门店、设备、Token 是否齐全", 80, 220),
        ("2 登录商户后台", "确认门店资料和设备归属", 420, 220),
        ("3 设置点播系统密码", "保护管理模式不被顾客误触", 760, 220),
        ("4 配置价格和会员活动", "单次消费、充值套餐、优惠券", 1100, 220),
        ("5 在点播系统写入 Token", "完成门店绑定与内容同步", 280, 500),
        ("6 下载并安装主推游戏", "验证主机游戏和一体机内容", 700, 500),
        ("7 做首单测试", "测试支付、开局和结束回传", 1120, 500),
    ]
    for title, desc, x, y in items:
        draw_box(draw, x, y, 260, 120, title, desc, "#FFFFFF", "#86EFAC" if y > 300 else "#93C5FD")
    for sx, ex in [(340, 420), (680, 760), (1020, 1100)]:
        draw_arrow(draw, (sx, 280), (ex, 280))
    draw_arrow(draw, (1230, 340), (1230, 430))
    draw_arrow(draw, (1230, 430), (410, 430))
    draw_arrow(draw, (410, 430), (410, 500))
    draw_arrow(draw, (540, 560), (700, 560))
    draw_arrow(draw, (960, 560), (1120, 560))
    img.save(path)


def make_merchant_system_map(path: Path) -> None:
    img = Image.new("RGB", (1600, 820), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "门店后台与点播系统协同示意", "后台负责规则，终端负责体验")
    draw_box(draw, 150, 220, 500, 220, "商户后台", "设备列表：核对主机/头显\n系统设置：修改点播系统密码\n商品管理：单次消费项目\n会员/运营：充值套餐、会员级别、优惠券", "#EFF6FF", "#60A5FA")
    draw_box(draw, 950, 220, 500, 220, "PC 点播系统", "写入 Token\n同步门店内容\n下载和安装游戏\n现场选游戏、扫码支付、开局引导", "#ECFDF5", "#34D399")
    draw_arrow(draw, (650, 330), (950, 330), "#64748B")
    draw.rounded_rectangle([230, 540, 1370, 680], radius=20, fill="#F8FAFC", outline="#CBD5E1", width=2)
    tip_font = load_font(24)
    tips = [
        "后台改规则：价格、会员、优惠、密码。",
        "终端做执行：拉内容、装游戏、展示支付和开局引导。",
        "两边要联调：后台改完规则后，终端要重新同步或做实际下单测试。",
    ]
    for idx, tip in enumerate(tips):
        draw.text((280, 575 + idx * 32), tip, font=tip_font, fill="#334155")
    img.save(path)


def make_player_flow(path: Path) -> None:
    img = Image.new("RGB", (1600, 840), "#FFFDF8")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "玩家体验主流程", "从走近终端到完成体验的现场路径")
    items = [
        ("1 浏览游戏", "看封面、时长、价格和人数说明", 80, 240),
        ("2 进入详情页", "决定玩什么、玩多久、几个人玩", 380, 240),
        ("3 进入支付", "散客扫码 or 会员识别", 680, 240),
        ("4 佩戴头显", "按照编号和指引佩戴设备", 980, 240),
        ("5 开始游玩", "终端进入等待或开局状态", 1280, 240),
        ("6 结束返回首页", "设备恢复空闲，可继续复玩", 640, 520),
    ]
    for title, desc, x, y in items:
        draw_box(draw, x, y, 240, 120, title, desc, "#FFFFFF", "#FBBF24" if x > 1200 else "#93C5FD")
    for sx, ex in [(320, 380), (620, 680), (920, 980), (1220, 1280)]:
        draw_arrow(draw, (sx, 300), (ex, 300))
    draw_arrow(draw, (1400, 360), (1400, 470))
    draw_arrow(draw, (1400, 470), (760, 470))
    draw_arrow(draw, (760, 470), (760, 520))
    img.save(path)


def make_player_payment_map(path: Path) -> None:
    img = Image.new("RGB", (1600, 860), "white")
    draw = ImageDraw.Draw(img)
    draw_title(draw, "玩家支付路径示意", "散客与会员进入的是不同支付确认方式")
    draw_box(draw, 620, 130, 360, 110, "选好游戏并点击开始", "终端根据玩家身份进入不同支付分支", "#EFF6FF", "#60A5FA")
    draw_arrow(draw, (800, 240), (430, 340))
    draw_arrow(draw, (800, 240), (800, 340))
    draw_arrow(draw, (800, 240), (1170, 340))
    draw_box(draw, 180, 340, 500, 180, "散客直接扫码", "终端展示微信 / 支付宝二维码\n玩家直接用手机扫一扫完成付款\n支付完成后进入设备引导", "#FEF3C7", "#F59E0B")
    draw_box(draw, 550, 340, 500, 180, "会员主动扫码", "会员主动扫描终端码\n在小程序查看优惠、余额、游戏币和补差\n确认后完成支付", "#ECFDF5", "#34D399")
    draw_box(draw, 920, 340, 500, 180, "会员码反扫 / 店员协助", "玩家出示会员码由终端识别\n或店员手机端代确认\n支付成功后进入设备引导", "#EFF6FF", "#60A5FA")
    draw.rounded_rectangle([250, 620, 1350, 760], radius=20, fill="#F8FAFC", outline="#CBD5E1", width=2)
    font = load_font(24)
    lines = [
        "共同结果：支付成功后，终端才会分配设备编号并提示佩戴头显。",
        "如果支付成功但没有开局，应立即联系门店，由店员在点播系统订单中处理补开局或退款。",
    ]
    for i, line in enumerate(lines):
        draw.text((300, 665 + i * 34), line, font=font, fill="#334155")
    img.save(path)


def apply_font(run, size: float = 11, *, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = DOC_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.font.size = Pt(11)

    for name, size, color in [
        ("Title", 22, "0F172A"),
        ("Heading 1", 18, "1E3A8A"),
        ("Heading 2", 15, "0F766E"),
        ("Heading 3", 12.5, "334155"),
    ]:
        style = doc.styles[name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    apply_font(r, 22, bold=True, color="0F172A")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitle)
    apply_font(r2, 11.5, color="64748B")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for k, v in meta:
        row = table.rows[0] if table.rows[0].cells[0].text == "" and table.rows[0].cells[1].text == "" else table.add_row()
        row.cells[0].text = k
        row.cells[1].text = v
    style_table(table, header=False)


def style_table(table, *, header: bool = True, header_fill: str = "DBEAFE") -> None:
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if header and ridx == 0:
                set_cell_shading(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(3)
                for run in p.runs:
                    apply_font(run, 10.5, bold=(header and ridx == 0))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    apply_font(r, 18 if level == 1 else 15 if level == 2 else 12.5, bold=True, color="1E3A8A" if level == 1 else "0F766E" if level == 2 else "334155")


def add_para(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        r0 = p.add_run(bold_prefix)
        apply_font(r0, 11, bold=True, color="0F172A")
    r = p.add_run(text if not bold_prefix else "")
    if not bold_prefix:
        apply_font(r, 11, color="334155")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        tail = p.add_run(text)
        apply_font(tail, 11, color="334155")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        apply_font(r, 11, color="334155")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        apply_font(r, 11, color="334155")


def add_image(doc: Document, path: Path, width_cm: float = 16.0) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))


def add_simple_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def build_official_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 · 官方运营使用教程",
        "可编辑版说明文档｜适用角色：平台超管、运营交付、财务配置人员",
        [
            ("适用版本", "v2.19（基于仓库资料整理）"),
            ("覆盖系统", "总运营后台 admin-dashboard Platform"),
            ("文档目标", "指导官方从新建商家到门店具备营业条件的完整交付流程"),
        ],
    )
    add_para(doc, "这份文档面向官方运营和交付人员。需要特别强调的是：官方工作里至少存在两条并行链路，一条是商家/门店开通交付线，另一条是平台内容管理线；两条线最终在“内容分发到门店”这个节点汇合。")
    add_image(doc, ASSET_DIR / "official_flow.png")
    add_heading(doc, "一、开通前要先确认什么", 1)
    add_bullets(doc, [
        "商家主体信息是否齐全：名称、联系人、联系电话、负责区域、代理归属。",
        "门店现场设备是否已经统计清楚：主机数量、头显数量、设备型号和主机编号规则。",
        "计划上线的游戏清单是否明确：哪些是主机游戏，哪些是一体机游戏，哪些内容需要现场 USB 安装。",
        "门店是否需要同步完成收款测试：如需要，支付配置和测试联系人要提前准备。",
    ])
    add_heading(doc, "二、商家开通交付线", 1)
    add_heading(doc, "1. 新建商家", 2)
    add_para(doc, "操作路径：店铺管理 → 商家管理 → 新增商家")
    add_numbers(doc, [
        "填写商家名称、联系人、联系电话、负责区域、代理商和提现手续费率。",
        "创建商家后台管理员账号和密码，这套账号会交付给商家作为正式登录凭证。",
        "补充提现账户信息；如果现场暂时拿不到银行卡信息，可以先建档，后续再补。",
    ])
    add_simple_table(doc, ["字段", "建议口径"], [
        ["商家名称", "与合同或营业执照保持一致"],
        ["管理员账号", "便于识别商家主体，避免多店混淆"],
        ["商家状态", "交付前保持“正常”"],
    ])
    add_heading(doc, "2. 新建店铺", 2)
    add_para(doc, "操作路径：店铺管理 → 店铺列表 → 新增店铺")
    add_bullets(doc, [
        "必须先选对所属商家，再填写店铺名称、所属区域、详细地址、联系电话和营业状态。",
        "店铺创建后会生成注册码，后续玩家到店可通过二维码注册会员。",
        "如果门店要做新会员礼包，也建议在这一阶段同步确认注册送预存款或游戏币规则。",
    ])
    add_heading(doc, "三、平台内容管理线", 1)
    add_para(doc, "这一部分不是商家开通线的串行步骤，而是平台内容管理的独立前置工作。只要门店需要用到某个内容，平台侧就需要先把该内容准备好，后续才能在“内容分发”时下发到门店。")
    add_heading(doc, "1. 录入游戏并准备上架", 2)
    add_para(doc, "操作路径：内容中心 → 游戏库 → 添加游戏")
    add_bullets(doc, [
        "游戏名称、封面、介绍、题材要与终端展示口径一致，避免玩家选购时理解偏差。",
        "一定要区分运行平台：主机游戏和一体机游戏的现场安装动作不同。",
        "体验时长、游戏豆消耗和推荐售价要跟商务方案保持一致，避免门店不知道如何定价。",
    ])
    add_heading(doc, "2. 维护版本、状态与运营参数", 2)
    add_bullets(doc, [
        "确认内容当前是草稿、待审核、已上线还是已下架，避免门店分发时看不到内容。",
        "明确版本、时长、游戏豆消耗和推荐售价等运营参数，确保门店理解成本口径。",
        "如果游戏来自 CP 供应商后台，先走审核再进入门店分发，不要和商家开通交付混在一起处理。",
    ])
    add_heading(doc, "四、设备与支付交付线", 1)
    add_heading(doc, "1. 录入主机并生成 Token", 2)
    add_para(doc, "操作路径：数据中心 → 设备配置管理 → 主机管理")
    add_numbers(doc, [
        "录入主机编号、主机名称、设备类型、硬件配置、系统版本和 MAC 地址。",
        "把主机分配到正确的商家和店铺。",
        "分配完成后生成 Token，并交付给门店在点播系统中填写。",
    ])
    add_heading(doc, "关键原则", 3)
    add_bullets(doc, [
        "一台主机一个 Token，不建议多个主机共用同一个 Token。",
        "主机取消分配后，原 Token 应视为失效，重新分配时要重新生成或重新交付。",
        "MAC 地址是现场识别主机的重要依据，建议录入时认真核对。",
    ])
    add_heading(doc, "2. 录入头显并绑定主机", 2)
    add_para(doc, "操作路径：数据中心 → 设备配置管理 → 头显管理")
    add_bullets(doc, [
        "录入头显名称、型号、SN 码和固件版本。",
        "先把头显分配到正确门店，再绑定到同店铺下的在线主机。",
        "绑定时优先确认物理编号，避免后台绑定关系与现场设备摆放不一致。",
    ])
    add_image(doc, ASSET_DIR / "official_device_map.png")
    add_heading(doc, "五、两条线的汇合点：内容分发与支付配置", 1)
    add_heading(doc, "1. 内容分发", 2)
    add_para(doc, "操作路径：内容中心 → 内容分发")
    add_bullets(doc, [
        "内容分发是商家开通线与内容管理线真正汇合的节点：门店已经准备好，内容也已经准备好，才能下发到门店。",
        "智能分发适合日常更新或补发，速度更快。",
        "全量分发适合首次部署、版本修复或大范围重装。",
        "不要只看任务已提交，要确认门店端实际已经能下载或看到内容。",
    ])
    add_heading(doc, "2. 拉卡拉配置", 2)
    add_para(doc, "操作路径：平台财务 → 拉卡拉配置")
    add_bullets(doc, [
        "完成商户号、终端号、机构号和证书等基础配置后，再安排门店测试微信 / 支付宝收款。",
        "如门店涉及代理商或 CP 结算，也要同步检查结算配置是否齐全。",
    ])
    add_heading(doc, "六、最终交付检查", 1)
    add_numbers(doc, [
        "商家后台账号已经交付给店长或负责人。",
        "主机和头显的后台归属与现场编号一致。",
        "门店拿到了 Token，并知道在哪里填写。",
        "目标游戏已经完成分发，门店知道哪些内容是平台已下发、哪些内容仍需现场 USB 安装。",
        "支付已经做过测试，门店知道异常联系人是谁。",
    ])
    doc.save(OFFICIAL_DOC)
    return OFFICIAL_DOC


def build_merchant_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 · 商家门店使用教程",
        "可编辑版说明文档｜适用角色：老板、店长、店员、门店管理员",
        [
            ("适用版本", "v2.19（基于仓库资料整理）"),
            ("覆盖系统", "商户后台 admin-dashboard Shop + PC 点播系统管理模式"),
            ("文档目标", "帮助门店完成开业准备、设备核对、价格配置和终端上线"),
        ],
    )
    add_para(doc, "这份文档从门店视角出发，重点讲“官方交接以后，门店还需要自己做什么”，以及后台与点播系统之间怎样配合，才能真正开门营业。")
    add_image(doc, ASSET_DIR / "merchant_prep_flow.png")
    add_heading(doc, "一、接收官方交付后的第一步", 1)
    add_bullets(doc, [
        "先确认是否已经拿到后台账号、门店信息、主机编号、头显清单和 Token。",
        "让现场设备和后台设备列表一一对应，避免后面改密码或绑定时找错机器。",
        "确认官方已经告诉你：哪些游戏已经分发，哪些内容需要在现场继续安装。",
    ])
    add_heading(doc, "二、登录商户后台并核对设备", 1)
    add_para(doc, "操作路径：系统设置 → 设备列表")
    add_heading(doc, "1. 看主机设备", 2)
    add_bullets(doc, [
        "重点看主机编号、MAC 地址、在线状态和绑定头显数量。",
        "如果主机显示已经分配，但现场找不到对应主机，先不要继续装游戏，要先把设备清单核对清楚。",
    ])
    add_heading(doc, "2. 看头显设备", 2)
    add_bullets(doc, [
        "重点看 SN 码、型号、电量、绑定主机和使用状态。",
        "如果官方还没绑好，门店可以在后台补做同店铺下的头显绑定。",
        "如果系统提示头显和主机不在同一门店，先回官方后台修正归属，再继续操作。",
    ])
    add_heading(doc, "三、设置点播系统密码", 1)
    add_para(doc, "操作路径：系统设置 → 设备列表 → 主机设备 → 修改点播系统密码")
    add_bullets(doc, [
        "点播系统密码是给店员进入管理模式用的，不建议留空。",
        "建议每台主机都设置独立密码，并记录在交接班表上。",
        "如果密码忘了，由有权限的门店管理员在后台重新设置。",
    ])
    add_heading(doc, "四、配置门店售卖规则", 1)
    add_heading(doc, "1. 配价格", 2)
    add_para(doc, "操作路径：商品管理 → 单次消费项目")
    add_bullets(doc, [
        "建议按 30 分钟、60 分钟、单局体验等维度配置价格。",
        "终端展示的时长与后台定价要一致，避免顾客看到的和店员收的不是一回事。",
    ])
    add_heading(doc, "2. 配充值套餐、会员和优惠券", 2)
    add_bullets(doc, [
        "在运营管理中配置充值套餐，例如充 100 送 10、充 300 送 50。",
        "在会员级别里配置普通会员、金卡、钻石等折扣权益。",
        "在优惠券里配置满减券、代金券、折扣券或体验券，方便开业拉新和复购促销。",
    ])
    add_simple_table(doc, ["配置项", "推荐做法"], [
        ["单次消费项目", "按时长或单局体验维度设置售价"],
        ["充值套餐", "用少量常用档位，不要第一次就配太复杂"],
        ["会员级别", "至少先有普通会员和一个高阶会员折扣"],
        ["优惠券", "先配一到两个常用活动，不要一次上太多规则"],
    ])
    add_image(doc, ASSET_DIR / "merchant_system_map.png")
    add_heading(doc, "五、在点播系统里完成上线准备", 1)
    add_heading(doc, "1. 开机前准备", 2)
    add_bullets(doc, [
        "确认主机已经在官方后台完成分配，并拿到了专属 Token。",
        "确认现场网络正常，磁盘空间足够；大文件分发前最好先测下载速度。",
        "如有一体机内容需要安装，提前准备好 USB 3.0 数据线。",
    ])
    add_heading(doc, "2. 写入 Token", 2)
    add_numbers(doc, [
        "打开点播系统安装包或登录页，进入“基础配置 / 系统设置 / Token 设置”入口。",
        "把官方下发的 Token 填进去并保存。",
        "保存后重新加载终端，让系统完成门店绑定和内容同步。",
    ])
    add_para(doc, "提示：有些版本在登录页右上角“更多菜单”里进入基础配置，有些版本要先进入管理模式再改。")
    add_heading(doc, "3. 同步并安装游戏", 2)
    add_bullets(doc, [
        "Token 生效后，终端应该能看到当前门店已分发的游戏列表。",
        "主机游戏通常直接点“安装到主机”；一体机游戏需要先接入头显，再点“安装到一体机 / 头显”。",
        "营业前至少把主推内容装好，并随机抽查 1 到 2 款能正常进入启动流程。",
    ])
    add_heading(doc, "六、开业前最后巡检", 1)
    add_numbers(doc, [
        "后台设备列表与现场设备编号一致。",
        "点播系统密码已经设置并告知值班店员。",
        "主推游戏已经安装完成。",
        "测试过一笔支付和一笔开局，确认订单能在后台查到。",
        "头显电量、佩戴指引和门店注册码都已经准备好。",
    ])
    doc.save(MERCHANT_DOC)
    return MERCHANT_DOC


def build_player_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 · 玩家体验使用教程",
        "可编辑版说明文档｜适用角色：到店体验 VR 的顾客",
        [
            ("适用版本", "v2.19（基于仓库资料整理）"),
            ("覆盖系统", "PC 点播系统 + 微信小程序"),
            ("文档目标", "帮助玩家理解选游戏、支付、佩戴设备和结束体验的完整过程"),
        ],
    )
    add_para(doc, "这份文档不是后台操作手册，而是站在顾客视角，解释门店里的大屏点播终端是怎么工作的，玩家到店后应该怎么完成体验。")
    add_image(doc, ASSET_DIR / "player_flow.png")
    add_heading(doc, "一、玩家到店后会经历什么", 1)
    add_bullets(doc, [
        "先在终端上浏览可玩的游戏，看封面、时长、价格和支持人数。",
        "选定内容后进入支付环节，散客和会员的支付路径不同。",
        "支付成功后终端会提示分配到的头显编号，玩家按引导佩戴设备并开始游戏。",
        "游戏结束后终端回到首页，设备恢复空闲，玩家可以选择复玩或注册会员。",
    ])
    add_heading(doc, "二、如何选游戏", 1)
    add_bullets(doc, [
        "首页通常会展示推荐内容、游戏分类、搜索入口和游戏卡片。",
        "点击游戏卡片后，可查看玩法介绍、时长、价格和适合几个人玩。",
        "如果是多人游戏，终端或店员会在支付前确认需要几台设备。",
    ])
    add_image(doc, ASSET_DIR / "player_payment_map.png")
    add_heading(doc, "三、支付方式怎么走", 1)
    add_heading(doc, "1. 散客支付", 2)
    add_bullets(doc, [
        "散客不需要先注册会员。",
        "终端会展示微信或支付宝二维码，直接用手机扫一扫完成付款。",
        "支付成功后终端自动进入设备引导页。",
    ])
    add_heading(doc, "2. 会员支付", 2)
    add_bullets(doc, [
        "会员可以主动扫描终端上的支付码，在小程序中查看优惠、余额、游戏币和补差金额。",
        "也可以出示会员码，让终端识别后再到小程序里确认支付。",
        "如果是团客、测试或特殊场景，店员也可能用手机代玩家完成确认。",
    ])
    add_simple_table(doc, ["玩家身份", "常见支付方式", "说明"], [
        ["散客", "微信支付 / 支付宝支付", "终端展示二维码后直接扫码付款"],
        ["会员主动扫码", "预存款、游戏币、微信补差", "在小程序中查看明细后确认支付"],
        ["会员码反扫", "预存款、游戏币、微信补差", "先识别会员身份，再确认扣款"],
    ])
    add_heading(doc, "四、支付后怎么开始玩", 1)
    add_bullets(doc, [
        "终端会提示对应的头显编号，或由店员直接把设备递给玩家。",
        "按图示把头显戴好，调整绑带、瞳距和清晰度；有不舒服就立刻找店员协助。",
        "很多门店采用“终端选游戏，头显直接开局”的方式，头显内部不一定会显示复杂菜单。",
    ])
    add_heading(doc, "五、结束、复玩和会员注册", 1)
    add_bullets(doc, [
        "游戏结束后，终端会回到首页，可继续体验其他内容。",
        "如果想继续玩，可以重新选游戏；支持续费的内容也可能允许追加时长。",
        "如果想成为会员，可扫描门店注册码或打开头号空间小程序完成注册，领取新会员礼包。",
    ])
    add_heading(doc, "六、什么时候应该找店员", 1)
    add_bullets(doc, [
        "支付成功但终端没有进入佩戴引导。",
        "头显戴上后看不清、太松、太紧或身体不适。",
        "游戏中途异常退出、没有声音、没有画面或无法继续。",
        "想退款、补开局、换内容或需要多人协助体验。",
    ])
    doc.save(PLAYER_DOC)
    return PLAYER_DOC


def build_all() -> list[Path]:
    ensure_dirs()
    make_official_flow(ASSET_DIR / "official_flow.png")
    make_official_device_map(ASSET_DIR / "official_device_map.png")
    make_merchant_prep_flow(ASSET_DIR / "merchant_prep_flow.png")
    make_merchant_system_map(ASSET_DIR / "merchant_system_map.png")
    make_player_flow(ASSET_DIR / "player_flow.png")
    make_player_payment_map(ASSET_DIR / "player_payment_map.png")
    return [
        build_official_doc(),
        build_merchant_doc(),
        build_player_doc(),
    ]


if __name__ == "__main__":
    for path in build_all():
        print(path)
