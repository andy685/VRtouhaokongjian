from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "user-payment-prd-assets"
DOCX_PATH = OUT_DIR / "用户支付PRD_玩家玩游戏支付流程.docx"

FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
PAGE_SIZE = (1600, 2260)
PAGE_WIDTH, PAGE_HEIGHT = PAGE_SIZE
MARGIN = 96
CONTENT_W = PAGE_WIDTH - MARGIN * 2


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(FONT_PATH, size=size)
    except Exception:
        return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> str:
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            bbox = draw.textbbox((0, 0), trial, font=font)
            if bbox[2] - bbox[0] > max_width and current:
                lines.append(current)
                current = ch
            else:
                current = trial
        if current:
            lines.append(current)
    return "\n".join(lines)


def draw_heading(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, size: int = 36, color: str = "#1D4ED8") -> int:
    font = load_font(size)
    draw.text((x, y), text, font=font, fill=color)
    bbox = draw.textbbox((x, y), text, font=font)
    return bbox[3] + 8


def draw_para(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, width: int = CONTENT_W, size: int = 24, color: str = "#334155", spacing: int = 10) -> int:
    font = load_font(size)
    wrapped = wrap_text(draw, text, font, width)
    draw.multiline_text((x, y), wrapped, font=font, fill=color, spacing=spacing)
    bbox = draw.multiline_textbbox((x, y), wrapped, font=font, spacing=spacing)
    return bbox[3] + 8


def draw_bullets(draw: ImageDraw.ImageDraw, x: int, y: int, bullets: list[str], width: int = CONTENT_W, size: int = 22) -> int:
    for bullet in bullets:
        text_y = y
        y = draw_para(draw, x + 28, y, bullet, width - 28, size=size)
        draw.ellipse([x, text_y + 12, x + 10, text_y + 22], fill="#0F172A")
        y += 6
    return y


def draw_table(
    draw: ImageDraw.ImageDraw,
    x: int,
    y: int,
    width: int,
    headers: list[str],
    rows: list[list[str]],
    ratios: list[float],
    font_size: int = 18,
    header_fill: str = "#DBEAFE",
) -> int:
    body_font = load_font(font_size)
    head_font = load_font(font_size + 1)
    col_widths = [int(width * r) for r in ratios]
    col_widths[-1] += width - sum(col_widths)

    def row_height(cells: list[str], font) -> int:
        heights = []
        for cell, cw in zip(cells, col_widths):
            wrapped = wrap_text(draw, cell, font, cw - 24)
            bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6)
            heights.append((bbox[3] - bbox[1]) + 26)
        return max(heights)

    current_y = y
    height = row_height(headers, head_font)
    cx = x
    for header, cw in zip(headers, col_widths):
        draw.rectangle([cx, current_y, cx + cw, current_y + height], fill=header_fill, outline="#475569", width=2)
        wrapped = wrap_text(draw, header, head_font, cw - 24)
        bbox = draw.multiline_textbbox((0, 0), wrapped, font=head_font, spacing=6)
        draw.multiline_text((cx + 12, current_y + (height - (bbox[3] - bbox[1])) / 2), wrapped, font=head_font, fill="#0F172A", spacing=6)
        cx += cw
    current_y += height

    for row in rows:
        height = row_height(row, body_font)
        cx = x
        for cell, cw in zip(row, col_widths):
            draw.rectangle([cx, current_y, cx + cw, current_y + height], fill="white", outline="#64748B", width=1)
            wrapped = wrap_text(draw, cell, body_font, cw - 24)
            bbox = draw.multiline_textbbox((0, 0), wrapped, font=body_font, spacing=6)
            draw.multiline_text((cx + 12, current_y + (height - (bbox[3] - bbox[1])) / 2), wrapped, font=body_font, fill="#334155", spacing=6)
            cx += cw
        current_y += height
    return current_y + 4


def draw_box(draw: ImageDraw.ImageDraw, x: int, y: int, w: int, h: int, text: str, fill: str, outline: str) -> None:
    font = load_font(22)
    draw.rounded_rectangle([x, y, x + w, y + h], radius=18, fill=fill, outline=outline, width=3)
    wrapped = wrap_text(draw, text, font, w - 24)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6)
    tx = x + (w - (bbox[2] - bbox[0])) / 2
    ty = y + (h - (bbox[3] - bbox[1])) / 2
    draw.multiline_text((tx, ty), wrapped, font=font, fill="#0F172A", spacing=6, align="center")


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str | None = None) -> None:
    color = "#2563EB"
    draw.line([start, end], fill=color, width=4)
    if start[0] == end[0]:
        direction = 1 if end[1] > start[1] else -1
        p1 = (end[0] - 8, end[1] - 12 * direction)
        p2 = (end[0] + 8, end[1] - 12 * direction)
    else:
        direction = 1 if end[0] > start[0] else -1
        p1 = (end[0] - 12 * direction, end[1] - 8)
        p2 = (end[0] - 12 * direction, end[1] + 8)
    draw.polygon([end, p1, p2], fill=color)
    if label:
        font = load_font(18)
        bbox = draw.textbbox((0, 0), label, font=font)
        lx = (start[0] + end[0]) / 2 - (bbox[2] - bbox[0]) / 2
        ly = (start[1] + end[1]) / 2 - 28
        draw.rounded_rectangle([lx - 8, ly - 4, lx + (bbox[2] - bbox[0]) + 8, ly + (bbox[3] - bbox[1]) + 4], radius=8, fill="#DBEAFE", outline="#93C5FD")
        draw.text((lx, ly), label, font=font, fill="#1E3A8A")


def draw_overall_flow(path: Path) -> None:
    img = Image.new("RGB", (1440, 900), "#F8FBFF")
    draw = ImageDraw.Draw(img)
    draw_heading(draw, 40, 28, "玩家玩游戏支付主流程", 34, "#0F172A")
    top = [
        ("1. 到达设备\n选择内容 / 设备", 50),
        ("2. 触发支付\n主动扫码 / 反扫 / 店员扫码", 285),
        ("3. 展示支付明细\n金额、优惠、支付资产", 570),
        ("4. 确认支付\n扣款 / 外部支付", 865),
        ("5. 订单创建\n写入支付流水", 1120),
    ]
    for text, x in top:
        draw_box(draw, x, 170, 210, 110, text, "#EFF6FF", "#2563EB")
    for start_x in [260, 545, 830, 1080]:
        draw_arrow(draw, (start_x, 225), (start_x + 25, 225))
    draw_arrow(draw, (1225, 280), (1225, 350))
    draw_arrow(draw, (1225, 350), (235, 350), "支付成功后进入设备控制")
    draw_arrow(draw, (235, 350), (235, 470))
    bottom = [
        ("6. 下发开局指令", 160, "#ECFDF5", "#10B981"),
        ("7. 设备开局成功\n进入游玩中", 430, "#ECFDF5", "#10B981"),
        ("8. 游玩结束\n回传结束原因", 735, "#EFF6FF", "#2563EB"),
        ("9. 正常完成 / 异常闭环\n补开局 / 退款 / 复核", 1015, "#FEF3C7", "#F59E0B"),
    ]
    for text, x, fill, outline in bottom:
        draw_box(draw, x, 470, 220 if x != 1015 else 280, 110, text, fill, outline)
    draw_arrow(draw, (380, 525), (430, 525))
    draw_arrow(draw, (650, 525), (735, 525))
    draw_arrow(draw, (955, 525), (1015, 525))
    draw_para(draw, 50, 700, "说明：\n• 支付环节以“支付成功、开局成功、结束回传、订单归档”为主链路。\n• 任一环节失败都需要有异常订单记录，并能进入补开局、人工复核或退款流程。", 1320, 20, "#475569", 8)
    img.save(path)


def draw_reverse_flow(path: Path) -> None:
    img = Image.new("RGB", (1440, 980), "#FFFDF8")
    draw = ImageDraw.Draw(img)
    draw_heading(draw, 40, 28, "会员码反扫支付闭环（重点流程）", 34, "#0F172A")
    left = [
        "A1. 玩家出示会员码\n设备端 / 扫码枪完成反扫",
        "A2. 系统识别会员身份\n校验会员状态、资产、套餐、次数",
        "A3. 小程序展示支付明细\n展示扣减资产、补差金额、优惠信息",
        "A4. 玩家确认扣款\n进入支付处理中",
        "A5. 扣款成功\n生成支付流水与点播订单",
        "A6. 系统下发开局指令\n设备回传开局结果",
    ]
    right = [
        "B1. 开局成功\n订单状态=进行中，进入游玩阶段",
        "B2. 正常结束\n回传结束原因=正常完成",
        "B3. 订单完结\n记录时长、内容、支付方式、设备",
        "B4. 异常场景一：支付成功未开局\n自动转异常订单，待补开局或退款",
        "B5. 异常场景二：中途退出 / 人工强停\n根据规则处理，不退款或进入人工复核",
        "B6. 退款闭环\n未结算走原路退回；已结算走线下退款+凭证",
    ]
    y = 120
    for text in left:
        draw_box(draw, 70, y, 520, 88, text, "#EFF6FF", "#2563EB")
        y += 118
    y = 120
    for idx, text in enumerate(right):
        fill = "#ECFDF5" if idx < 3 else "#FEF3C7"
        outline = "#10B981" if idx < 3 else "#F59E0B"
        draw_box(draw, 820, y, 550, 88, text, fill, outline)
        y += 118
    for i in range(5):
        yy = 120 + i * 118 + 88
        draw_arrow(draw, (330, yy), (330, yy + 30))
        draw_arrow(draw, (1095, yy), (1095, yy + 30))
    draw_arrow(draw, (590, 605), (820, 165), "开局成功")
    draw_arrow(draw, (590, 487), (820, 515), "扣款成功但开局失败")
    draw_para(draw, 70, 840, "闭环要求：\n1. 扫码成功不等于支付成功，必须经过会员识别、资产校验、确认扣款。\n2. 扣款成功不等于业务完成，必须继续校验开局结果。\n3. 结束时必须有结束原因；异常结束要进入异常订单或退款流程。\n4. 退款必须能回溯到订单、支付方式、设备、操作人和退款凭证。", 1300, 20, "#475569", 8)
    img.save(path)


def draw_architecture(path: Path) -> None:
    img = Image.new("RGB", (1440, 900), "#F8FAFC")
    draw = ImageDraw.Draw(img)
    draw_heading(draw, 40, 28, "支付与点播能力工程结构示意", 34, "#0F172A")
    draw_box(draw, 70, 120, 260, 100, "用户侧\n微信小程序 / H5", "#EFF6FF", "#2563EB")
    draw_box(draw, 420, 120, 320, 100, "门店侧\n点播终端 / 扫码枪 / 头显主机", "#F0FDF4", "#10B981")
    draw_box(draw, 850, 120, 280, 100, "后台运营侧\n商家后台 / 总运营后台", "#FEF3C7", "#F59E0B")
    draw_box(draw, 90, 380, 230, 100, "订单服务\n创建订单 / 状态流转 / 结算", "#FFFFFF", "#2563EB")
    draw_box(draw, 380, 380, 250, 100, "支付服务\n金额拆分 / 扣款 / 第三方支付", "#FFFFFF", "#2563EB")
    draw_box(draw, 690, 380, 260, 100, "会员资产服务\n预存款 / 游戏币 / 套票 / 次数", "#FFFFFF", "#2563EB")
    draw_box(draw, 1020, 380, 250, 100, "设备编排服务\n下发开局 / 结束回传 / 心跳", "#FFFFFF", "#2563EB")
    draw_box(draw, 270, 640, 330, 95, "异常订单中心\n支付异常 / 开局异常 / 人工复核", "#FEF2F2", "#DC2626")
    draw_box(draw, 760, 640, 320, 95, "退款中心\n原路退回 / 线下退款 / 凭证留存", "#FFF7ED", "#EA580C")
    draw_arrow(draw, (200, 220), (200, 380))
    draw_arrow(draw, (580, 220), (505, 380))
    draw_arrow(draw, (990, 220), (1145, 380))
    draw_arrow(draw, (320, 430), (380, 430))
    draw_arrow(draw, (630, 430), (690, 430))
    draw_arrow(draw, (950, 430), (1020, 430))
    draw_arrow(draw, (505, 480), (430, 640), "异常回流")
    draw_arrow(draw, (1145, 480), (920, 640), "退款 / 补偿")
    draw_para(draw, 70, 790, "设计原则：支付、订单、会员资产、设备控制四个能力必须解耦，但要通过订单号贯穿全链路，保证支付成功、开局、结束、异常、退款可追踪。", 1300, 20, "#334155")
    img.save(path)


def compose_pages(overall: Path, reverse: Path, architecture: Path) -> list[Path]:
    pages: list[Path] = []

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 110
    y = draw_heading(draw, MARGIN, y, "用户支付 PRD", 54, "#0F172A")
    y = draw_heading(draw, MARGIN, y + 8, "玩家玩游戏支付流程", 34, "#1E3A8A")
    y = draw_para(draw, MARGIN, y + 20, "文档目标：梳理玩家从发起支付到游玩结束的完整链路，覆盖支付方式、反扫闭环、异常处理、退款闭环和工程结构边界。", size=24)
    y = draw_bullets(draw, MARGIN, y + 10, [
        "重点场景：小程序主动扫码、会员码反扫、店员扫码点播。",
        "重点能力：扫码成功、扣款、开局、结束、异常、退款必须闭环。",
        "目标读者：产品、研发、测试、运营、设备联调人员。"
    ])
    y = draw_heading(draw, MARGIN, y + 20, "一、玩家玩游戏支付主流程", 30, "#2563EB")
    y = draw_para(draw, MARGIN, y + 6, "所有支付场景都应落到统一主链路：触发支付 → 明细确认 → 扣款成功 → 创单 → 开局 → 游玩结束 → 订单归档 / 异常处置。", size=22)
    flow_img = Image.open(overall).convert("RGB")
    flow_img.thumbnail((CONTENT_W, 900))
    img.paste(flow_img, (MARGIN, y + 12))
    page = ASSET_DIR / "prd-page-1.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "二、支付触发方式与可用支付方式", 32, "#2563EB")
    headers = ["触发方式", "前端载体", "适用对象", "可用支付方式", "支付成功后动作", "退款口径"]
    rows = [
        ["小程序主动扫码", "微信小程序", "散客 / 会员", "微信支付、预存款、混合支付（预存款+微信）、可扩展游戏币 / 套票", "创建点播订单并开局", "未开局可退款；正常完成不退款"],
        ["会员码反扫", "会员码 + 小程序确认", "会员", "预存款、会员余额、游戏币、预存次数、套票抵扣、资产不足时补差支付", "确认扣款后创单并自动开局", "支付成功未开局可退款；中途退出不退款；设备异常可复核 / 退款"],
        ["店员扫码点播", "店员 H5", "会员 / 团体客", "预存款、游戏币、套票、预存次数、微信补差", "店员代操作创建订单并开局", "按订单是否已结算决定原路退回或线下退款"],
    ]
    y = draw_table(draw, MARGIN, y + 16, CONTENT_W, headers, rows, [0.12, 0.11, 0.12, 0.27, 0.18, 0.20], 18)
    y = draw_heading(draw, MARGIN, y + 18, "三、支付方式定义", 30, "#2563EB")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "微信支付：玩家直接支付现金类金额，适用于主动扫码和补差支付。",
        "预存款 / 会员余额：直接扣减会员账户余额，不产生外部现金流水。",
        "游戏币：适用于按游戏币计价的内容或活动。",
        "预存次数：按次扣减，适合固定项目包次消费。",
        "套票抵扣：优先消耗玩家已购买套餐权益。",
        "混合支付：会员资产不足时，先扣会员资产，再补差微信支付。"
    ])
    y = draw_heading(draw, MARGIN, y + 8, "四、推荐订单状态机", 30, "#2563EB")
    state_headers = ["状态", "进入条件", "退出条件", "备注"]
    state_rows = [
        ["待确认支付", "扫码成功并识别到有效支付对象", "玩家确认支付 / 取消", "尚未扣款"],
        ["支付中", "已发起扣款或第三方支付", "支付成功 / 支付失败 / 超时", "需防重复提交"],
        ["待开局", "支付成功且订单创建完成", "设备开局成功 / 开局失败", "关键闭环节点"],
        ["进行中", "设备回传开局成功", "结束回传 / 人工强停", "记录过程状态"],
        ["已完成", "正常结束并回传时长", "-", "正常闭环"],
        ["异常待处理", "开局失败、结束异常、支付状态不一致", "补开局 / 退款 / 人工结案", "需运营介入"],
        ["退款中", "已发起退款请求", "退款成功 / 退款失败", "支持异步回调"],
        ["已退款", "退款成功", "-", "形成退款闭环"],
    ]
    draw_table(draw, MARGIN, y + 12, CONTENT_W, state_headers, state_rows, [0.16, 0.28, 0.28, 0.28], 19, "#DBEAFE")
    page = ASSET_DIR / "prd-page-2.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "五、反扫支付闭环（重点）", 32, "#2563EB")
    y = draw_para(draw, MARGIN, y + 8, "反扫支付不是“扫到码就结束”，而是“识别会员 → 确认明细 → 扣款成功 → 开局成功 → 结束归档 → 异常可退款”的完整闭环。", size=22)
    reverse_img = Image.open(reverse).convert("RGB")
    reverse_img.thumbnail((CONTENT_W, 900))
    img.paste(reverse_img, (MARGIN, y + 12))
    y = y + 12 + reverse_img.height + 20
    y = draw_heading(draw, MARGIN, y, "5.1 反扫详细步骤", 28, "#0F766E")
    reverse_headers = ["阶段", "动作发起方", "系统动作", "用户可见内容", "产出结果"]
    reverse_rows = [
        ["1", "玩家 / 扫码枪", "扫码枪读取会员码，设备端解析会员身份", "设备提示“扫码中 / 识别中”", "进入会员识别"],
        ["2", "系统", "校验会员状态、会员资产、套票、次数、设备可用性", "若可用，则跳转小程序确认明细", "形成待确认支付明细"],
        ["3", "小程序", "展示内容、金额、资产抵扣、补差支付、优惠信息", "玩家可查看扣款明细并确认", "形成支付确认动作"],
        ["4", "支付服务", "扣减预存款 / 游戏币 / 次数 / 套票，或发起微信补差", "支付处理中、禁止重复提交", "支付成功 / 失败"],
        ["5", "订单服务", "支付成功后创建点播订单并写入支付流水", "用户可看到支付成功状态", "订单状态=待开局"],
        ["6", "设备编排", "向设备下发开局指令，等待设备回传", "前端显示“准备开局 / 正在进入游戏”", "开局成功或开局失败"],
        ["7", "设备", "游玩中回传状态，结束时回传结束原因与时长", "玩家完成体验或收到异常提示", "订单状态=已完成 / 异常"],
        ["8", "退款中心", "若支付成功未开局或设备异常，则进入退款或补开局", "玩家收到退款或补开局通知", "退款 / 复核闭环完成"],
    ]
    draw_table(draw, MARGIN, y + 10, CONTENT_W, reverse_headers, reverse_rows, [0.08, 0.16, 0.29, 0.23, 0.24], 18, "#D1FAE5")
    page = ASSET_DIR / "prd-page-3.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "六、支付异常情况与应对策略", 32, "#2563EB")
    exception_headers = ["阶段", "异常场景", "用户感知", "系统策略", "后台策略", "是否退款"]
    exception_rows = [
        ["扫码前", "会员码无效 / 过期", "提示扫码失败", "不创建订单，不扣款", "记录失败日志", "否"],
        ["确认前", "会员被冻结 / 资产不可用", "提示无法支付", "终止流程", "记录会员异常原因", "否"],
        ["扣款中", "第三方支付超时", "显示支付处理中或失败", "禁止重复扣款，轮询支付结果", "异常订单待复核", "视最终结果"],
        ["扣款后", "扣款成功但创单失败", "提示处理中", "自动补单或转异常订单", "运营可人工补单", "必要时退款"],
        ["开局中", "支付成功但设备未开局", "提示设备准备失败", "自动转异常订单，支持补开局 / 退款", "商家 / 平台可复核", "是"],
        ["游玩中", "玩家中途退出", "结束体验", "按规则结束订单", "标记结束原因=中途退出", "否"],
        ["游玩中", "人工强停 / 设备故障", "体验被打断", "结束订单并写异常原因", "进入异常处理或退款", "视规则"],
        ["结束后", "结束回传缺失", "玩家侧已离场", "订单保持异常待复核", "后台补录或人工处理", "视排查"],
        ["退款中", "原路退回失败", "提示退款处理中", "保留退款单号，重复回调防重", "转人工或线下退款", "是"],
        ["退款中", "已结算订单需退款", "等待店员处理", "要求上传线下退款凭证", "商家后台留痕，平台可查", "是"],
    ]
    y = draw_table(draw, MARGIN, y + 16, CONTENT_W, exception_headers, exception_rows, [0.11, 0.18, 0.17, 0.20, 0.20, 0.14], 17, "#FDE68A")
    y = draw_heading(draw, MARGIN, y + 16, "七、退款闭环策略", 30, "#2563EB")
    refund_headers = ["场景", "触发条件", "退款方式", "必要留痕", "备注"]
    refund_rows = [
        ["正常完成", "玩家正常游玩结束", "不退款", "结束原因、时长、设备", "默认闭环"],
        ["支付成功未开局", "设备未回传开局成功", "原路退回或线下退款", "异常订单、退款单号", "优先补开局，失败再退款"],
        ["中途退出", "玩家主动中止", "不退款", "结束原因=中途退出", "按已确认规则执行"],
        ["人工强停 / 设备故障", "店员或系统中断体验", "视规则退款", "异常原因、操作人、凭证", "可部分或全额退款"],
        ["已结算订单", "订单已进入结算", "线下退款", "退款凭证、退款原因、操作人", "后台须可审计"],
        ["未结算订单", "仍可调用支付渠道退款", "原路退回", "退款流水、回调状态", "支持异步回调"],
    ]
    draw_table(draw, MARGIN, y + 12, CONTENT_W, refund_headers, refund_rows, [0.16, 0.22, 0.17, 0.23, 0.22], 18, "#DBEAFE")
    page = ASSET_DIR / "prd-page-4.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "八、工程结构与模块边界", 32, "#2563EB")
    y = draw_para(draw, MARGIN, y + 8, "支付链路不是单一页面功能，而是用户侧、设备侧、订单服务、支付服务、会员资产服务、设备编排服务和退款中心共同协作的结果。", size=22)
    arch = Image.open(architecture).convert("RGB")
    arch.thumbnail((CONTENT_W, 900))
    img.paste(arch, (MARGIN, y + 14))
    y = y + 14 + arch.height + 20
    y = draw_heading(draw, MARGIN, y, "九、产品结论", 32, "#2563EB")
    draw_bullets(draw, MARGIN, y + 4, [
        "支付设计必须以“支付成功 + 开局成功 + 结束回传 + 可退款”四段闭环为准，不应只停留在支付结果页。",
        "反扫支付是会员消费重点链路，必须强化扫码成功、扣款成功、开局成功之间的状态衔接。",
        "异常订单中心和退款中心属于支付设计本体，而不是事后补救模块。",
        "商家后台与总运营后台都需要围绕订单、设备、支付方式、异常原因形成可查询、可复核、可追溯的数据闭环。"
    ])
    page = ASSET_DIR / "prd-page-5.png"
    img.save(page)
    pages.append(page)
    return pages


def build_docx() -> Path:
    ensure_dirs()
    overall = ASSET_DIR / "overall_flow.png"
    reverse = ASSET_DIR / "reverse_scan_flow.png"
    architecture = ASSET_DIR / "architecture.png"
    draw_overall_flow(overall)
    draw_reverse_flow(reverse)
    draw_architecture(architecture)
    page_paths = compose_pages(overall, reverse, architecture)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    content_width = section.page_width - section.left_margin - section.right_margin

    for idx, page in enumerate(page_paths):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.add_run().add_picture(str(page), width=content_width)
        if idx < len(page_paths) - 1:
            doc.add_page_break()

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    ensure_dirs()
    print(build_docx())
