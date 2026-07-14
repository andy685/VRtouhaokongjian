from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "user-payment-prd-assets"
DOCX_PATH = OUT_DIR / "用户支付PRD_玩家玩游戏支付流程_可编辑版.docx"
REVISED_DOCX_PATH = OUT_DIR / "用户支付PRD_玩家玩游戏支付流程_可编辑版_修正版.docx"
FONT_NAME = "Arial Unicode MS"


def apply_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, header_fill: str = "DBEAFE") -> None:
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                shade_cell(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    apply_font(run, 10.5)


def add_p(doc: Document, text: str, style: str | None = None, *, size: float = 10.5, bold: bool = False, color: str | None = None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    apply_font(r, size, bold, color)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    apply_font(r, 10.5)
    return p


def set_col_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    normal.font.size = Pt(10.5)

    for name, size, color in [("Heading 1", 16, "2563EB"), ("Heading 2", 13, "1D4ED8"), ("Heading 3", 11, "0F766E")]:
        st = doc.styles[name]
        st.font.name = FONT_NAME
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    add_p(doc, "用户支付 PRD", size=22, bold=True, color="0F172A", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "玩家玩游戏支付流程", size=15, bold=True, color="2563EB", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_p(doc, "适用于：点播系统 · 玩家支付链路 · 散客直付 / 会员主动扫码 / 会员码反扫 / 店员扫码点播", size=10.5, color="475569", align=WD_ALIGN_PARAGRAPH.CENTER)

    meta = doc.add_table(rows=3, cols=2)
    meta.style = "Table Grid"
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.cell(0, 0).text = "文档目标"
    meta.cell(0, 1).text = "梳理玩家从发起支付到游玩结束的完整链路，覆盖支付方式、反扫闭环、异常处理、退款闭环和工程结构边界。"
    meta.cell(1, 0).text = "重点范围"
    meta.cell(1, 1).text = "重点介绍会员码反扫支付流程，并覆盖扫码成功、扣款、开局、结束、退款全闭环。"
    meta.cell(2, 0).text = "目标读者"
    meta.cell(2, 1).text = "产品、研发、测试、运营、设备联调人员。"
    set_col_widths(meta, [1.4, 4.8])
    style_table(meta, "E8EEF5")

    add_p(doc, "一、文档概述", "Heading 1")
    add_p(doc, "本 PRD 聚焦“玩家要玩一局游戏”在支付层面的完整产品设计。文档不只描述支付成功页面，而是明确支付前置条件、不同触发方式、设备开局协同、订单归档、异常转人工与退款闭环。")
    add_bullet(doc, "主业务对象：玩家、店员、设备终端、商家后台、总运营后台。")
    add_bullet(doc, "主业务目标：让“支付成功”与“真正开局并可追溯”形成同一条闭环链路。")
    add_bullet(doc, "设计原则：支付清晰、设备协同清晰、异常可追踪、退款有依据。")

    add_p(doc, "二、支付能力全景", "Heading 1")
    add_p(doc, "2.1 玩家玩游戏主流程", "Heading 2")
    add_p(doc, "所有支付场景都应落到统一主链路：触发支付 → 明细确认 → 扣款成功 → 创单 → 开局 → 游玩结束 → 订单归档 / 异常处置。")
    flow1 = ASSET_DIR / "overall_flow.png"
    if flow1.exists():
        doc.add_picture(str(flow1), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_p(doc, "2.2 支付触发方式与可用支付方式", "Heading 2")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["触发方式", "前端载体", "适用对象", "可用支付方式", "支付成功后动作", "退款口径"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    rows = [
        ["散客直接支付", "设备端收款码 / 支付宝或微信付款码", "散客", "微信支付、支付宝支付", "支付成功后创建点播订单并启动游戏", "未开局可退款；正常完成不退款"],
        ["会员主动扫码", "微信小程序", "会员", "微信支付、预存款、混合支付（预存款+微信）", "支付成功后创建点播订单并启动游戏", "未开局可退款；正常完成不退款"],
        ["会员码反扫", "会员码 + 小程序确认", "会员", "预存款、会员余额、游戏币、资产不足时补差支付", "确认扣款后创单并自动开局", "支付成功未开局可退款；中途退出不退款；设备异常可复核 / 退款"],
        ["店员扫码点播", "店员 H5", "会员 / 团体客", "预存款、游戏币、微信补差", "店员代操作创建订单并开局", "按订单是否已结算决定原路退回或线下退款"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(table, [0.9, 0.9, 1.0, 2.3, 1.4, 1.5])
    style_table(table)

    add_p(doc, "2.3 支付方式定义", "Heading 2")
    add_bullet(doc, "散客直接支付：散客全程不使用头号空间小程序，不使用会员资产，直接通过微信或支付宝完成扫码支付，支付成功后启动游戏。")
    add_bullet(doc, "微信支付：玩家直接支付现金类金额，适用于散客直接支付、会员主动扫码和补差支付。")
    add_bullet(doc, "支付宝支付：适用于散客直接支付场景，与微信支付同属外部现金支付。")
    add_bullet(doc, "预存款 / 会员余额：直接扣减会员账户余额，不产生外部支付流水。")
    add_bullet(doc, "游戏币：适用于按游戏币计价的内容或活动。")
    add_bullet(doc, "混合支付：会员资产不足时，先扣会员资产，再补差微信支付。")

    add_p(doc, "2.4 玩家类型与支付入口边界", "Heading 2")
    boundary = doc.add_table(rows=1, cols=5)
    boundary.style = "Table Grid"
    boundary.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["玩家类型 / 场景", "是否使用头号空间小程序", "支付确认载体", "主要支付方式", "支付成功后动作"]):
        boundary.cell(0, i).text = h
    boundary_rows = [
        ["散客直接支付", "否", "微信 / 支付宝原生支付页", "微信支付、支付宝支付", "创建点播订单并启动游戏"],
        ["会员主动扫码", "是", "头号空间小程序", "预存款、游戏币、微信补差", "创建点播订单并启动游戏"],
        ["会员码反扫", "是", "头号空间小程序", "预存款、会员余额、游戏币、微信补差", "确认扣款后创单并自动开局"],
        ["店员扫码点播", "否（玩家侧）/ 是（店员侧 H5）", "店员手机 H5", "预存款、游戏币、微信补差", "店员代操作创单并启动游戏"],
    ]
    for row in boundary_rows:
        cells = boundary.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(boundary, [1.2, 1.0, 1.4, 1.8, 1.3])
    style_table(boundary, "E0F2FE")

    add_p(doc, "2.5 支付流程状态", "Heading 2")
    status_table = doc.add_table(rows=1, cols=3)
    status_table.style = "Table Grid"
    status_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["状态", "说明", "下一步"]):
        status_table.cell(0, i).text = h
    status_rows = [
        ["待支付", "已生成支付请求，等待玩家扫码或确认支付", "进入正在支付 / 取消支付"],
        ["正在支付", "已发起扣款或第三方支付请求，等待结果返回", "进入支付成功 / 支付失败"],
        ["支付成功", "支付结果确认成功，可继续创单并下发开局", "进入开局流程"],
        ["支付失败", "支付未完成，本次支付链路结束", "允许重新发起支付"],
    ]
    for row in status_rows:
        cells = status_table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(status_table, [1.0, 3.0, 1.6])
    style_table(status_table, "DBEAFE")

    add_p(doc, "支付失败原因", "Heading 3")
    add_bullet(doc, "支付超时：支付通道或回调在约定时间内未返回结果。")
    add_bullet(doc, "余额不足：微信 / 支付宝支付账户余额不足。")
    add_bullet(doc, "订单已失效：订单超时、已被使用或已被取消。")
    add_bullet(doc, "支付通道异常：第三方支付通道返回异常或网络中断。")

    add_p(doc, "三、玩家玩游戏主流程设计", "Heading 1")
    add_p(doc, "3.1 散客直接支付流程", "Heading 2")
    add_bullet(doc, "散客在设备端选择游戏或内容。")
    add_bullet(doc, "设备端展示应付金额与收款码，散客直接使用微信或支付宝扫码支付。")
    add_bullet(doc, "支付成功后，系统创建点播订单，并立即向设备下发开局指令。")
    add_bullet(doc, "设备开局成功后进入游玩中；游玩结束后回传时长、结束原因、设备信息。")
    add_bullet(doc, "若支付成功但未开局，则进入异常订单，并支持退款或补开局。")
    add_bullet(doc, "该链路不经过头号空间小程序，玩家支付确认、支付结果都在微信 / 支付宝侧完成。")
    add_bullet(doc, "散客链路不参与会员资产抵扣，不展示会员余额、游戏币、会员优惠券。")

    add_p(doc, "3.2 会员支付主流程", "Heading 2")
    steps = [
        "玩家在设备端、小程序或店员协助下选择内容和设备。",
        "系统根据入口判断支付方式与可用资产，并展示支付明细。",
        "玩家确认后，系统执行扣款或第三方支付。",
        "支付成功后立即创建点播订单，并将支付流水与设备、内容、会员绑定。",
        "订单创建成功后向设备下发开局指令，等待设备回执。",
        "游玩结束后回传结束原因、体验时长、设备信息，订单进入已完成或异常状态。",
    ]
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(step)
        apply_font(r, 10.5)

    add_p(doc, "3.3 各触发方式说明", "Heading 2")
    flow_desc = doc.add_table(rows=1, cols=4)
    flow_desc.style = "Table Grid"
    flow_desc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["方式", "玩家感知", "系统关键动作", "产品关注点"]):
        flow_desc.cell(0, i).text = h
    for row in [
        ["散客直接支付", "玩家看到设备端收款码后，直接用微信 / 支付宝扫码", "支付成功后立即创单并开局", "必须明确散客不走头号空间小程序，也不走会员资产扣减"],
        ["会员主动扫码", "会员主动扫设备码进入支付页", "设备先创待支付订单，小程序确认支付后回调终端", "页面清晰、支付中提示、支付结果回跳"],
        ["会员码反扫", "玩家出示会员码，被设备反扫后在小程序确认", "先识别会员，再校验资产，再扣款并开局", "支付不是终点，必须校验开局与结束闭环"],
        ["店员扫码点播", "店员代玩家完成支付确认", "H5 展示明细、代玩家确认、同步设备开局", "店员误操作防呆、订单可追溯"],
    ]:
        cells = flow_desc.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(flow_desc, [1.0, 1.6, 2.0, 1.8])
    style_table(flow_desc)

    add_p(doc, "四、反扫支付闭环（重点）", "Heading 1")
    add_p(doc, "反扫支付不是“扫到码就结束”，而是“识别会员 → 确认明细 → 扣款成功 → 开局成功 → 结束归档 → 异常可退款”的完整闭环。若缺少开局、结束、退款任一环节，都会导致支付链路不可追溯。")
    flow2 = ASSET_DIR / "reverse_scan_flow.png"
    if flow2.exists():
        doc.add_picture(str(flow2), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    detail = doc.add_table(rows=1, cols=5)
    detail.style = "Table Grid"
    detail.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["阶段", "动作发起方", "系统动作", "用户可见内容", "产出结果"]):
        detail.cell(0, i).text = h
    detail_rows = [
        ["1", "玩家 / 扫码枪", "扫码枪读取会员码，设备端解析会员身份", "设备提示“扫码中 / 识别中”", "进入会员识别"],
        ["2", "系统", "校验会员状态、会员资产、设备可用性", "若可用，则跳转小程序确认明细", "形成待确认支付明细"],
        ["3", "小程序", "展示内容、金额、资产抵扣、补差支付、优惠信息", "玩家可查看扣款明细并确认", "形成支付确认动作"],
        ["4", "支付服务", "扣减预存款 / 游戏币，或发起微信补差", "支付处理中、禁止重复提交", "支付成功 / 失败"],
        ["5", "订单服务", "支付成功后创建点播订单并写入支付流水", "用户可看到支付成功状态", "订单状态=待开局"],
        ["6", "设备编排", "向设备下发开局指令，等待设备回传", "前端显示“准备开局 / 正在进入游戏”", "开局成功或开局失败"],
        ["7", "设备", "游玩中回传状态，结束时回传结束原因与时长", "玩家完成体验或收到异常提示", "订单状态=已完成 / 异常"],
        ["8", "退款中心", "若支付成功未开局或设备异常，则进入退款或补开局", "玩家收到退款或补开局通知", "退款 / 复核闭环完成"],
    ]
    for row in detail_rows:
        cells = detail.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(detail, [0.5, 1.1, 2.2, 1.7, 1.2])
    style_table(detail, "D1FAE5")

    add_p(doc, "五、关键功能讲解", "Heading 1")
    add_bullet(doc, "扫码成功：仅代表会员码识别成功，不代表已经扣款。")
    add_bullet(doc, "扣款成功：必须生成订单与支付流水，并锁定本次设备与内容。")
    add_bullet(doc, "开局成功：必须有设备回执，才能把支付真正视为“已履约”。")
    add_bullet(doc, "结束回传：必须有结束原因和时长，才能把订单归档为“已完成”。")
    add_bullet(doc, "退款闭环：支付成功但未开局、设备故障、人工强停等场景，必须能回到退款中心完成处置。")

    add_p(doc, "六、支付异常情况与应对策略", "Heading 1")
    exc = doc.add_table(rows=1, cols=6)
    exc.style = "Table Grid"
    exc.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["阶段", "异常场景", "用户感知", "系统策略", "后台策略", "是否退款"]):
        exc.cell(0, i).text = h
    exc_rows = [
        ["扫码前", "会员码无效 / 过期", "提示扫码失败", "不创建订单，不扣款", "记录失败日志", "否"],
        ["确认前", "会员被冻结 / 资产不可用", "提示无法支付", "终止流程", "记录会员异常原因", "否"],
        ["扣款中", "第三方支付超时", "显示支付处理中或失败", "禁止重复扣款，轮询支付结果", "异常订单待复核", "视最终结果"],
        ["扣款后", "扣款成功但创单失败", "提示处理中", "自动补单或转异常订单", "运营可人工补单", "必要时退款"],
        ["开局中", "支付成功但设备未开局", "提示设备准备失败", "自动转异常订单，支持补开局 / 退款", "商家 / 平台可复核", "是"],
        ["游玩中", "玩家中途退出", "结束体验", "按规则结束订单", "标记结束原因=中途退出", "否"],
        ["游玩中", "人工强停 / 设备故障", "体验被打断", "结束订单并写异常原因", "进入异常处理或退款", "视规则"],
        ["结束后", "结束回传缺失", "玩家侧已离场", "订单保持异常待复核", "后台补录或人工处理", "视排查"],
        ["退款中", "原路退回失败", "提示退款处理中", "保留退款单号，重复回调防重", "转人工或线下退款", "是"],
        ["退款中", "已结算订单需退款", "等待店员处理", "要求上传线下退款凭证", "商家后台留痕，平台可查", "是"],
    ]
    for row in exc_rows:
        cells = exc.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(exc, [0.8, 1.2, 1.2, 1.5, 1.4, 0.7])
    style_table(exc, "FDE68A")

    add_p(doc, "七、退款闭环策略", "Heading 1")
    refund = doc.add_table(rows=1, cols=5)
    refund.style = "Table Grid"
    refund.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["场景", "触发条件", "退款方式", "必要留痕", "备注"]):
        refund.cell(0, i).text = h
    refund_rows = [
        ["正常完成", "玩家正常游玩结束", "不退款", "结束原因、时长、设备", "默认闭环"],
        ["支付成功未开局", "设备未回传开局成功", "原路退回或线下退款", "异常订单、退款单号", "优先补开局，失败再退款"],
        ["中途退出", "玩家主动中止", "不退款", "结束原因=中途退出", "按已确认规则执行"],
        ["人工强停 / 设备故障", "店员或系统中断体验", "视规则退款", "异常原因、操作人、凭证", "可部分或全额退款"],
        ["已结算订单", "订单已进入结算", "线下退款", "退款凭证、退款原因、操作人", "后台须可审计"],
        ["未结算订单", "仍可调用支付渠道退款", "原路退回", "退款流水、回调状态", "支持异步回调"],
    ]
    for row in refund_rows:
        cells = refund.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(refund, [1.1, 1.6, 1.2, 1.6, 1.3])
    style_table(refund, "DBEAFE")

    add_p(doc, "八、工程结构与模块边界", "Heading 1")
    add_p(doc, "支付链路不是单一页面功能，而是用户侧、设备侧、订单服务、支付服务、会员资产服务、设备编排服务和退款中心共同协作的结果。")
    flow3 = ASSET_DIR / "architecture.png"
    if flow3.exists():
        doc.add_picture(str(flow3), width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_bullet(doc, "订单服务负责全链路主键：支付后必须有唯一订单号贯穿支付、设备、异常、退款。")
    add_bullet(doc, "支付服务负责拆分支付方式、保证防重、防重复扣款、接收第三方支付回调。")
    add_bullet(doc, "会员资产服务负责预存款、会员余额、游戏币等资产校验与扣减。")
    add_bullet(doc, "设备编排服务负责开局、游玩中状态、结束回传。")
    add_bullet(doc, "异常订单中心和退款中心负责兜底，确保问题订单有后续处理路径。")

    doc.save(REVISED_DOCX_PATH)
    return REVISED_DOCX_PATH


if __name__ == "__main__":
    print(build())
