from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "three-role-guide-assets"
DOCX_PATH = OUT_DIR / "头号空间-三系统三角色使用教程.docx"

FONT_PATH = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
PAGE_SIZE = (1600, 2260)
PAGE_WIDTH, PAGE_HEIGHT = PAGE_SIZE
MARGIN = 96
CONTENT_W = PAGE_WIDTH - MARGIN * 2


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    try:
        return ImageFont.truetype(FONT_PATH, size=size)
    except Exception:
        return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> str:
    lines: list[str] = []
    for raw in text.split("\n"):
        current = ""
        for ch in raw:
            trial = current + ch
            bbox = draw.textbbox((0, 0), trial, font=font)
            if bbox[2] - bbox[0] > max_width and current:
                lines.append(current)
                current = ch
            else:
                current = trial
        lines.append(current or "")
    return "\n".join(lines)


def draw_heading(draw: ImageDraw.ImageDraw, x: int, y: int, text: str, size: int = 34, color: str = "#1E3A8A") -> int:
    font = load_font(size)
    draw.text((x, y), text, font=font, fill=color)
    bbox = draw.textbbox((x, y), text, font=font)
    return bbox[3] + 10


def draw_para(
    draw: ImageDraw.ImageDraw,
    x: int,
    y: int,
    text: str,
    *,
    width: int = CONTENT_W,
    size: int = 22,
    color: str = "#334155",
    spacing: int = 8,
) -> int:
    font = load_font(size)
    wrapped = wrap_text(draw, text, font, width)
    draw.multiline_text((x, y), wrapped, font=font, fill=color, spacing=spacing)
    bbox = draw.multiline_textbbox((x, y), wrapped, font=font, spacing=spacing)
    return bbox[3] + 10


def draw_bullets(draw: ImageDraw.ImageDraw, x: int, y: int, bullets: list[str], *, size: int = 21, width: int = CONTENT_W) -> int:
    for bullet in bullets:
        text_y = y
        y = draw_para(draw, x + 28, y, bullet, width=width - 28, size=size)
        draw.ellipse([x, text_y + 11, x + 10, text_y + 21], fill="#0F172A")
        y += 4
    return y


def draw_numbered(draw: ImageDraw.ImageDraw, x: int, y: int, items: list[str], *, size: int = 21, width: int = CONTENT_W) -> int:
    number_font = load_font(size)
    for idx, item in enumerate(items, start=1):
        label = f"{idx}."
        draw.text((x, y), label, font=number_font, fill="#0F172A")
        y = draw_para(draw, x + 34, y, item, width=width - 34, size=size)
        y += 2
    return y


def draw_table(
    draw: ImageDraw.ImageDraw,
    x: int,
    y: int,
    width: int,
    headers: list[str],
    rows: list[list[str]],
    ratios: list[float],
    *,
    font_size: int = 18,
    header_fill: str = "#DBEAFE",
) -> int:
    body_font = load_font(font_size)
    head_font = load_font(font_size + 1)
    col_widths = [int(width * ratio) for ratio in ratios]
    col_widths[-1] += width - sum(col_widths)

    def row_height(values: list[str], font) -> int:
        heights = []
        for value, col_width in zip(values, col_widths):
            wrapped = wrap_text(draw, value, font, col_width - 24)
            bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=6)
            heights.append((bbox[3] - bbox[1]) + 24)
        return max(heights)

    current_y = y
    header_h = row_height(headers, head_font)
    current_x = x
    for header, col_width in zip(headers, col_widths):
        draw.rectangle([current_x, current_y, current_x + col_width, current_y + header_h], fill=header_fill, outline="#475569", width=2)
        wrapped = wrap_text(draw, header, head_font, col_width - 24)
        bbox = draw.multiline_textbbox((0, 0), wrapped, font=head_font, spacing=6)
        draw.multiline_text(
            (current_x + 12, current_y + (header_h - (bbox[3] - bbox[1])) / 2),
            wrapped,
            font=head_font,
            fill="#0F172A",
            spacing=6,
        )
        current_x += col_width
    current_y += header_h

    for row in rows:
        height = row_height(row, body_font)
        current_x = x
        for value, col_width in zip(row, col_widths):
            draw.rectangle([current_x, current_y, current_x + col_width, current_y + height], fill="white", outline="#64748B", width=1)
            wrapped = wrap_text(draw, value, body_font, col_width - 24)
            bbox = draw.multiline_textbbox((0, 0), wrapped, font=body_font, spacing=6)
            draw.multiline_text(
                (current_x + 12, current_y + (height - (bbox[3] - bbox[1])) / 2),
                wrapped,
                font=body_font,
                fill="#334155",
                spacing=6,
            )
            current_x += col_width
        current_y += height
    return current_y + 6


def draw_callout(draw: ImageDraw.ImageDraw, x: int, y: int, width: int, title: str, body: str, *, fill: str = "#EFF6FF") -> int:
    title_font = load_font(21)
    body_font = load_font(20)
    title_wrapped = wrap_text(draw, f"{title}：", title_font, width - 28)
    body_wrapped = wrap_text(draw, body, body_font, width - 28)
    title_bbox = draw.multiline_textbbox((0, 0), title_wrapped, font=title_font, spacing=6)
    body_bbox = draw.multiline_textbbox((0, 0), body_wrapped, font=body_font, spacing=6)
    height = (title_bbox[3] - title_bbox[1]) + (body_bbox[3] - body_bbox[1]) + 30
    draw.rounded_rectangle([x, y, x + width, y + height], radius=18, fill=fill, outline="#93C5FD", width=2)
    draw.multiline_text((x + 14, y + 10), title_wrapped, font=title_font, fill="#0F172A", spacing=6)
    draw.multiline_text((x + 14, y + 16 + (title_bbox[3] - title_bbox[1])), body_wrapped, font=body_font, fill="#475569", spacing=6)
    return y + height + 8


def draw_page_no(draw: ImageDraw.ImageDraw, page_no: int) -> None:
    font = load_font(18)
    text = str(page_no)
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((PAGE_WIDTH - MARGIN - (bbox[2] - bbox[0]), PAGE_HEIGHT - 52), text, font=font, fill="#64748B")


def page_title(draw: ImageDraw.ImageDraw, y: int, title: str, subtitle: str | None = None) -> int:
    y = draw_heading(draw, MARGIN, y, title, 34, "#0F172A")
    if subtitle:
        y = draw_para(draw, MARGIN, y + 4, subtitle, size=20, color="#475569")
    return y + 4


def compose_pages() -> list[Path]:
    pages: list[Path] = []

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "头号空间三系统三角色使用教程", 44, "#0F172A")
    y = draw_para(draw, MARGIN, y + 8, "覆盖总运营后台、商户后台与点播系统，串联官方、商家/门店、玩家的完整业务流程。", size=24, color="#475569")
    y = draw_table(
        draw,
        MARGIN,
        y + 20,
        CONTENT_W,
        ["项目", "说明"],
        [
            ["适用版本", "v2.19（基于 2026年7月16日 仓库资料整理）"],
            ["目标角色", "官方运营、商家/门店管理员、到店玩家"],
            ["覆盖系统", "总运营后台（Platform） / 商户后台（Shop） / PC 点播系统"],
            ["文档目标", "把新门店从建档、设备配置、内容分发一路串到玩家点播和结算闭环。"],
        ],
        [0.20, 0.80],
        font_size=21,
    )
    y = draw_callout(draw, MARGIN, y + 8, CONTENT_W, "阅读建议", "先看总流程，再按角色阅读对应章节；如果门店正在准备开业，建议官方、商家和现场店员按本文顺序联动执行。")
    y = draw_heading(draw, MARGIN, y + 10, "1. 系统与角色映射", 30, "#2563EB")
    y = draw_table(
        draw,
        MARGIN,
        y + 6,
        CONTENT_W,
        ["系统", "主要使用角色", "核心目标", "重点章节"],
        [
            ["总运营后台", "官方运营 / 平台超管", "完成商家、门店、设备、内容和支付底座配置", "第 3 章"],
            ["商户后台", "老板 / 店长 / 店员", "管理设备、价格、会员、优惠券、订单和营收", "第 4 章"],
            ["PC 点播系统", "商家/门店（管理模式） + 玩家（体验模式）", "配置 Token、下载安装游戏、现场选游戏与开局", "第 5 章和第 6 章"],
        ],
        [0.18, 0.22, 0.40, 0.20],
        font_size=19,
    )
    y = draw_heading(draw, MARGIN, y + 10, "2. 全流程总览", 30, "#2563EB")
    y = draw_table(
        draw,
        MARGIN,
        y + 6,
        CONTENT_W,
        ["步骤", "责任角色", "使用系统", "关键动作", "完成结果"],
        [
            ["1", "官方", "总运营后台", "新建商家并开通后台账号", "商家拿到可登录的账号"],
            ["2", "官方", "总运营后台", "新建店铺并维护门店资料", "店铺归属、地址和注册码建立"],
            ["3", "官方", "总运营后台", "录入游戏并维护上线状态", "内容进入可分发游戏库"],
            ["4", "官方", "总运营后台", "录入主机并分配到门店", "门店获得对应主机"],
            ["5", "官方", "总运营后台", "生成主机 Token", "点播系统拥有激活凭证"],
            ["6", "官方", "总运营后台", "录入头显并绑定主机", "设备关系就绪"],
            ["7", "官方", "总运营后台", "内容分发 + 拉卡拉配置", "门店具备下载和收款条件"],
            ["8", "商家/门店", "商户后台", "设置密码、价格、会员和优惠", "门店完成营业准备"],
            ["9", "商家/门店", "PC 点播系统", "填写 Token、同步内容、下载安装游戏", "终端可供玩家点播"],
            ["10", "玩家", "PC 点播系统 + 小程序", "选游戏、扫码支付、佩戴头显、完成游玩", "订单和体验闭环"],
        ],
        [0.08, 0.15, 0.18, 0.31, 0.28],
        font_size=18,
    )
    draw_page_no(draw, 1)
    page = ASSET_DIR / "page-1.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = page_title(draw, y, "3. 官方角色：总运营后台操作教程", "目标：把一间新门店从“未建档”推进到“可营业可点播”。")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["常用模块", "后台路径", "用途"],
        [
            ["商家管理", "店铺管理 → 商家管理", "新建商家、维护后台账号和结算信息"],
            ["店铺列表", "店铺管理 → 店铺列表", "创建门店、维护门店状态和注册码"],
            ["游戏库", "内容中心 → 游戏库", "录入 / 编辑可分发游戏"],
            ["设备配置管理", "数据中心 → 设备配置管理", "录入主机、录入头显、分配门店、生成 Token、绑定设备"],
            ["内容分发", "内容中心 → 内容分发", "把游戏推送到指定门店"],
            ["拉卡拉配置", "平台财务 → 拉卡拉配置", "维护支付和分账基础信息"],
        ],
        [0.18, 0.28, 0.54],
        font_size=19,
    )
    y = draw_heading(draw, MARGIN, y + 10, "3.1 新建商家", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：店铺管理 → 商家管理 → 新增商家", size=20, color="#334155")
    y = draw_numbered(draw, MARGIN, y + 4, [
        "填写商家名称、联系人、联系电话、负责区域、代理商、手续费率和商家状态。",
        "创建商家管理员账号与密码；这套账号后续会交付给商家登录后台。",
        "补充提现账户：开户银行、银行卡号、开户人姓名、身份证号；如暂未确认，可后补。",
    ], size=20)
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["关键字段", "建议填写口径", "为什么重要"],
        [
            ["商家名称", "与合同或营业执照一致", "减少结算、对账和门店归属歧义"],
            ["管理员账号", "便于识别商家主体", "这是商家首次使用系统的核心凭证"],
            ["商家状态", "交付前保持“正常”", "状态异常会影响后续门店使用"],
            ["提现手续费率", "按商务方案填写", "影响打款成本和财务口径"],
        ],
        [0.18, 0.30, 0.52],
        font_size=18,
    )
    y = draw_heading(draw, MARGIN, y + 10, "3.2 新建店铺", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：店铺管理 → 店铺列表 → 新增店铺", size=20, color="#334155")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "必须先选对所属商家，再填写店铺名称、城市、地址、联系电话和营业状态。",
        "店铺创建后会生成门店注册码，可用于玩家现场注册会员。",
        "如门店有新会员礼包，建议同步约定注册送预存款或游戏币规则。",
    ], size=20)
    draw_page_no(draw, 2)
    page = ASSET_DIR / "page-2.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "3.3 录入游戏并准备上架", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：内容中心 → 游戏库 → 添加游戏", size=20, color="#334155")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["字段", "建议说明", "门店会直接受影响的点"],
        [
            ["游戏名称 / 题材", "与封面、宣传文案保持一致", "影响玩家搜索和点播理解"],
            ["运行平台", "明确主机游戏或一体机游戏", "直接决定现场安装路径"],
            ["付费模式", "按次 / 按时长", "影响玩家结算和门店定价"],
            ["体验时长 / 游戏豆消耗", "与商务方案一致", "影响门店成本和售价"],
            ["状态", "保持可上线 / 可分发", "否则门店看不到内容"],
        ],
        [0.18, 0.30, 0.52],
        font_size=18,
    )
    y = draw_heading(draw, MARGIN, y + 10, "3.4 录入主机、分配门店并生成 Token", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：数据中心 → 设备配置管理 → 主机管理", size=20, color="#334155")
    y = draw_numbered(draw, MARGIN, y + 4, [
        "录入主机编号、主机名称、设备类型、硬件配置、系统版本、MAC 地址。",
        "为主机分配所属商家和所属店铺。",
        "分配完成后生成 Token，并把 Token 交给门店现场配置到点播系统。",
    ], size=20)
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["规则", "解释"],
        [
            ["一台主机一个 Token", "Token 是点播系统的激活凭证，不建议多台主机共用。"],
            ["取消分配会吊销 Token", "主机换店或回收后，旧 Token 应视为失效。"],
            ["MAC 地址必填", "用于识别现场主机，减少错绑与误发。"],
        ],
        [0.28, 0.72],
        font_size=19,
        header_fill="#E0F2FE",
    )
    y = draw_heading(draw, MARGIN, y + 10, "3.5 录入头显、分配门店并绑定主机", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：数据中心 → 设备配置管理 → 头显管理", size=20, color="#334155")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "录入头显名称、设备型号、SN 码和固件版本。",
        "先分配到正确门店，再执行绑定主机。",
        "绑定时优先选择同店铺下的在线主机，避免跨店错误绑定。",
    ], size=20)
    draw_page_no(draw, 3)
    page = ASSET_DIR / "page-3.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "3.6 内容分发与拉卡拉配置", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "内容分发路径：内容中心 → 内容分发；支付配置路径：平台财务 → 拉卡拉配置", size=20, color="#334155")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["动作", "适用场景", "说明"],
        [
            ["智能分发", "日常更新或补发", "只下发变更文件，速度快，推荐为常规方式。"],
            ["全量分发", "首次分发、版本修复、大范围重装", "重新推送全部文件，适合首装或问题修复。"],
        ],
        [0.18, 0.24, 0.58],
        font_size=19,
    )
    y = draw_bullets(draw, MARGIN, y + 10, [
        "分发后要确认状态确实到了“已分发 / 可下载”，不要只看任务已提交。",
        "支付配置完成后，要安排门店做微信 / 支付宝测试，避免首单才发现收款未开通。",
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 10, "3.7 交付给商家的清单", 28, "#0F766E")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["交付项", "建议内容"],
        [
            ["账号资料", "商户后台管理员账号、初始密码、建议改密时间"],
            ["设备资料", "主机编号、头显清单、Token、建议设备密码"],
            ["内容资料", "已分发游戏清单、重点推荐内容、是否需要 USB 安装的一体机游戏"],
            ["支付资料", "门店是否已开通拉卡拉、测试订单结果、异常联系人"],
        ],
        [0.22, 0.78],
        font_size=19,
    )
    y = draw_heading(draw, MARGIN, y + 16, "4. 商家/门店角色：商户后台操作教程", 32, "#0F172A")
    y = draw_para(draw, MARGIN, y + 6, "前提：官方已经完成商家、店铺、主机、Token、内容分发和支付底座配置。", size=21, color="#475569")
    y = draw_table(
        draw,
        MARGIN,
        y + 10,
        CONTENT_W,
        ["常用模块", "后台路径", "主要作用"],
        [
            ["设备列表", "系统设置 → 设备列表", "查看主机和头显、核对状态、补做绑定"],
            ["单次消费项目", "商品管理 → 单次消费项目", "配置玩家现场可购买的体验项目和价格"],
            ["充值套餐", "运营管理 → 充值套餐", "设置会员储值活动"],
            ["会员级别 / 优惠券", "会员管理 / 运营管理", "设置会员权益、礼包和促销活动"],
        ],
        [0.18, 0.28, 0.54],
        font_size=18,
    )
    draw_page_no(draw, 4)
    page = ASSET_DIR / "page-4.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "4.1 首次登录后的三件事", 28, "#0F766E")
    y = draw_numbered(draw, MARGIN, y + 4, [
        "核对店铺名称、地址、联系电话和营业状态是否正确。",
        "进入设备列表确认主机数量、头显数量、所属门店和在线状态是否与现场一致。",
        "确认价格、支付、会员和优惠活动是否准备完毕，不要在玩家到店后临时新增商品。",
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 8, "4.2 设备管理与点播系统密码", 28, "#0F766E")
    y = draw_para(draw, MARGIN, y + 4, "操作路径：系统设置 → 设备列表", size=20, color="#334155")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "在“主机设备”页查看主机编号、MAC 地址、硬件信息、在线状态和已绑定头显数量。",
        "在“头显设备”页查看 SN 码、型号、电量、绑定主机和使用状态。",
        "如官方未完成绑定，门店可在后台补做同店铺下的头显绑定。",
        "建议每台主机都设置点播系统密码，避免顾客误入管理模式。",
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 8, "4.3 配置价格、会员与优惠", 28, "#0F766E")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["配置项", "后台路径", "建议配置内容"],
        [
            ["单次消费项目", "商品管理 → 单次消费项目", "按 30 分钟 / 60 分钟 / 单局体验等维度设置售价"],
            ["充值套餐", "运营管理 → 充值套餐", "如充 100 送 10、充 300 送 50 等会员储值活动"],
            ["会员级别", "会员管理 → 会员级别", "设置普通会员、金卡、钻石等折扣权益"],
            ["优惠券", "运营管理 → 优惠券", "配置满减券、折扣券、代金券、体验券"],
        ],
        [0.18, 0.28, 0.54],
        font_size=18,
    )
    y = draw_heading(draw, MARGIN, y + 8, "4.4 日常查看与交接班", 28, "#0F766E")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["关注事项", "建议查看页面", "你要看什么"],
        [
            ["点播是否正常开局", "点播系统订单", "支付方式、内容名称、设备名称、结束原因、是否异常"],
            ["当天收银情况", "收银订单 / 店铺销售日报", "现金、微信、支付宝、会员资产消耗情况"],
            ["活动效果", "会员消费排行 / 订单查询", "优惠券核销量、充值转化、复购情况"],
            ["交接班", "交接班记录", "班次销售额、现金留存、异常订单说明"],
        ],
        [0.22, 0.24, 0.54],
        font_size=18,
    )
    draw_page_no(draw, 5)
    page = ASSET_DIR / "page-5.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "5. 商家/门店角色：PC 点播系统管理模式教程", 32, "#0F172A")
    y = draw_para(draw, MARGIN, y + 6, "适用场景：新主机首次部署、重新配置 Token、同步新游戏、安装一体机内容或营业前巡检。", size=21, color="#475569")
    y = draw_heading(draw, MARGIN, y + 10, "5.1 开机前准备", 28, "#0F766E")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "确认当前主机已经在官方后台完成分配，并拿到了专属 Token。",
        "确认主机能联网，磁盘空间足够；大文件分发前最好测试下载速度。",
        "如门店要安装一体机内容，提前准备 USB 3.0 数据线。"
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 8, "5.2 配置 Token", 28, "#0F766E")
    y = draw_numbered(draw, MARGIN, y + 4, [
        "打开点播系统安装包或登录页，进入“基础配置 / 系统设置 / Token 设置”入口。",
        "把官方分配的主机 Token 粘贴进去并保存。",
        "保存后重新加载终端，让系统完成设备认证、门店绑定和内容同步。"
    ], size=20)
    y = draw_callout(draw, MARGIN, y + 6, CONTENT_W, "入口提示", "如果当前版本在登录页右上角提供“更多菜单”，通常可先进入“基础配置”再填写 Token；部署版则从管理模式进入系统设置后填写。")
    y = draw_heading(draw, MARGIN, y + 10, "5.3 同步内容并下载安装游戏", 28, "#0F766E")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "Token 生效后，终端应能拉到所属门店和已分发的游戏列表。",
        "如平台侧刚做完分发，现场可点击同步、刷新或重新进入内容页。",
        "开业前至少要把主推内容下载完成，并随机抽查 1 到 2 款能正常进入启动流程。"
    ], size=20)
    draw_page_no(draw, 6)
    page = ASSET_DIR / "page-6.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "5.4 安装规则：主机游戏与一体机游戏", 28, "#0F766E")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["内容类型", "终端上的动作", "现场要求", "安装结果"],
        [
            ["主机游戏", "点击“安装到主机”", "无需额外接入头显即可先装好主机内容", "安装完成后可直接走主机侧启动"],
            ["一体机游戏", "等待识别头显后点击“安装到一体机 / 头显”", "必须把头显通过 USB 3.0 接到电脑", "安装完成后才能真正下发到头显端"],
        ],
        [0.18, 0.28, 0.24, 0.30],
        font_size=18,
    )
    y = draw_bullets(draw, MARGIN, y + 8, [
        "如果提示“等待一体机接入”或“未连接数据线”，先检查 USB 线和头显连接状态。",
        "如果平台已分发但终端没有安装按钮，先确认内容运行平台是否匹配当前设备。",
        "门店应使用点播系统密码进入管理模式；顾客只负责选游戏和支付。"
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 10, "6. 玩家角色：点播系统体验教程", 32, "#0F172A")
    y = draw_para(draw, MARGIN, y + 6, "目标：让玩家从看到终端到完成支付、佩戴头显、开始游戏，整个过程简单清晰。", size=21, color="#475569")
    y = draw_table(
        draw,
        MARGIN,
        y + 10,
        CONTENT_W,
        ["阶段", "玩家看到什么", "需要做什么"],
        [
            ["进入终端", "首页游戏列表、分类、搜索、推荐位", "浏览想玩的游戏"],
            ["查看详情", "游戏封面、时长、价格、人数说明", "选择内容、时长和人数"],
            ["进入支付", "二维码或身份识别提示", "按自己的身份完成支付"],
            ["设备引导", "头显编号和佩戴提示", "按图示佩戴设备"],
            ["开始游玩", "终端进入等待或开局状态", "沉浸式体验，按工作人员提示操作"],
            ["结束返回", "终端回到首页，设备恢复空闲", "决定是否复玩或注册会员"],
        ],
        [0.16, 0.42, 0.42],
        font_size=18,
    )
    draw_page_no(draw, 7)
    page = ASSET_DIR / "page-7.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = draw_heading(draw, MARGIN, y, "6.1 到店后的支付方式说明", 28, "#0F766E")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["玩家身份", "常见支付方式", "操作说明"],
        [
            ["散客", "微信支付 / 支付宝支付", "终端展示二维码后，直接用手机扫一扫完成付款，无需先注册会员。"],
            ["会员主动扫码", "预存款、游戏币、微信补差", "会员主动扫描终端码，在小程序中确认优惠、余额和补差。"],
            ["会员码反扫", "预存款、游戏币、微信补差", "玩家出示会员码，终端识别后由小程序展示订单明细并确认支付。"],
            ["员工协助点播", "店员扫码代操作", "适合团客、测试或特殊场景，由店员手机端完成确认。"],
        ],
        [0.18, 0.22, 0.60],
        font_size=18,
    )
    y = draw_callout(draw, MARGIN, y + 8, CONTENT_W, "会员自动抵扣顺序", "常见口径为：优惠券 → 会员折扣 / 活动优惠 → 预存款 → 游戏币 → 微信补差。玩家看到的最终金额可能因此低于标价。")
    y = draw_heading(draw, MARGIN, y + 10, "6.2 佩戴设备与结束流程", 28, "#0F766E")
    y = draw_bullets(draw, MARGIN, y + 4, [
        "支付成功后，终端会提示分配到的头显编号或由店员直接递交设备。",
        "按图示佩戴头显，调整绑带、瞳距和清晰度；有问题及时示意店员。",
        "游戏结束后终端回到首页，设备恢复空闲，可继续选择下一款内容。",
        "若遇到支付成功未开局、设备故障或体验异常，应立即找店员处理补开局或退款。"
    ], size=20)
    y = draw_heading(draw, MARGIN, y + 10, "6.3 如何注册成为会员", 28, "#0F766E")
    y = draw_numbered(draw, MARGIN, y + 4, [
        "扫描门店注册码或打开头号空间小程序。",
        "填写手机号和基础信息，完成注册。",
        "领取门店设置的新会员礼包，如预存款、游戏币或体验券。",
        "后续到店时，可直接走会员扫码支付并享受折扣。"
    ], size=20)
    draw_page_no(draw, 8)
    page = ASSET_DIR / "page-8.png"
    img.save(page)
    pages.append(page)

    img = Image.new("RGB", PAGE_SIZE, "white")
    draw = ImageDraw.Draw(img)
    y = 90
    y = page_title(draw, y, "7. 常见问题与排查顺序")
    y = draw_table(
        draw,
        MARGIN,
        y + 8,
        CONTENT_W,
        ["问题现象", "先看哪里", "建议处理动作"],
        [
            ["点播系统提示 Token 无效", "平台侧主机分配与 Token 状态", "确认主机是否已分配门店、Token 是否被吊销或复制错误，再重新保存。"],
            ["门店看不到新游戏", "平台侧内容分发状态", "确认游戏已上线并已分发到该店，再在终端手动同步。"],
            ["一体机游戏装不上", "USB 连接与运行平台", "确认这是头显内容，并使用 USB 3.0 连接头显后再安装。"],
            ["头显无法绑定到主机", "设备归属门店", "确保头显和主机在同一店铺下；若归属错了，先回官方后台修正。"],
            ["支付成功但没有开局", "门店点播系统订单", "记录订单号和设备号，由门店判断是否补开局或退款，并同步官方排查异常。"],
            ["忘记管理密码", "商户后台设备列表", "由有权限的门店管理员重新设置点播系统密码。"],
        ],
        [0.26, 0.24, 0.50],
        font_size=19,
    )
    y = draw_callout(draw, MARGIN, y + 10, CONTENT_W, "推荐排查顺序", "先查平台侧有没有“建档、分配、分发”的基础问题，再查商户后台配置是否齐全，最后查现场网络、USB 连接、电量和玩家操作。", fill="#F8FAFC")
    y = draw_para(draw, MARGIN, y + 24, "交付结论：只要官方侧完成建档与分发、商家侧完成价格与密码配置、现场终端完成 Token 和安装校验，玩家就能在点播系统中顺畅完成选游戏、支付和游玩。", size=22, color="#0F172A")
    draw_page_no(draw, 9)
    page = ASSET_DIR / "page-9.png"
    img.save(page)
    pages.append(page)

    return pages


def build_docx() -> Path:
    ensure_dirs()
    page_paths = compose_pages()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    content_width = section.page_width - section.left_margin - section.right_margin

    for idx, page in enumerate(page_paths):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = 0
        paragraph.paragraph_format.space_after = 0
        paragraph.add_run().add_picture(str(page), width=content_width)
        if idx < len(page_paths) - 1:
            doc.add_page_break()

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    ensure_dirs()
    print(build_docx())
