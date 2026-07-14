from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "deliverables" / "用户支付PRD_玩家玩游戏支付流程.docx"
OUT = ROOT / "deliverables" / "用户支付PRD_玩家玩游戏支付流程_v2.docx"
ASSET_DIR = ROOT / "deliverables" / "user-payment-prd-assets"


REPLACEMENTS = {
    "预存款、会员余额、游戏币、预存次数、套票抵扣、资产不足时补差支付": "会员余额、游戏币、资产不足时补差支付",
    "预存款、游戏币、套票、预存次数、微信补差": "会员余额、游戏币、微信补差",
    "预存款、会员余额、游戏币、套票、次数、微信补差": "会员余额、游戏币、微信补差",
    "预存款、游戏币、套票、次数等资产校验与扣减": "会员余额、游戏币等资产校验与扣减",
    "校验会员状态、会员资产、套票、次数、设备可用性": "校验会员状态、会员资产、设备可用性",
    "扣减预存款 / 游戏币 / 次数 / 套票，或发起微信补差": "扣减会员余额 / 游戏币，或发起微信补差",
    "不展示会员余额、游戏币、套票、次数、会员优惠券。": "不展示会员余额、游戏币、会员优惠券。",
}


def replace_text(value: str) -> str:
    updated = value
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    return updated


def patch_docx_text(src: Path, out: Path) -> None:
    doc = Document(src)

    for para in doc.paragraphs:
        text = para.text
        new_text = replace_text(text)
        if new_text != text:
            para.text = new_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                new_text = replace_text(text)
                if new_text != text:
                    cell.text = new_text

    doc.save(out)


def replace_media(docx_path: Path) -> None:
    replacements = {
        "word/media/image1.png": ASSET_DIR / "overall_flow.png",
        "word/media/image2.png": ASSET_DIR / "reverse_scan_flow.png",
        "word/media/image3.png": ASSET_DIR / "architecture.png",
    }

    tmp = docx_path.with_suffix(".tmp.docx")
    with ZipFile(docx_path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename in replacements:
                zout.writestr(item, replacements[item.filename].read_bytes())
            else:
                zout.writestr(item, zin.read(item.filename))
    tmp.replace(docx_path)


def main() -> None:
    patch_docx_text(SRC, OUT)
    replace_media(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
