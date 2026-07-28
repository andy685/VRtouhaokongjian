from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
ASSET_DIR = OUT_DIR / "usage-guide-editable-assets"
DOC_FONT = "Songti SC"

OFFICIAL_DOC = OUT_DIR / "头号空间-官方运营测试执行手册.docx"
MERCHANT_DOC = OUT_DIR / "头号空间-商家门店测试执行手册.docx"
PLAYER_DOC = OUT_DIR / "头号空间-玩家全流程测试执行手册.docx"


def set_font(run, size: float = 11, *, bold: bool = False, color: str = "1F2937") -> None:
    run.font.name = DOC_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    normal = doc.styles["Normal"]
    normal.font.name = DOC_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 22, "0F172A"),
        ("Heading 1", 16, "1E3A8A"),
        ("Heading 2", 13.5, "0F766E"),
        ("Heading 3", 11.5, "334155"),
    ]:
        style = doc.styles[name]
        style.font.name = DOC_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, *, header_fill: str = "DBEAFE") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ridx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ridx == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(3)
                for run in paragraph.runs:
                    set_font(run, 10, bold=(ridx == 0))


def add_title(doc: Document, title: str, subtitle: str, meta: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(title), 21, bold=True, color="0F172A")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p2.add_run(subtitle), 11, color="64748B")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = meta[0][0]
    table.rows[0].cells[1].text = meta[0][1]
    for key, value in meta[1:]:
        row = table.add_row().cells
        row[0].text = key
        row[1].text = value
    style_table(table, header_fill="E0F2FE")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    size = 16 if level == 1 else 13.5 if level == 2 else 11.5
    color = "1E3A8A" if level == 1 else "0F766E" if level == 2 else "334155"
    set_font(p.add_run(text), size, bold=True, color=color)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), 10.5, color="334155")


def add_note(doc: Document, label: str, text: str, *, fill: str = "FEF3C7") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    shade_cell(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    set_font(p.add_run(f"{label}："), 10.5, bold=True, color="92400E")
    set_font(p.add_run(text), 10.5, color="78350F")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_font(p.add_run(item), 10.5, color="334155")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_font(p.add_run(item), 10.5, color="334155")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], *, header_fill: str = "DBEAFE") -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, head in enumerate(headers):
        table.cell(0, idx).text = head
    for row_data in rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_data):
            row[idx].text = value
    style_table(table, header_fill=header_fill)


def add_image(doc: Document, path: Path, caption: str, *, width_cm: float = 16.2) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.add_run(caption), 9.5, color="64748B")


def add_case(
    doc: Document,
    code: str,
    title: str,
    objective: str,
    steps: list[str],
    expected: list[str],
    *,
    path_text: str | None = None,
    exceptions: list[str] | None = None,
) -> None:
    add_heading(doc, f"{code} {title}", 2)
    add_para(doc, f"测试目标：{objective}")
    if path_text:
        add_para(doc, f"操作路径：{path_text}")
    add_para(doc, "操作步骤：")
    add_numbers(doc, steps)
    add_para(doc, "预期结果：")
    add_bullets(doc, expected)
    if exceptions:
        add_para(doc, "异常分支：")
        add_bullets(doc, exceptions)


def build_merchant_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 - 商家门店测试执行手册",
        "Word 交付版｜面向测试人员、产品验收人员、门店联调人员",
        [
            ("适用版本", "v2.20"),
            ("更新日期", "2026年7月28日"),
            ("覆盖系统", "商家后台 admin-dashboard Shop + PC 收银系统 cashier-ui + PC 点播系统"),
            ("文档目标", "测试人员按步骤执行，完成门店后台、收银、点播与首单开局的核心验收"),
        ],
    )
    add_para(doc, "本文为测试执行版，不是培训材料。门店侧不能只测后台和收银，还要把后台配置结果真正串到点播系统，确认内容同步、游戏下载和首单开局都能走通。建议测试按“后台基线 → 设备 → 收银终端 → 点播系统 → 会员 → 收银 → 首单开局 → 异常订单”的顺序推进，并记录每个用例的实际结果与截图。")
    add_note(doc, "截图建议", "门店版文档应该配截图。当前版本已先放入流程图和关键页面示意图，后续如果你们要做正式测试包，建议再补真实后台页面截图。")

    add_heading(doc, "一、测试前准备", 1)
    add_table(
        doc,
        ["检查项", "要求"],
        [
            ["商家账号", "已创建，状态正常，可登录商家后台"],
            ["店铺", "至少 1 家门店可用"],
            ["主机", "至少 1 台主机已分配到门店"],
            ["头显", "至少 1 台头显已分配到同门店"],
            ["内容", "至少 1 个游戏已完成分发"],
            ["收银终端", "商家后台可新建收银终端并复制 Token"],
            ["测试账号", "至少准备 1 个店长账号、1 个收银员账号"],
        ],
    )
    add_table(
        doc,
        ["测试会员场景", "手机号", "预期"],
        [
            ["已是本店会员", "17600110765", "不允许重复添加"],
            ["全局已存在未关联本店", "15812345678", "自动关联到本店"],
            ["全新会员", "自备未使用手机号", "创建最小档案"],
        ],
        header_fill="DCFCE7",
    )

    add_heading(doc, "二、测试顺序总览", 1)
    add_image(doc, ASSET_DIR / "merchant_prep_flow.png", "图 1：门店开业准备流程示意图")
    add_image(doc, ASSET_DIR / "merchant_system_map.png", "图 2：商家后台与收银/点播终端协同示意图")
    add_table(
        doc,
        ["链路阶段", "测试目标", "通过标准"],
        [
            ["商家后台", "确认门店、设备、价格、会员和收银终端配置可用", "后台配置项可见、可保存、可回查"],
            ["收银终端", "确认收银 Token、登录和会员 / 支付链路可用", "可绑定、可登录、可下单"],
            ["点播系统", "确认主机 Token、内容同步、游戏下载与管理模式可用", "可识别门店、可看到内容、可装游戏"],
            ["玩家开局", "确认点播成功后能进入设备引导或开局闭环", "至少完成 1 次首单测试"],
        ],
        header_fill="E0F2FE",
    )
    add_note(doc, "重点提醒", "商家文档必须把“商家后台配置”与“点播系统执行”串起来。只有当门店在点播系统里完成 Token 写入、内容同步、游戏下载和首单开局后，门店链路才算真正通过。", fill="DBEAFE")

    add_heading(doc, "三、菜单与入口基线", 1)
    add_table(
        doc,
        ["模块", "关键页面"],
        [
            ["商品管理", "单次消费项目、实物商品"],
            ["运营管理", "充值套餐、优惠券、活动赠送"],
            ["会员管理", "会员列表、会员级别、会员储值变更、会员游戏币查询"],
            ["订单管理", "收银订单、点播系统订单、异常订单"],
            ["系统设置", "设备列表、收银终端、收银小票、用户列表、角色列表"],
        ],
    )
    add_table(
        doc,
        ["当前版本不应出现的入口", "预期"],
        [
            ["IC 卡管理", "不展示"],
            ["员工卡设置", "不展示"],
            ["套票", "不展示"],
            ["第三方外围设备", "不展示"],
            ["商家后台支付设置 / 系统参数", "不作为当前门店测试入口"],
        ],
        header_fill="FEE2E2",
    )

    add_heading(doc, "四、核心测试用例", 1)
    add_case(
        doc,
        "M01",
        "商家后台登录成功",
        "确认门店账号可正常进入后台，首页和菜单加载正确。",
        ["打开商家后台登录页。", "输入店长或管理员账号密码。", "点击登录。"],
        ["登录成功，进入“今日概况”。", "页面无空白、无报错。", "左侧菜单与本手册基线一致。"],
    )
    add_case(
        doc,
        "M02",
        "设备列表与绑定关系正确",
        "确认门店可看到官方已分配的主机和头显，且归属关系正确。",
        [
            "进入“主机设备”页签。",
            "核对主机编号、名称、MAC 地址、所属门店、在线状态。",
            "切换到“头显设备”页签，核对头显名称、SN 码、型号、绑定主机。",
        ],
        [
            "主机和头显都可见。",
            "主机与头显所属门店一致。",
            "绑定关系与现场预期一致。",
            "头显详情页可查看电量、IPD、绑定主机等信息。",
        ],
        path_text="系统设置 → 设备列表",
        exceptions=["异店头显不能绑定到当前主机。", "当店铺无主机时，绑定主机弹窗应给出明确提示。"],
    )
    add_case(
        doc,
        "M03",
        "门店补绑头显成功",
        "确认商家可在同店铺范围内补做设备绑定。",
        ["在“主机设备”中找到目标主机。", "点击“绑定头显”。", "输入同门店头显 SN 码并确认绑定。"],
        ["绑定成功。", "主机绑定头显数增加。", "头显详情中的绑定主机更新为目标主机。"],
        path_text="系统设置 → 设备列表",
        exceptions=["输入不存在的 SN 时提示未找到设备。", "输入异店头显 SN 时提示不同门店，禁止绑定。"],
    )
    add_case(
        doc,
        "M04",
        "设置点播系统密码成功",
        "确认门店可为主机设置点播系统管理密码。",
        ["选择一台主机。", "点击“修改点播系统密码”。", "输入新密码并确认。", "保存。"],
        ["保存成功。", "页面给出成功提示。", "后续点播系统进入管理模式时要求输入该密码。"],
        path_text="系统设置 → 设备列表 → 主机设备 → 修改点播系统密码",
        exceptions=["两次密码不一致时不允许保存。", "留空保存的行为应与产品口径一致并可被记录。"],
    )
    add_case(
        doc,
        "M05",
        "新建收银终端并复制 Token",
        "确认商家后台可以为收银机生成 Token。",
        ["进入“收银终端”页面。", "点击“新建收银设备”。", "选择店铺并填写终端名称。", "记录系统自动生成的 Token。", "点击创建并复制 Token。"],
        ["新设备创建成功。", "列表出现新终端。", "初始状态为“未使用”。", "Token 可复制。"],
        path_text="系统设置 → 收银终端",
        exceptions=["终端名称为空时不允许创建。", "重复 Token 应被阻止或自动重新生成。"],
    )
    add_case(
        doc,
        "M06",
        "收银系统绑定本机成功",
        "确认 PC 收银系统必须先绑定收银 Token，才能登录。",
        ["打开收银系统登录页。", "不绑定 Token，直接尝试登录。", "进入“更多菜单 → 基础配置”。", "粘贴从商家后台“收银终端”复制的 Token。", "点击“绑定本机”。", "绑定成功后再次登录。"],
        ["未绑定时，登录被拦截。", "绑定成功后，显示设备名称、门店、脱敏 Token。", "再次登录可成功进入销售页。"],
        path_text="PC 收银系统登录页 → 更多菜单 → 基础配置",
        exceptions=["无效 Token 提示 Token 无效。", "别机已绑定 Token 提示已被其他设备使用。", "已禁用 Token 提示该收银设备已被禁用。"],
    )
    add_case(
        doc,
        "M07",
        "点播系统 Token 写入与门店识别成功",
        "确认商家能从后台拿到主机 / 点播系统 Token，并在点播系统完成门店识别。",
        ["进入商家后台“系统设置 → 设备列表 → 主机设备”。", "找到目标主机并复制点播系统 Token。", "打开 PC 点播系统，进入管理模式或基础配置。", "粘贴点播系统 Token 并保存。", "重启或刷新点播系统。"],
        ["点播系统成功识别当前门店。", "可看到门店名称或对应设备身份。", "Token 保存后不会报无效或未分配。"],
        path_text="商家后台设备列表 → 主机设备 Token；PC 点播系统 → 基础配置 / 管理模式",
        exceptions=["Token 无效时应提示重新核对。", "主机未分配到门店时，点播系统不能正常识别门店。"],
    )
    add_case(
        doc,
        "M08",
        "点播系统内容同步与游戏下载成功",
        "确认后台分发结果可以真正同步到点播系统，门店能看到并安装目标内容。",
        [
            "完成点播系统 Token 写入。",
            "在点播系统触发内容同步或重新加载。",
            "查看是否出现平台已分发的游戏列表。",
            "选择至少 1 个主推游戏执行安装。",
            "如涉及一体机内容，再验证头显侧安装链路。",
        ],
        [
            "点播系统能看到目标门店已分发内容。",
            "至少 1 个游戏安装成功。",
            "安装状态、下载状态和完成提示清晰可见。",
        ],
        path_text="PC 点播系统 → 同步内容 / 游戏列表 / 安装",
        exceptions=["看不到新游戏时，先检查平台是否已分发。", "安装失败时，应能记录失败状态并支持重试。"],
    )
    add_case(
        doc,
        "M09",
        "新增会员仅允许手机号建档",
        "确认 v2.20 的会员建档规则已经生效。",
        ["在收银系统进入“销售 → 新增会员”。", "输入已是本店会员手机号。", "输入全局已存在但未关联本店的手机号。", "输入全新手机号。"],
        ["本店已存在时提示该手机号已是本店会员。", "全局已存在未关联时自动关联到本店。", "全新手机号时创建最小档案。", "整个流程中不要求填写姓名、性别、生日、备注。"],
    )
    add_case(
        doc,
        "M10",
        "散客收银与会员收银成功",
        "确认无会员和有会员两类场景都可完成收银闭环。",
        ["在销售页添加商品或体验项目。", "分别执行一次散客收银和一次会员收银。", "会员场景观察优惠和资产抵扣结果。", "完成结算。"],
        ["两类订单都能创建成功。", "成功页显示支付完成。", "订单列表可查询到对应订单。", "会员场景下优惠和补差结果正确。"],
    )
    add_case(
        doc,
        "M11",
        "点播首单开局闭环成功",
        "确认门店完成后台配置、内容同步和支付后，至少能成功走通一次点播开局。",
        [
            "在点播系统选择已安装游戏。",
            "完成 1 次散客或会员支付确认。",
            "观察点播系统是否进入设备分配、佩戴引导或开局状态。",
            "如有头显编号提示，核对与后台绑定关系一致。",
            "完成体验结束后的返回或订单回查。",
        ],
        [
            "支付成功后能进入开局链路，而不是停留在待支付。",
            "设备编号或引导信息正确。",
            "订单在后台或相关列表中可回查。",
        ],
        path_text="PC 点播系统 → 选游戏 → 支付 → 开局 / 设备引导",
        exceptions=["支付成功但未开局时，应记录订单号并回查点播系统订单。", "设备编号错误时，优先回查主机与头显绑定关系。"],
    )
    add_case(
        doc,
        "M12",
        "订单、营收、交班与异常订单闭环",
        "确认交易结果可回溯，异常订单流程可用。",
        ["完成至少 1 笔散客单、1 笔会员单、1 笔充值单。", "到订单页查询。", "到营收页查看汇总。", "执行一次交班。", "对一笔订单执行标记异常、取消标记或重新处理。"],
        ["三类订单均可查到。", "营收统计有数据变化。", "交班记录生成成功。", "异常订单支持标记、取消标记、查看驳回原因和重新处理。"],
    )

    add_heading(doc, "五、测试完成标准", 1)
    add_bullets(
        doc,
        [
            "商家后台登录、设备查看、收银终端、主机 Token、商品、会员、订单、营收、交班均通过。",
            "收银系统完成 Token 绑定、会员建档、散客收银、会员收银、充值测试。",
            "点播系统完成 Token 写入、内容同步、游戏下载和至少 1 次首单开局测试。",
            "当前版本暂隐能力未误开放。",
            "异常订单链路至少验证 1 次。",
        ],
    )

    add_heading(doc, "六、截图清单建议", 1)
    add_table(
        doc,
        ["截图名称", "建议来源"],
        [
            ["商家后台首页", "登录后今日概况页"],
            ["设备列表页", "主机设备 + 头显设备页签"],
            ["收银终端页", "新建终端并复制 Token 的页面"],
            ["收银绑定页", "登录页基础配置弹层"],
            ["点播系统基础配置页", "写入主机 / 点播系统 Token 的页面"],
            ["点播系统内容列表页", "同步后的游戏列表与安装状态"],
            ["点播系统开局页", "支付成功后设备引导 / 开局状态"],
            ["新增会员页", "仅手机号建档页面"],
            ["异常订单页", "标记异常 / 已驳回 / 重新处理状态"],
        ],
        header_fill="FEF3C7",
    )

    doc.save(MERCHANT_DOC)
    return MERCHANT_DOC


def build_official_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 - 官方运营测试执行手册",
        "Word 交付版｜面向测试人员、运营验收人员、实施交付人员",
        [
            ("适用版本", "v2.20"),
            ("更新日期", "2026年7月28日"),
            ("覆盖系统", "总运营后台 admin-dashboard Platform"),
            ("文档目标", "测试人员按步骤执行，完成官方交付链路核心验收"),
        ],
    )
    add_para(doc, "本文用于测试执行，不用于角色培训。建议按“商家开通线 → 设备交付线 → 内容分发线 → 支付配置线 → 下游联动验证”的顺序执行。")
    add_note(doc, "截图建议", "官方版也建议配截图。当前版本先放入流程图和关键交付示意图，后续可继续补真实平台后台页面截图。")

    add_heading(doc, "一、测试前准备", 1)
    add_table(
        doc,
        ["检查项", "要求"],
        [
            ["平台账号", "可登录 Platform 角色"],
            ["商家样例数据", "可新建 1 个全新商家"],
            ["店铺样例数据", "可在新商家下创建至少 1 家店铺"],
            ["设备资料", "至少准备 1 台主机、1 台头显的编号或 SN 信息"],
            ["游戏内容", "至少准备 1 个可上线或可分发游戏"],
            ["支付资料", "至少准备 1 组可测试的支付配置示例"],
        ],
    )

    add_heading(doc, "二、官方交付总览", 1)
    add_image(doc, ASSET_DIR / "official_flow.png", "图 1：官方交付流程示意图")
    add_image(doc, ASSET_DIR / "official_device_map.png", "图 2：官方交付包示意图")

    add_heading(doc, "三、平台菜单基线", 1)
    add_table(
        doc,
        ["模块", "关键页面"],
        [
            ["店铺管理", "商家管理、店铺列表、代理商"],
            ["内容中心", "游戏库、内容分发、游戏审核"],
            ["数据中心", "设备配置管理、设备运行总览"],
            ["平台财务", "营收总览、异常订单、拉卡拉配置、结算管理"],
            ["平台通知", "公告管理、消息推送、系统通知"],
            ["帮助与运维", "帮助文档、FAQ、版本发布、操作日志"],
        ],
    )

    add_heading(doc, "四、核心测试用例", 1)
    add_case(
        doc,
        "O01",
        "平台后台登录成功",
        "确认平台超管可正常进入总运营后台。",
        ["打开平台登录页。", "输入平台账号密码。", "点击登录。"],
        ["成功进入“大屏看板”。", "左侧菜单正常展示。", "页面无报错、无空白。"],
    )
    add_case(
        doc,
        "O02",
        "新建商家成功",
        "确认平台可创建新的商家主体，并为商家生成后台账号。",
        ["点击“新增商家”。", "填写商家名称、联系人、联系电话、区域、代理归属、手续费率。", "填写商家管理员账号和密码。", "保存。"],
        ["商家创建成功。", "列表可看到新商家。", "状态可设置为“正常”。", "商家拥有后台登录凭证。"],
        path_text="店铺管理 → 商家管理",
    )
    add_case(
        doc,
        "O03",
        "新建店铺成功",
        "确认平台可在商家下创建门店，并生成注册码。",
        ["点击“新增店铺”。", "选择所属商家。", "填写店铺名称、所属区域、详细地址、联系电话、状态。", "保存。"],
        ["店铺创建成功。", "店铺可挂到目标商家下。", "店铺列表中可见新门店。", "可生成或展示注册码。"],
        path_text="店铺管理 → 店铺列表",
    )
    add_case(
        doc,
        "O04",
        "录入主机并分配到店铺",
        "确认平台可录入点播主机，并完成门店分配。",
        ["点击“录入主机”。", "填写主机编号、主机名称、设备类型、配置、系统版本、MAC 地址。", "保存后执行“分配”。", "选择商家和店铺并确认。"],
        ["主机录入成功。", "主机分配成功。", "分配后归属商家和店铺正确。"],
        path_text="数据中心 → 设备配置管理 → 主机管理",
        exceptions=["MAC 地址为空时不应保存。", "未分配店铺时，下游门店不应可正常使用。"],
    )
    add_case(
        doc,
        "O05",
        "生成 Token 并可交付下游",
        "确认平台侧主机 Token 可生成并用于门店点播系统配置。",
        ["在已分配主机的行内点击“生成 Token”。", "记录 Token。", "确认 Token 状态为可用。"],
        ["Token 成功生成。", "Token 可复制。", "该 Token 可交给门店在点播系统填写。"],
        exceptions=["此处验证的是主机 / 点播系统 Token，不是商家后台的收银 Token。"],
    )
    add_case(
        doc,
        "O06",
        "录入头显并绑定主机",
        "确认平台可录入头显，并绑定到同店铺主机。",
        ["点击“录入头显”。", "填写头显名称、型号、SN、固件版本。", "保存后执行“分配”。", "选择目标商家和店铺。", "执行“绑定到主机”，选择同店铺主机。"],
        ["头显录入成功。", "头显分配成功。", "绑定成功后，主机与头显建立关系。"],
        path_text="数据中心 → 设备配置管理 → 头显管理",
        exceptions=["异店主机不应出现在可绑定范围内。", "SN 重复应被拦截或提示异常。"],
    )
    add_case(
        doc,
        "O07",
        "内容分发成功",
        "确认平台可将游戏分发到指定门店，门店后续可拉取内容。",
        ["找到目标游戏。", "点击“分发”。", "选择目标店铺。", "选择“智能分发”作为主测方式。", "提交任务并查看分发详情。"],
        ["分发任务创建成功。", "状态能从未分发进入分发中，再进入已分发。", "门店后续应能看到对应内容。"],
        path_text="内容中心 → 内容分发",
        exceptions=["任务失败时应可重试。", "单店铺分发状态应可单独查看。"],
    )
    add_case(
        doc,
        "O08",
        "拉卡拉配置保存成功",
        "确认平台可维护支付配置，为门店支付测试提供前提。",
        ["进入拉卡拉配置页。", "录入商户号、终端号、机构号等信息。", "上传证书或补充必要字段。", "保存。"],
        ["配置保存成功。", "页面无校验异常。", "可作为门店微信 / 支付宝测试前置条件。"],
        path_text="平台财务 → 拉卡拉配置",
    )
    add_case(
        doc,
        "O09",
        "平台交付结果可被门店消费",
        "确认平台交付结果不只是“保存成功”，而是真的能驱动下游使用。",
        ["用新建商家账号登录商家后台。", "确认店铺在商家侧可见。", "确认主机、头显在设备列表可见。", "使用主机 Token 在门店侧进行点播系统接入验证。", "确认门店能看到已分发内容。", "确认门店可执行首单支付测试。"],
        ["上游交付结果全部能被下游正确消费。", "只有当下游可见时，才算本轮官方交付通过。"],
    )
    add_case(
        doc,
        "O10",
        "异常定位能力可支持下游排查",
        "确认当门店反馈问题时，平台测试人员知道先查哪里。",
        ["模拟门店看不到游戏。", "模拟头显无法绑定主机。", "模拟点播系统无法识别 Token。", "模拟门店不能支付。"],
        ["能分别回到内容分发、设备配置管理、主机管理、拉卡拉配置进行定位。"],
    )

    add_heading(doc, "五、完成标准", 1)
    add_bullets(
        doc,
        [
            "商家和店铺可成功创建。",
            "主机和头显可成功录入、分配、绑定。",
            "Token 可成功生成并交付。",
            "内容可成功分发到门店。",
            "支付配置可成功保存。",
            "下游门店可基于上述结果完成基础使用验证。",
        ],
    )

    add_heading(doc, "六、截图清单建议", 1)
    add_table(
        doc,
        ["截图名称", "建议来源"],
        [
            ["平台首页", "登录后大屏看板"],
            ["商家管理页", "新增商家弹窗与列表结果"],
            ["店铺列表页", "新增店铺弹窗与列表结果"],
            ["主机管理页", "录入主机、分配、生成 Token 的操作结果"],
            ["头显管理页", "录入头显、分配、绑定主机结果"],
            ["内容分发页", "分发任务状态与详情"],
            ["拉卡拉配置页", "配置保存结果"],
        ],
        header_fill="FEF3C7",
    )

    doc.save(OFFICIAL_DOC)
    return OFFICIAL_DOC


def build_player_doc() -> Path:
    doc = Document()
    style_doc(doc)
    add_title(
        doc,
        "头号空间 - 玩家全流程测试执行手册",
        "Word 交付版｜面向测试人员、产品验收人员、现场联调人员",
        [
            ("适用版本", "v2.20"),
            ("更新日期", "2026年7月28日"),
            ("覆盖系统", "PC 点播系统 + 微信小程序 + VR 头显终端"),
            ("文档目标", "测试人员按步骤执行，完成玩家到店、支付、开局、体验结束的全流程验收"),
        ],
    )
    add_para(doc, "本文关注玩家真实体验路径，不能只验证“能支付”，还要验证“支付后能进入设备引导/开局，结束后能正确回流”。建议按“浏览游戏 → 选择内容 → 支付确认 → 佩戴设备 → 开局体验 → 结束返回 → 会员注册/复玩”的顺序执行。")
    add_note(doc, "截图建议", "玩家版适合同时放流程图和真实页面截图。当前版本已放入流程图，并补充了小程序支付确认页、支付成功页、支付失败页等现有素材。")

    add_heading(doc, "一、测试前准备", 1)
    add_table(
        doc,
        ["检查项", "要求"],
        [
            ["门店环境", "PC 点播系统可正常打开，至少 1 台头显可用"],
            ["内容环境", "至少 1 个已安装、可体验的游戏"],
            ["支付环境", "可发起散客扫码支付和会员支付"],
            ["会员环境", "至少准备 1 个已有余额 / 游戏币会员，1 个可注册新会员"],
            ["观察点", "要能看到支付确认结果、设备引导结果、体验结束回流结果"],
        ],
    )
    add_table(
        doc,
        ["玩家场景", "重点验证内容", "通过标准"],
        [
            ["散客", "浏览、扫码支付、开局、结束返回", "无需登录也可完成闭环"],
            ["会员主动扫码", "优惠、预存款、游戏币、补差", "扣费和展示结果正确"],
            ["会员码反扫 / 店员协助", "识别会员、确认扣费、进入开局", "会员身份和支付链路正确"],
            ["新玩家注册", "注册二维码 / 小程序入会", "可成功建档并进入会员链路"],
        ],
        header_fill="DCFCE7",
    )

    add_heading(doc, "二、玩家全流程总览", 1)
    add_image(doc, ASSET_DIR / "player_flow.png", "图 1：玩家体验主流程示意图")
    add_image(doc, ASSET_DIR / "player_payment_map.png", "图 2：玩家支付路径示意图")

    add_heading(doc, "三、支付页面素材参考", 1)
    add_para(doc, "以下截图用于帮助测试人员理解玩家在小程序或支付确认页会看到的关键状态。")
    add_image(doc, ROOT / "miniapp-payment" / "screenshots" / "08_付款结算_优惠券+折扣+全额覆盖.png", "图 3：会员支付确认页示例（优惠券 + 折扣 + 全额覆盖）", width_cm=13.0)
    add_image(doc, ROOT / "miniapp-payment" / "screenshots" / "19_支付成功_优惠券+折扣+不足→补差.png", "图 4：会员支付成功页示例（优惠后补差）", width_cm=13.0)
    add_image(doc, ROOT / "miniapp-payment" / "screenshots" / "23_支付失败.png", "图 5：支付失败页示例", width_cm=13.0)

    add_heading(doc, "四、核心测试用例", 1)
    add_case(
        doc,
        "P01",
        "玩家可从首页浏览游戏并进入详情页",
        "确认玩家到店后可以顺利从终端首页浏览内容，并进入目标游戏详情页。",
        ["打开 PC 点播系统首页。", "观察首页是否展示可玩游戏。", "点击任意一款游戏卡片。", "进入详情页后查看标题、价格、时长、人数等信息。"],
        ["首页可正常展示游戏列表。", "详情页能正常打开。", "价格、时长、支持人数等基础信息可见。"],
        path_text="PC 点播系统首页 → 游戏详情页",
    )
    add_case(
        doc,
        "P02",
        "玩家可选择体验配置并进入支付环节",
        "确认玩家可以从详情页选择时长、人数或设备数，并进入支付确认页。",
        ["在详情页选择体验时长。", "如为多人项目，选择人数或设备数。", "点击“开始游戏”或等效按钮。"],
        ["能进入支付环节。", "支付前展示的项目名称、价格、时长与选择结果一致。"],
        path_text="游戏详情页 → 支付确认页",
    )
    add_case(
        doc,
        "P03",
        "散客扫码支付成功",
        "确认未注册玩家可以直接扫码支付并进入后续开局链路。",
        ["在支付页选择散客支付。", "使用微信或支付宝扫一扫终端二维码。", "在手机侧完成支付。", "观察终端支付完成后的状态变化。"],
        ["散客无需先注册或登录。", "支付成功后，终端不再停留在待支付状态。", "终端进入设备引导、开局准备或等效下一步。"],
        path_text="PC 点播系统支付页 → 微信 / 支付宝扫码 → 支付成功",
        exceptions=["二维码无效或支付失败时，终端要有清晰失败状态或可重试提示。"],
    )
    add_case(
        doc,
        "P04",
        "会员主动扫码支付成功",
        "确认会员主动扫终端码时，小程序能正确展示优惠、预存款、游戏币和补差结果。",
        ["在终端选择会员主动扫码支付。", "使用头号空间小程序扫一扫。", "观察小程序支付确认页中的优惠、预存款、游戏币和补差金额。", "完成支付。"],
        ["小程序能识别会员身份。", "优惠、会员资产和补差金额展示正确。", "支付成功后，终端进入设备引导或开局链路。"],
        path_text="终端支付页 → 小程序扫一扫 → 小程序确认支付",
        exceptions=["余额或游戏币不足时，应正确显示补差。", "支付失败时应停留在可重试状态。"],
    )
    add_case(
        doc,
        "P05",
        "会员码反扫 / 店员协助支付成功",
        "确认玩家出示会员码或由店员协助时，仍可完成正确的会员支付闭环。",
        ["打开玩家小程序会员码。", "由终端或店员侧识别会员码。", "确认支付明细。", "完成支付并观察终端状态。"],
        ["会员身份识别成功。", "支付明细正确。", "支付成功后可继续进入开局链路。"],
        path_text="小程序会员码 → 终端识别 → 支付确认",
    )
    add_case(
        doc,
        "P06",
        "支付成功后进入设备引导或开局状态",
        "确认玩家支付成功后不会断链，而是能继续进入佩戴设备和开局环节。",
        ["完成任意一种支付成功场景。", "观察终端是否展示设备编号、佩戴提示或开局倒计时。", "核对设备编号是否与门店现场安排一致。"],
        ["支付成功后有明确后续动作。", "玩家能知道自己该佩戴哪台设备。", "不会出现支付成功但终端无响应的状态。"],
        path_text="支付成功页 → 设备引导 / 开局准备",
        exceptions=["支付成功但未开局时，应记录订单号并回查门店点播系统订单。"],
    )
    add_case(
        doc,
        "P07",
        "佩戴设备后成功开始体验",
        "确认玩家戴上头显后，可以进入实际体验状态。",
        ["根据终端提示佩戴头显。", "如门店需要店员辅助，完成辅助后等待系统启动。", "观察头显或终端是否进入游戏中状态。"],
        ["玩家能进入实际游戏体验。", "头显不会停留在未启动或错误状态。", "终端或后台可识别当前已开局。"],
        path_text="设备引导 → 佩戴头显 → 开始游玩",
    )
    add_case(
        doc,
        "P08",
        "体验结束后正确返回首页或结束状态",
        "确认游戏到时或结束后，玩家链路能正常收口。",
        ["完成一次完整体验或模拟到时结束。", "观察头显结束状态。", "观察终端是否返回首页、待机页或结束确认状态。"],
        ["体验结束后有明确结束反馈。", "终端回到可继续下一单的状态。", "设备恢复为空闲或可继续服务状态。"],
        path_text="游戏中 → 结束 → 首页 / 待机页",
    )
    add_case(
        doc,
        "P09",
        "新玩家可注册会员并进入会员链路",
        "确认未注册玩家到店后，可以通过门店注册码或小程序完成注册，并在后续使用会员支付。",
        ["扫描门店注册码或打开小程序注册入口。", "填写手机号并完成注册。", "回到终端重新选择游戏。", "验证会员身份可被识别。"],
        ["注册成功。", "后续可按会员身份支付。", "新会员礼包或基础权益展示与配置一致。"],
        path_text="门店注册码 / 小程序 → 注册会员 → 回到终端支付",
    )
    add_case(
        doc,
        "P10",
        "玩家异常场景可被正确识别和记录",
        "确认测试人员知道哪些玩家异常需要记录并回查。",
        ["模拟支付失败。", "模拟支付成功但未开局。", "模拟设备编号错误或引导错误。", "模拟中途退出或结束异常。"],
        ["每种异常都有可记录的现象和订单信息。", "测试人员知道应回查支付结果、门店点播订单、设备绑定关系。"],
        exceptions=["支付失败重点看失败页和终端状态。", "支付成功未开局重点记录订单号、设备号、发生时间。", "设备编号错误重点回查主机与头显绑定关系。"],
    )

    add_heading(doc, "五、通过标准", 1)
    add_bullets(
        doc,
        [
            "至少完成 1 次散客支付闭环测试。",
            "至少完成 1 次会员支付闭环测试。",
            "至少完成 1 次支付成功后的设备引导 / 开局测试。",
            "至少完成 1 次体验结束回流测试。",
            "至少完成 1 次新玩家注册会员测试。",
            "至少记录 1 组异常场景样例。",
        ],
    )

    add_heading(doc, "六、截图清单建议", 1)
    add_table(
        doc,
        ["截图名称", "建议来源"],
        [
            ["点播系统首页", "玩家浏览游戏入口"],
            ["游戏详情页", "价格、时长、人数展示"],
            ["散客支付页", "终端二维码支付页面"],
            ["会员支付确认页", "小程序确认支付页面"],
            ["支付成功页", "支付完成后的状态页"],
            ["支付失败页", "支付失败或取消场景"],
            ["设备引导页", "支付成功后的设备编号 / 佩戴提示"],
            ["体验结束页", "结束后回到首页或待机页的状态"],
            ["会员注册页", "门店注册码或小程序注册入口"],
        ],
        header_fill="FEF3C7",
    )

    doc.save(PLAYER_DOC)
    return PLAYER_DOC


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_official_doc()
    build_merchant_doc()
    build_player_doc()
    print(OFFICIAL_DOC)
    print(MERCHANT_DOC)
    print(PLAYER_DOC)


if __name__ == "__main__":
    main()
