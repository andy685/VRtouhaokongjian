from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "deliverables" / "用户支付PRD_玩家玩游戏支付流程_v2.docx"
OUT = ROOT / "deliverables" / "用户支付PRD_玩家玩游戏支付流程_v3.docx"
FONT_NAME = "Arial Unicode MS"


def apply_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_table(table, header_fill: str = "DBEAFE") -> None:
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if r_idx == 0:
                shade_cell(cell, header_fill)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    apply_font(run, 10.5)


def set_col_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def delete_table(table) -> None:
    tbl = table._element
    tbl.getparent().remove(tbl)


def find_paragraph(doc: Document, text: str):
    for para in doc.paragraphs:
        if para.text.strip() == text:
            return para
    return None


def find_table_by_first_header(doc: Document, header: str):
    for table in doc.tables:
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == header:
            return table
    return None


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_para._p = new_p
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        apply_font(run, 10.5)
    return new_para


def move_table_after(paragraph, table) -> None:
    paragraph._p.addnext(table._tbl)


def optimize() -> Path:
    doc = Document(SRC)

    # 1. Remove repetitive 2.4 section and boundary table
    p_boundary = find_paragraph(doc, "2.4 玩家类型与支付入口边界")
    if p_boundary is not None:
        delete_paragraph(p_boundary)
    t_boundary = find_table_by_first_header(doc, "玩家类型 / 场景")
    if t_boundary is not None:
        delete_table(t_boundary)

    # 2. Replace 3.3 section with payment states instead of repetitive trigger explanation
    p_states = find_paragraph(doc, "3.3 各触发方式说明")
    if p_states is None:
        raise RuntimeError("Could not find section 3.3")
    p_states.text = "3.3 支付流程状态"
    for run in p_states.runs:
        apply_font(run, 13, bold=True, color="1D4ED8")

    t_trigger = find_table_by_first_header(doc, "方式")
    if t_trigger is not None:
        delete_table(t_trigger)

    status_table = doc.add_table(rows=1, cols=3)
    headers = ["状态", "说明", "下一步"]
    for i, h in enumerate(headers):
        status_table.cell(0, i).text = h
    rows = [
        ["待支付", "已生成支付请求，等待玩家扫码或确认支付", "进入正在支付 / 取消支付"],
        ["正在支付", "已发起扣款或第三方支付请求，等待结果返回", "进入支付成功 / 支付失败"],
        ["支付成功", "支付结果确认成功，可继续创单并下发开局", "进入开局流程"],
        ["支付失败", "支付未完成，本次支付链路结束", "允许重新发起支付"],
    ]
    for row in rows:
        cells = status_table.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = v
    set_col_widths(status_table, [1.0, 3.0, 1.6])
    style_table(status_table, "DBEAFE")
    move_table_after(p_states, status_table)

    p_fail = insert_paragraph_after(p_states, style="Heading 3")
    p_fail.text = "支付失败原因"
    for run in p_fail.runs:
        apply_font(run, 11, bold=True, color="0F766E")

    reasons = [
        "支付超时：支付通道或回调在约定时间内未返回结果。",
        "余额不足：微信 / 支付宝支付账户余额不足。",
        "订单已失效：订单超时、已被使用或已被取消。",
        "支付通道异常：第三方支付通道返回异常或网络中断。",
    ]

    anchor = p_fail
    for text in reasons:
        p = insert_paragraph_after(anchor)
        p.style = "List Bullet"
        run = p.add_run(text)
        apply_font(run, 10.5)
        anchor = p

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(optimize())
