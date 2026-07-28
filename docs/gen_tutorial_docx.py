#!/usr/bin/env python3
"""生成头号空间三个角色的使用说明 .docx 文档，含流程图"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Arc
import matplotlib.patches as mpatches
import numpy as np

# 设置中文字体
plt.rcParams['font.sans-serif'] = ['Heiti SC', 'PingFang SC', 'Hiragino Sans GB', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================
#  通用流程图绘制工具
# ============================================================

COLORS = {
    'primary': '#2563EB',    # 蓝色
    'secondary': '#7C3AED',  # 紫色
    'success': '#059669',    # 绿色
    'warning': '#D97706',    # 橙色
    'danger': '#DC2626',     # 红色
    'info': '#0891B2',       # 青色
    'gray': '#6B7280',       # 灰色
    'light': '#F3F4F6',      # 浅灰
    'dark': '#1F2937',       # 深色
    'white': '#FFFFFF',
    'border': '#D1D5DB',
}

def draw_box(ax, x, y, w, h, text, color='primary', fontsize=10, text_color='white', bold=True):
    """绘制圆角矩形框"""
    box = FancyBboxPatch(
        (x - w/2, y - h/2), w, h,
        boxstyle="round,pad=0.08", 
        facecolor=COLORS[color],
        edgecolor=COLORS[color],
        linewidth=1.5, alpha=0.95
    )
    ax.add_patch(box)
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=text_color, fontweight='bold' if bold else 'normal',
            fontfamily='sans-serif')

def draw_arrow(ax, x1, y1, x2, y2, color='#9CA3AF', lw=2):
    """绘制箭头"""
    ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                arrowprops=dict(arrowstyle='->', color=color, lw=lw,
                               connectionstyle='arc3,rad=0'))

def draw_label(ax, x, y, text, fontsize=9, color='#6B7280'):
    """绘制标签文字"""
    ax.text(x, y, text, ha='center', va='center', fontsize=fontsize,
            color=color, fontfamily='sans-serif')

def draw_title_bar(ax, x, y, w, text, color='primary'):
    """绘制标题栏"""
    box = FancyBboxPatch(
        (x - w/2, y - 0.25), w, 0.5,
        boxstyle="round,pad=0.05",
        facecolor=COLORS[color],
        edgecolor=COLORS[color],
        linewidth=1, alpha=0.15
    )
    ax.add_patch(box)
    ax.text(x, y, text, ha='center', va='center', fontsize=13,
            color=COLORS[color], fontweight='bold', fontfamily='sans-serif')

# ============================================================
#  流程图 1: 官方运营 - 完整业务链路
# ============================================================

def draw_official_flow():
    """官方运营完整业务链路 - 横向流程图"""
    fig, ax = plt.subplots(1, 1, figsize=(16, 6))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 6)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    # 标题
    draw_title_bar(ax, 8, 5.6, 6, '官方运营 · 业务链路全景图', 'primary')

    steps = [
        (1.5, 4.0, '① 新建商家', 'primary'),
        (3.8, 4.0, '② 新建店铺', 'primary'),
        (6.1, 4.0, '③ 录入游戏', 'info'),
        (8.4, 4.0, '④ 录入主机\n生成Token', 'secondary'),
        (10.7, 4.0, '⑤ 录入头显\n绑定主机', 'secondary'),
        (13.0, 4.0, '⑥ 内容分发', 'success'),
        (14.5, 2.5, '⑦ 配置\n拉卡拉', 'warning'),
    ]

    for x, y, text, color in steps:
        w = 2.0
        h = 1.5 if '\n' in text else 1.2
        # 判断是否两行
        lines = text.split('\n')
        if len(lines) == 1:
            draw_box(ax, x, y, w, 1.2, text, color, fontsize=9)
        else:
            draw_box(ax, x, y, w, 1.6, text, color, fontsize=8)

    # 横向箭头
    for i in range(len(steps) - 2):
        x1 = steps[i][0] + 1.0
        y1 = steps[i][1]
        x2 = steps[i+1][0] - 1.0
        y2 = steps[i+1][1]
        draw_arrow(ax, x1, y1, x2, y2, '#9CA3AF', lw=2.5)

    # ⑥→⑦ 斜箭头
    draw_arrow(ax, 13.0, 3.3, 14.2, 3.0, '#9CA3AF', lw=2.5)

    # 底部说明
    draw_label(ax, 8, 1.0, '商家后台地址 + 管理员账号交付商家 → 店铺上线 → 游戏入库 → 设备就绪 → 门店下载游戏 → 支持收款',
               fontsize=9, color='#6B7280')

    # 图例
    draw_box(ax, 2.5, 2.0, 1.8, 0.5, '商家/店铺', 'primary', fontsize=7)
    draw_box(ax, 5.5, 2.0, 1.8, 0.5, '内容', 'info', fontsize=7)
    draw_box(ax, 8.5, 2.0, 1.8, 0.5, '设备', 'secondary', fontsize=7)
    draw_box(ax, 11.5, 2.0, 1.8, 0.5, '分发/支付', 'success', fontsize=7)

    path = os.path.join(OUTPUT_DIR, 'flow_official_main.png')
    plt.tight_layout(pad=0.5)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 2: 官方运营 - 新建商家三步
# ============================================================

def draw_official_merchant_steps():
    """新建商家三步流程"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 7, 4.6, 10, '新建商家 · 三步流程', 'primary')

    steps_data = [
        (2.3, 2.5, 'Step 1\n基本信息', ['商家名称', '联系人/电话', '负责区域', '手续费率', '商家状态']),
        (7.0, 2.5, 'Step 2\n管理员账号', ['管理员账号', '管理员密码', '→ 交付商家', '登录后台+收银']),
        (11.7, 2.5, 'Step 3\n提现账户', ['开户银行', '银行卡号', '开户人姓名', '身份证号']),
    ]

    for x, y, title, fields in steps_data:
        draw_box(ax, x, y + 0.6, 3.2, 0.8, title, 'primary', fontsize=10)
        # 字段列表
        for j, f in enumerate(fields):
            color_tag = 'warning' if '→' in f else 'gray'
            draw_box(ax, x, y - 0.4 - j * 0.55, 3.0, 0.4, f, 'light',
                     fontsize=7, text_color=COLORS['dark'], bold=False)

    # 箭头
    draw_arrow(ax, 4.0, 2.5, 5.2, 2.5, COLORS['primary'], lw=3)
    draw_arrow(ax, 8.7, 2.5, 9.9, 2.5, COLORS['primary'], lw=3)

    # 底部标签
    draw_label(ax, 7, 0.5, '完成后：商家获得独立后台访问权限 + PC收银系统登录权限',
               fontsize=10, color=COLORS['success'])

    path = os.path.join(OUTPUT_DIR, 'flow_official_merchant.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 3: 官方运营 - 设备录入绑定
# ============================================================

def draw_official_device():
    """主机+头显 录入绑定流程"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 5.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 5.5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 7, 5.2, 10, '设备录入与绑定流程', 'secondary')

    # 主机线
    draw_box(ax, 2.0, 3.6, 2.5, 1.0, '① 录入主机\n编号·MAC·类型', 'secondary', fontsize=8)
    draw_arrow(ax, 3.4, 3.6, 5.0, 3.6, COLORS['secondary'], lw=2.5)
    draw_box(ax, 6.5, 3.6, 2.5, 1.0, '② 分配到\n商家+店铺', 'secondary', fontsize=8)
    draw_arrow(ax, 7.9, 3.6, 9.5, 3.6, COLORS['secondary'], lw=2.5)
    draw_box(ax, 11.0, 3.6, 2.2, 1.0, '③ 生成\nToken', 'danger', fontsize=9)

    # 头显线
    draw_box(ax, 2.0, 1.5, 2.5, 1.0, '④ 录入头显\nSN·型号·版本', 'info', fontsize=8)
    draw_arrow(ax, 3.4, 1.5, 5.0, 1.5, COLORS['info'], lw=2.5)
    draw_box(ax, 6.5, 1.5, 2.5, 1.0, '⑤ 分配到\n商家+店铺', 'info', fontsize=8)
    draw_arrow(ax, 7.9, 1.5, 9.5, 1.5, COLORS['info'], lw=2.5)
    draw_box(ax, 11.0, 1.5, 2.2, 1.0, '⑥ 绑定\n到主机', 'success', fontsize=9)

    # 竖线连接
    draw_arrow(ax, 11.0, 2.6, 11.0, 3.0, '#9CA3AF', lw=1.5)
    draw_label(ax, 12.3, 2.8, '一台主机\n可绑定\n多台头显', fontsize=8, color='#6B7280')

    # 关键词
    draw_box(ax, 7, 0.3, 12, 0.5, 'Token 是 PC点播系统激活凭证，一个主机一个Token，不可共用', 'warning',
             fontsize=8, text_color=COLORS['dark'])

    path = os.path.join(OUTPUT_DIR, 'flow_official_device.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 4: 商家门店 - 运营全流程
# ============================================================

def draw_shop_flow():
    """商家门店运营全流程"""
    fig, ax = plt.subplots(1, 1, figsize=(16, 7))
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 7)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 8, 6.6, 12, '商家门店 · 运营全流程', 'success')

    # 第一行
    top = [
        (1.5, 5.2, '① 登录\n后台', 'primary'),
        (3.5, 5.2, '② 查看/\n绑定设备', 'secondary'),
        (5.5, 5.2, '③ 设置\n设备密码', 'secondary'),
        (7.5, 5.2, '④ 配置\n商品价格', 'info'),
        (9.5, 5.2, '⑤ 创建\n会员体系', 'info'),
        (11.5, 5.2, '⑥ 创建\n优惠券', 'success'),
        (13.5, 5.2, '⑦ 收银\n接客', 'warning'),
    ]

    for x, y, text, color in top:
        draw_box(ax, x, y, 1.7, 1.2, text, color, fontsize=8)

    for i in range(len(top) - 1):
        draw_arrow(ax, top[i][0] + 0.9, top[i][1], top[i+1][0] - 0.9, top[i+1][1], '#9CA3AF', lw=2)

    # 收银展开
    cashier_items = [
        (5.3, 2.5, '散客收银\n微信/支付宝/现金', 'primary'),
        (8.0, 2.5, '会员收银\n储值+游戏币+优惠券', 'info'),
        (10.7, 2.5, '充值/退款', 'warning'),
        (13.3, 2.5, '营收查看\n交班对账', 'success'),
    ]

    for x, y, text, color in cashier_items:
        w = 2.3
        h = 1.5
        draw_box(ax, x, y, w, h, text, color, fontsize=8)

    for i in range(len(cashier_items) - 1):
        draw_arrow(ax, cashier_items[i][0] + 1.2, cashier_items[i][1],
                   cashier_items[i+1][0] - 1.2, cashier_items[i+1][1], '#9CA3AF', lw=2)

    # 从收银接客到散客收银的竖线
    draw_arrow(ax, 13.5, 4.5, 13.3, 3.3, '#9CA3AF', lw=2)

    # 玩家端虚线框
    rect = FancyBboxPatch((10.5, 0.3), 4.5, 1.3, boxstyle="round,pad=0.1",
                          facecolor='none', edgecolor=COLORS['primary'], 
                          linewidth=1.5, linestyle='--')
    ax.add_patch(rect)
    ax.text(12.75, 1.2, '玩家端流程', ha='center', va='center', fontsize=8,
            color=COLORS['primary'], fontweight='bold')
    ax.text(12.75, 0.7, 'PC终端选游戏 → 扫码支付 → 佩戴体验',
            ha='center', va='center', fontsize=7.5, color=COLORS['dark'])

    path = os.path.join(OUTPUT_DIR, 'flow_shop_main.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 5: 商家门店 - 会员支付抵扣优先级
# ============================================================

def draw_shop_payment_priority():
    """会员消费抵扣优先级"""
    fig, ax = plt.subplots(1, 1, figsize=(14, 3.5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 3.5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 7, 3.2, 12, '会员消费 · 自动抵扣优先级', 'info')

    priorities = [
        (1.4, 1.5, '① 优惠券\n自动匹配', 'success'),
        (3.7, 1.5, '② 会员折扣\n等级折扣', 'info'),
        (6.0, 1.5, '③ 预存款\n人民币余额', 'primary'),
        (8.3, 1.5, '④ 游戏币\n1币 = 1元', 'secondary'),
        (10.6, 1.5, '⑤ 微信补差\n不足部分', 'warning'),
    ]

    for x, y, text, color in priorities:
        draw_box(ax, x, y, 2.0, 1.2, text, color, fontsize=8)

    for i in range(len(priorities) - 1):
        draw_arrow(ax, priorities[i][0] + 1.0, priorities[i][1],
                   priorities[i+1][0] - 1.0, priorities[i+1][1], '#DC2626', lw=2.5)

    # 底部标签
    draw_label(ax, 7, 0.3, '→ 先花了再补：优惠券先扣 → 折扣先算 → 余额先扣 → 游戏币补齐 → 最后微信买单',
               fontsize=9, color='#6B7280')

    path = os.path.join(OUTPUT_DIR, 'flow_shop_payment.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 6: 玩家 - 体验全流程
# ============================================================

def draw_player_flow():
    """玩家体验全流程"""
    fig, ax = plt.subplots(1, 1, figsize=(15, 4.5))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 4.5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 7.5, 4.2, 11, '玩家体验 · 全流程图', 'primary')

    steps = [
        (1.2, 2.5, '① 到店\nPC终端首页', 'primary'),
        (3.3, 2.5, '② 选游戏\n浏览+详情', 'info'),
        (5.4, 2.5, '③ 选配置\n时长+人数', 'info'),
        (7.5, 2.5, '④ 扫码\n支付', 'success'),
        (9.6, 2.5, '⑤ 佩戴\nVR头显', 'warning'),
        (11.7, 2.5, '⑥ 沉浸\n游戏中', 'secondary'),
        (13.8, 2.5, '⑦ 结束\n返回首页', 'gray'),
    ]

    for x, y, text, color in steps:
        draw_box(ax, x, y, 1.8, 1.2, text, color, fontsize=8)

    for i in range(len(steps) - 1):
        draw_arrow(ax, steps[i][0] + 0.9, steps[i][1],
                   steps[i+1][0] - 0.9, steps[i+1][1], '#9CA3AF', lw=2.5)

    # 续费回路
    ax.annotate('', xy=(11.1, 1.1), xytext=(10.2, 1.1),
                arrowprops=dict(arrowstyle='->', color=COLORS['warning'], lw=1.5,
                               connectionstyle='arc3,rad=0.4'))
    draw_label(ax, 10.65, 0.7, '续费加时', fontsize=7.5, color=COLORS['warning'])

    # 支付方式标注
    draw_label(ax, 7.5, 1.2, '支持: 微信 · 支付宝 · 会员储值 · 会员码',
               fontsize=8, color='#6B7280')

    # 底部提示
    draw_label(ax, 7.5, 0.3, 'VR头显内无任何UI → 全沉浸 → 摘盔自动暂停 → 超3分钟自动结束',
               fontsize=8, color='#6B7280')

    path = os.path.join(OUTPUT_DIR, 'flow_player_main.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  流程图 7: 玩家 - 支付方式决策树
# ============================================================

def draw_player_payment():
    """玩家支付方式决策树"""
    fig, ax = plt.subplots(1, 1, figsize=(13, 5.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 5.5)
    ax.axis('off')
    ax.set_facecolor('#FAFBFC')
    fig.patch.set_facecolor('#FAFBFC')

    draw_title_bar(ax, 6.5, 5.2, 10, '支付方式选择', 'success')

    # 起始
    draw_box(ax, 6.5, 4.0, 3.0, 0.7, 'PC终端展示二维码', 'dark', fontsize=9)

    # 分支竖线
    ax.plot([6.5, 6.5], [3.6, 3.1], color='#9CA3AF', lw=2)
    ax.plot([6.5, 3.0], [3.1, 3.1], color='#9CA3AF', lw=2)
    ax.plot([6.5, 10.0], [3.1, 3.1], color='#9CA3AF', lw=2)
    ax.plot([3.0, 3.0], [3.1, 2.8], color='#9CA3AF', lw=2)
    ax.plot([10.0, 10.0], [3.1, 2.8], color='#9CA3AF', lw=2)

    # 是否会员
    draw_label(ax, 4.7, 3.15, '散客', fontsize=9, color=COLORS['gray'])
    draw_label(ax, 8.3, 3.15, '会员', fontsize=9, color=COLORS['primary'])

    # 散客
    draw_box(ax, 3.0, 2.2, 3.5, 1.0, '微信/支付宝 扫码\n→ 确认金额 → 支付成功',
             'info', fontsize=8)

    # 会员 - 两条路
    draw_box(ax, 8.0, 2.2, 2.5, 1.0, '小程序\n扫码支付', 'primary', fontsize=8)
    draw_box(ax, 11.0, 2.2, 2.5, 1.0, '出示会员码\n终端反扫', 'secondary', fontsize=8)

    # 箭头
    ax.annotate('', xy=(8.0, 2.8), xytext=(10.0, 3.1),
                arrowprops=dict(arrowstyle='->', color='#9CA3AF', lw=1.5))
    ax.annotate('', xy=(11.0, 2.8), xytext=(10.0, 3.1),
                arrowprops=dict(arrowstyle='->', color='#9CA3AF', lw=1.5))

    # 抵扣顺序
    draw_box(ax, 9.5, 1.2, 6.0, 1.0,
             '抵扣顺序: 优惠券 → 会员折扣 → 预存款 → 游戏币 → 微信补差',
             'light', fontsize=8, text_color=COLORS['dark'], bold=False)

    # 注册入口
    draw_box(ax, 3.0, 0.6, 3.5, 0.7, '非会员? 扫码注册 → 送游戏币+预存款',
             'warning', fontsize=7.5)

    path = os.path.join(OUTPUT_DIR, 'flow_player_payment.png')
    plt.tight_layout(pad=0.3)
    plt.savefig(path, dpi=180, bbox_inches='tight', facecolor='#FAFBFC')
    plt.close()
    return path

# ============================================================
#  通用 docx 样式设置
# ============================================================

def setup_styles(doc):
    """设置 Word 文档样式"""
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.35
    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Microsoft YaHei'
        h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        h_style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        if level == 1:
            h_style.font.size = Pt(22)
            h_style.font.bold = True
        elif level == 2:
            h_style.font.size = Pt(16)
            h_style.font.bold = True
        else:
            h_style.font.size = Pt(13)
            h_style.font.bold = True

def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
        # 蓝色背景
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2563EB"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后间距
    return table

def add_path_hint(doc, text):
    """添加带引用的操作路径提示"""
    p = doc.add_paragraph()
    run = p.add_run('▸ 操作路径：')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    run.font.italic = True

def add_note(doc, text):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run('⚠ ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)

def add_image_centered(doc, path, width_inches=6.0):
    """居中插入图片"""
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
        doc.add_paragraph()

# ============================================================
#  文档 1: 官方运营篇
# ============================================================

def build_official_docx():
    flow_main = draw_official_flow()
    flow_merchant = draw_official_merchant_steps()
    flow_device = draw_official_device()

    doc = Document()
    setup_styles(doc)

    # ===== 封面区域 =====
    doc.add_heading('头号空间 · 官方运营使用说明', level=1)
    p = doc.add_paragraph()
    run = p.add_run('适用版本：v2.19  |  更新日期：2026年7月16日  |  目标角色：平台超管、运营人员')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    add_image_centered(doc, flow_main, 6.0)
    doc.add_paragraph()

    # ===== 1. 角色定位 =====
    doc.add_heading('1. 角色定位与工作台', level=2)
    doc.add_heading('你的职责', level=3)
    duties = [
        '商家管理：审核入驻、开通商家后台',
        '店铺管理：为商家创建门店',
        '内容管理：录入和管理VR游戏库、审核CP上传的游戏',
        '设备管理：录入主机和头显、生成Token、绑定设备',
        '内容分发：把游戏推送到指定门店',
        '支付配置：开通拉卡拉支付，配置分账比例',
        '全局监控：查看全平台数据、处理财务结算',
    ]
    for d in duties:
        p = doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('登录工作台', level=3)
    doc.add_paragraph('1. 打开总运营后台地址，输入平台超管账号密码登录')
    doc.add_paragraph('2. 登录后进入「大屏看板」，查看全平台数据总览：订单数、营收、活跃设备、会员增长')
    doc.add_paragraph('3. 左侧菜单可访问所有管理模块')

    # ===== 2. 新建商家 =====
    doc.add_heading('2. 新建商家', level=2)
    add_path_hint(doc, '左侧菜单 → 店铺管理 → 商家管理 → 点击「新增商家」')
    doc.add_paragraph('商家是平台中的经营主体，每个商家拥有独立的商家后台。创建流程分三步：')

    add_image_centered(doc, flow_merchant, 5.8)

    doc.add_heading('Step 1：基本信息', level=3)
    add_styled_table(doc,
        ['字段', '说明'],
        [['商家名称', '如「恒然集团」「幻影星空」'],
         ['联系人', '商家主要负责人姓名'],
         ['联系电话', '商家联系电话'],
         ['负责区域', '深圳/广州/北京/上海/成都/杭州/武汉'],
         ['对应代理商', '如该商家通过代理商入驻，选择对应代理商'],
         ['提现手续费率', '默认 0.5%（0.005），可调整'],
         ['商家状态', '正常（审核通过）/ 待审核 / 停用']],
        [4, 8])

    doc.add_heading('Step 2：管理员账号', level=3)
    add_styled_table(doc,
        ['字段', '说明'],
        [['管理员账号', '商家后台登录用的用户名'],
         ['管理员密码', '商家后台登录密码']],
        [4, 8])
    add_note(doc, '此账号密码需要交付给商家，用于登录商家后台和PC收银系统。')

    doc.add_heading('Step 3：提现账户', level=3)
    add_styled_table(doc,
        ['字段', '说明'],
        [['开户银行', '工商/建设/农业/中国/招商等'],
         ['银行卡号', '最多23位'],
         ['开户人姓名', '对公或对私账户姓名'],
         ['身份证号', '开户人身份证']],
        [4, 8])

    p = doc.add_paragraph()
    run = p.add_run('完成后：商家获得独立的后台访问权限和PC收银系统登录权限。')
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)

    # ===== 3. 新建店铺 =====
    doc.add_heading('3. 新建店铺', level=2)
    add_path_hint(doc, '左侧菜单 → 店铺管理 → 店铺列表 → 点击「新增店铺」')
    doc.add_paragraph('一个商家下可以有多个门店（如「福田旗舰店」「南山科技园店」）。')

    add_styled_table(doc,
        ['字段', '说明'],
        [['店铺名称', '如「深圳福田旗舰店」'],
         ['所属商家', '下拉选择已创建的商家'],
         ['所属区域', '下拉选择城市'],
         ['详细地址', '门店实际地址'],
         ['联系电话', '门店联系电话'],
         ['店铺状态', '营业中 / 已打烊 / 维护中'],
         ['开通支付', '开启后生成拉卡拉支付二维码']],
        [4, 8])

    doc.add_paragraph('注册码：店铺创建后自动生成，格式 REG-XXXX-XXXX，顾客扫码注册成为该店会员。')
    doc.add_paragraph('注册规则：可设置新注册用户自动赠送游戏币和预存款。')

    # ===== 4. 录入游戏 =====
    doc.add_heading('4. 录入游戏', level=2)
    add_path_hint(doc, '左侧菜单 → 内容中心 → 游戏库 → 点击「添加游戏」')
    doc.add_paragraph('在游戏库中录入可供门店使用的VR游戏内容。')

    add_styled_table(doc,
        ['字段', '说明'],
        [['游戏名称', '如「节奏光剑」「半衰期：爱莉克斯」'],
         ['游戏类型', '联机/单机'],
         ['付费模式', '按次 / 按时长'],
         ['体验时长', '单次体验分钟数'],
         ['游戏豆消耗', '每次体验消耗的游戏豆数量'],
         ['推荐价格', '平台建议零售价'],
         ['运行平台', '头显一体机 / 主机游戏'],
         ['状态', '已上线 / 草稿 / 待审核 / 已下架'],
         ['供应商', '选择CP游戏供应商']],
        [4, 8])
    add_note(doc, '游戏还可以来自CP供应商后台自行上传，运营人员在「游戏审核」中审核通过后入库。')

    # ===== 5. 录入主机与Token =====
    doc.add_heading('5. 录入主机设备与生成Token', level=2)
    add_path_hint(doc, '左侧菜单 → 数据中心 → 设备配置管理 → 主机管理 → 点击「录入主机」')
    doc.add_paragraph('这是整个链路中最关键的步骤，主机是承载PC点播系统的电脑设备。')

    add_image_centered(doc, flow_device, 5.8)

    doc.add_heading('5.1 录入主机', level=3)
    add_styled_table(doc,
        ['字段', '说明'],
        [['主机编号', '如 PCT-001'],
         ['主机名称', '如「大厅主机 #01」'],
         ['设备类型', '悬浮骑兵/暗黑行者/暗黑机甲/幻影飞碟/通用主机'],
         ['硬件配置', 'CPU/内存/存储/显卡'],
         ['系统版本', '如 Windows 11 Kiosk v2.1'],
         ['MAC地址', '网卡物理地址（必填）']],
        [4, 8])

    doc.add_heading('5.2 分配与生成Token', level=3)
    doc.add_paragraph('1. 在主机列表中找到该主机，点击「分配」')
    doc.add_paragraph('2. 选择所属商家 → 选择所属店铺')
    doc.add_paragraph('3. 分配成功后点击「生成Token」')
    add_note(doc, 'Token是关键：PC点播系统启动时需要输入此Token完成设备认证和绑定。一个主机一个Token，不可共用。')

    # ===== 6. 录入头显 =====
    doc.add_heading('6. 录入头显设备与绑定主机', level=2)
    add_path_hint(doc, '左侧菜单 → 数据中心 → 设备配置管理 → 头显管理 → 点击「录入头显」')

    doc.add_heading('6.1 录入头显', level=3)
    add_styled_table(doc,
        ['字段', '说明'],
        [['头显名称', '如「Pico 4 #01」'],
         ['设备型号', 'Pico 4 / Pico 4 Ultra / Meta Quest 3 等'],
         ['SN码', '设备唯一序列号（必填）'],
         ['固件版本', '如 v5.9.0'],
         ['备注', '补充说明']],
        [4, 8])

    doc.add_heading('6.2 分配与绑定', level=3)
    doc.add_paragraph('1. 点击「分配」→ 选择商家和店铺')
    doc.add_paragraph('2. 点击「绑定到主机」→ 选择同店铺下在线的主机')
    doc.add_paragraph('3. 一台主机可绑定多台头显')

    # ===== 7. 内容分发 =====
    doc.add_heading('7. 内容分发（游戏推送到门店）', level=2)
    add_path_hint(doc, '左侧菜单 → 内容中心 → 内容分发')
    doc.add_paragraph('游戏录入到平台后，需要通过「内容分发」推送到指定门店，门店的PC点播系统才会自动下载。')

    doc.add_heading('7.1 分发操作', level=3)
    doc.add_paragraph('1. 切换到「游戏列表」Tab')
    doc.add_paragraph('2. 找到目标游戏（状态为「未分发」），点击「分发」')
    doc.add_paragraph('3. 在弹窗中选择目标店铺和分发方式：')
    doc.add_paragraph('    智能分发（推荐）：仅下发变更的构件文件，速度快', style='List Bullet')
    doc.add_paragraph('    全量分发：重新下发全部文件（首次分发或修复用）', style='List Bullet')

    doc.add_heading('7.2 分发状态说明', level=3)
    add_styled_table(doc,
        ['状态', '含义'],
        [['未分发', '游戏未被推送到任何门店'],
         ['分发中', '正在向门店推送游戏文件'],
         ['已分发', '推送完成，门店点播系统已可下载'],
         ['有更新', '游戏有新版本，可重新分发'],
         ['失败', '分发过程出错，可重试'],
         ['已撤回', '运营主动撤销了分发']],
        [4, 8])

    # ===== 8. 拉卡拉 =====
    doc.add_heading('8. 配置拉卡拉支付', level=2)
    add_path_hint(doc, '左侧菜单 → 平台财务 → 拉卡拉配置')
    doc.add_paragraph('配置拉卡拉商户号、终端号和证书后，门店才能接收微信/支付宝付款。需要为平台、商家、代理商、CP分别配置各自的商户信息，包括：')
    doc.add_paragraph('商户号 / 终端号 / 机构号', style='List Bullet')
    doc.add_paragraph('支付证书上传', style='List Bullet')
    doc.add_paragraph('分账比例设置', style='List Bullet')

    # ===== 9. 其他功能 =====
    doc.add_heading('9. 其他运营功能', level=2)
    add_styled_table(doc,
        ['模块', '功能'],
        [['会员中心', '查看全平台会员数据、储值流水、游戏币流水、消费排行'],
         ['订单流水', '收银订单、点播订单、储值变更、游戏币兑换、活动赠送记录'],
         ['平台财务', '营收总览、商家/代理商/CP分账结算、异常订单处理'],
         ['平台账号', '管理系统账号和RBAC角色权限'],
         ['系统运维', '版本发布、告警中心、操作日志、工单系统'],
         ['公告通知', '向全平台或指定商家推送公告']],
        [4, 10])

    # ===== 附录 =====
    doc.add_heading('附录：常见问题与术语表', level=2)
    doc.add_heading('常见问题', level=3)
    
    faqs = [
        ('Q：商家反馈登录不了后台？', 'A：检查商家状态是否为「正常」，确认管理员账号密码是否正确。'),
        ('Q：内容分发后门店还是看不到游戏？', 'A：确认分发状态为「已完成」，检查门店PC点播系统是否在线并完成Token配置。'),
        ('Q：设备怎么回收或重新分配？', 'A：在设备配置管理中，找到设备点击「回收」，可解除分配后再重新分配到其他店铺。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.bold = True
        p = doc.add_paragraph(a)

    doc.add_heading('术语表', level=3)
    add_styled_table(doc,
        ['术语', '解释'],
        [['Token', '设备认证令牌，PC点播系统启动时用于绑定店铺'],
         ['游戏豆', '商家端运营代币，用于采购和管理游戏'],
         ['游戏币', '会员端消费代币，1:1抵扣消费金额'],
         ['预存款', '会员人民币储值余额'],
         ['内容分发', '将游戏从平台推送到指定门店'],
         ['RBAC', '基于角色的权限控制系统'],
         ['拉卡拉', '第三方支付服务商，处理微信/支付宝支付']],
        [3, 10])

    out_path = os.path.join(OUTPUT_DIR, '头号空间-使用教程-官方运营篇.docx')
    doc.save(out_path)
    print(f'[OK] {out_path}')
    return out_path

# ============================================================
#  文档 2: 商家门店篇
# ============================================================

def build_shop_docx():
    flow_main = draw_shop_flow()
    flow_payment = draw_shop_payment_priority()

    doc = Document()
    setup_styles(doc)

    doc.add_heading('头号空间 · 商家门店使用说明', level=1)
    p = doc.add_paragraph()
    run = p.add_run('适用版本：v2.19  |  更新日期：2026年7月16日  |  目标角色：店长、收银员、门店管理员')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    add_image_centered(doc, flow_main, 6.0)
    doc.add_paragraph()

    # ===== 1. 角色定位 =====
    doc.add_heading('1. 角色定位与业务全景', level=2)
    doc.add_heading('你的职责', level=3)
    duties = [
        '设备管理：查看已分配设备、绑定头显、设置管理密码',
        '商品定价：设置单次消费项目价格和充值套餐',
        '会员运营：创建会员等级、设置折扣、发优惠券',
        '收银操作：散客收银、会员收银、充值与退款',
        '营收管理：查看营收数据、交接班对账',
    ]
    for d in duties:
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('你使用的系统', level=3)
    add_styled_table(doc,
        ['系统', '用途', '入口'],
        [['商家后台 (Web)', '设备管理、商品配置、会员管理、营收查看', '浏览器访问'],
         ['PC收银系统 (桌面)', '日常收银、会员识别、支付操作', '门店电脑桌面']],
        [4, 5, 5])
    add_note(doc, '两个系统使用同一套账号密码（由官方分配）。')

    # ===== 2. 登录 =====
    doc.add_heading('2. 登录商家后台', level=2)
    doc.add_paragraph('1. 使用官方分配的管理员账号和密码登录商家后台')
    doc.add_paragraph('2. 登录后进入「今日概况」，可查看当日营收、订单数、活跃设备')
    doc.add_paragraph('3. 左侧菜单可进入各管理模块')

    # ===== 3. 设备 =====
    doc.add_heading('3. 管理店铺设备', level=2)
    add_path_hint(doc, '系统设置 → 设备列表')

    doc.add_heading('查看主机设备', level=3)
    doc.add_paragraph('Tab「主机设备」：查看主机编号、名称、MAC地址、状态（在线/离线/故障）', style='List Bullet')
    doc.add_paragraph('可通过筛选快速定位', style='List Bullet')
    doc.add_paragraph('每台主机显示其绑定的头显数量', style='List Bullet')

    doc.add_heading('绑定头显到主机', level=3)
    doc.add_paragraph('1. 在主机行点击「绑定头显」')
    doc.add_paragraph('2. 输入头显 SN 码搜索')
    doc.add_paragraph('3. 点击确认完成绑定')

    doc.add_heading('查看头显设备', level=3)
    doc.add_paragraph('Tab「头显设备」：查看所有头显，状态包括空闲/使用中/充电/离线/故障', style='List Bullet')
    doc.add_paragraph('点击头显可查看详情（电量、瞳距IPD等）', style='List Bullet')

    # ===== 4. 密码 =====
    doc.add_heading('4. 设置设备密码', level=2)
    add_path_hint(doc, '系统设置 → 设备列表 → 主机设备 → 点击「修改点播系统密码」')
    doc.add_paragraph('这是保护PC点播系统不被顾客随意操作的安全措施。设置密码后，任何试图进入点播系统管理模式的操作都需要验证密码。')

    add_styled_table(doc,
        ['操作', '说明'],
        [['设置密码', '输入新密码并确认，立即生效'],
         ['修改密码', '输入旧密码后设置新密码'],
         ['清除密码', '清空密码，不再需要验证（不推荐）']],
        [4, 10])
    add_note(doc, '建议每台主机都设置独立的密码，并记录在案。')

    # ===== 5. 商品价格 =====
    doc.add_heading('5. 配置商品和价格', level=2)

    doc.add_heading('单次消费项目', level=3)
    add_path_hint(doc, '商品管理 → 单次消费项目')
    add_styled_table(doc,
        ['项目名称', '时长', '价格', '说明'],
        [['30分钟体验', '30分钟', '¥59', '标准单人体验'],
         ['60分钟畅玩', '60分钟', '¥99', '超值畅玩套餐'],
         ['单游戏体验券', '一局', '¥39', '指定游戏一局']],
        [3.5, 2.5, 2.5, 5])

    doc.add_heading('充值套餐', level=3)
    add_path_hint(doc, '运营管理 → 充值套餐')
    add_styled_table(doc,
        ['套餐名称', '充值金额', '赠送金额', '实到余额'],
        [['青铜卡', '¥100', '¥10', '¥110'],
         ['白银卡', '¥300', '¥50', '¥350'],
         ['黄金卡', '¥500', '¥120', '¥620'],
         ['钻石卡', '¥1000', '¥300', '¥1300']],
        [3, 3, 3, 3])

    # ===== 6. 会员 =====
    doc.add_heading('6. 创建会员体系', level=2)

    doc.add_heading('设置会员等级', level=3)
    add_path_hint(doc, '会员管理 → 会员级别')
    add_styled_table(doc,
        ['等级', '折扣率', '权益建议'],
        [['普通会员', '9.5折', '基础折扣'],
         ['金卡会员', '9折', '专属折扣 + 优先预约'],
         ['钻石会员', '8.5折', '最高折扣 + 生日礼 + 专属通道']],
        [3, 3, 7])

    doc.add_heading('新会员礼包', level=3)
    add_path_hint(doc, '会员管理 → 设置会员折扣')
    doc.add_paragraph('可配置新注册会员自动赠送游戏币（如50币）和预存款（如¥10）。', style='List Bullet')

    # ===== 7. 优惠券 =====
    doc.add_heading('7. 创建优惠券', level=2)
    add_path_hint(doc, '运营管理 → 优惠券')

    add_styled_table(doc,
        ['类型', '示例', '使用场景'],
        [['折扣券', '满39元打8折', '新店开业促销'],
         ['满减券', '满100减20', '鼓励高消费'],
         ['代金券', '立减15元', '会员日福利'],
         ['体验券', '免费体验1次', '拉新引流']],
        [3, 4, 5])

    doc.add_paragraph('创建时需设置有效期、使用门槛（最低消费金额或指定游戏）、适用店铺范围。')

    # ===== 8. 收银 =====
    doc.add_heading('8. 收银系统日常操作', level=2)

    doc.add_heading('8.1 散客收银', level=3)
    doc.add_paragraph('顾客没有注册会员，使用微信/支付宝/现金支付。')
    doc.add_paragraph('1. 点击「加入商品」或直接选择VR体验项目')
    doc.add_paragraph('2. 输入体验人数和时长，系统自动计算金额')
    doc.add_paragraph('3. 在右侧支付区域选择支付方式：现金直接确认 / 微信支付宝扫码枪扫描付款码')
    doc.add_paragraph('4. 支付成功后弹出成功页面，引导顾客佩戴设备')

    doc.add_heading('8.2 会员收银', level=3)
    doc.add_paragraph('顾客已注册会员，使用会员储值或游戏币抵扣。')
    doc.add_paragraph('1. 点击「会员」按钮，输入会员手机号或扫描会员码')
    doc.add_paragraph('2. 系统自动识别会员身份，显示可用余额和游戏币')
    doc.add_paragraph('3. 加入商品，系统按固定优先级自动抵扣：')

    add_image_centered(doc, flow_payment, 5.5)

    doc.add_paragraph('4. 确认扣费，系统从会员账户扣款')
    doc.add_paragraph('5. 余额不足时，自动计算需要微信补差的金额')

    doc.add_heading('8.3 充值', level=3)
    doc.add_paragraph('路径：会员页面 → 选择会员 → 充值。选择预设套餐，确认后充值到会员预存款账户。')

    doc.add_heading('8.4 退款', level=3)
    doc.add_paragraph('1. 在订单列表中找到需要退款的订单')
    doc.add_paragraph('2. 点击退款，选择退款原因：客户不想要了 / 游戏体验异常 / 操作错误 / 设备故障 / 重复付款 / 其他')
    doc.add_paragraph('3. 确认退款金额（仅支持整单退款）')
    doc.add_paragraph('4. 已结算订单需上传退款凭证（转账截图等，1-3张）')

    # ===== 9. 营收 =====
    doc.add_heading('9. 营收查看与交班对账', level=2)

    doc.add_heading('营收查看', level=3)
    add_path_hint(doc, '营收 → 按收银员 / 按支付方式')
    doc.add_paragraph('按收银员：每位店员的收款明细', style='List Bullet')
    doc.add_paragraph('按支付方式：现金/微信/支付宝/会员储值分类统计', style='List Bullet')
    doc.add_paragraph('按时间段：小时/日/周/月趋势图', style='List Bullet')

    doc.add_heading('交班对账', level=3)
    add_path_hint(doc, '交班 → 交班对账')
    doc.add_paragraph('1. 查看本班次营收汇总')
    doc.add_paragraph('2. 复核现金金额（与实际现金核对）')
    doc.add_paragraph('3. 填写备注（如有差异说明原因）')
    doc.add_paragraph('4. 系统自动生成交班记录')
    add_note(doc, '交班后本班次数据锁定，无法修改。')

    # ===== 附录 =====
    doc.add_heading('附录：常见问题与术语表', level=2)
    
    faqs = [
        ('Q：PC点播系统显示「设备未激活」？', 'A：联系官方运营确认Token是否已生成。在商家后台 → 设备列表 → 主机设备中可查看Token。'),
        ('Q：支付失败怎么处理？', 'A：系统会显示具体原因（网络超时/余额不足/支付被拒/系统繁忙），根据提示重试或引导顾客更换支付方式。'),
        ('Q：顾客想续费加时？', 'A：在PC终端上选择续费操作，扫码支付追加时长即可，无需重新开始游戏。'),
        ('Q：头显一直显示「使用中」无法释放？', 'A：在商家后台设备列表中找到该头显，手动操作「释放设备」。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.bold = True
        doc.add_paragraph(a)

    doc.add_heading('术语表', level=3)
    add_styled_table(doc,
        ['术语', '解释'],
        [['Token', '设备认证令牌，PC点播系统绑定店铺的凭证'],
         ['游戏币', '会员端消费代币，1:1抵扣消费金额'],
         ['预存款', '会员人民币储值余额'],
         ['散客', '未注册会员的到店顾客'],
         ['交班', '店员换班时的账务交接']],
        [3, 10])

    out_path = os.path.join(OUTPUT_DIR, '头号空间-使用教程-商家门店篇.docx')
    doc.save(out_path)
    print(f'[OK] {out_path}')
    return out_path

# ============================================================
#  文档 3: 玩家体验篇
# ============================================================

def build_player_docx():
    flow_main = draw_player_flow()
    flow_payment = draw_player_payment()

    doc = Document()
    setup_styles(doc)

    doc.add_heading('头号空间 · 玩家使用说明', level=1)
    p = doc.add_paragraph()
    run = p.add_run('适用版本：v2.19  |  更新日期：2026年7月16日  |  目标角色：到店VR体验顾客')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    add_image_centered(doc, flow_main, 6.0)
    doc.add_paragraph()

    doc.add_paragraph('整个过程简单直观，工作人员会引导你完成每一步。')

    # ===== 1. 全流程 =====
    doc.add_heading('1. 到店体验全流程', level=2)

    doc.add_heading('PC点播系统是什么', level=3)
    doc.add_paragraph('门店里那台大屏幕触摸电脑就是PC点播系统，你可以自主浏览游戏、查看介绍、完成支付。')
    doc.add_paragraph('VR头显内不显示任何菜单和文字，全沉浸式体验', style='List Bullet')
    doc.add_paragraph('散客无需注册登录就能玩', style='List Bullet')
    doc.add_paragraph('会员扫码自动识别身份并享受折扣', style='List Bullet')

    # ===== 2. 选游戏 =====
    doc.add_heading('2. 浏览与选择游戏', level=2)

    doc.add_heading('查看游戏', level=3)
    doc.add_paragraph('1. 走近PC终端，首页展示所有可玩的游戏')
    doc.add_paragraph('2. 每款游戏显示：封面图、名称、时长、价格')
    doc.add_paragraph('3. 点击游戏卡片进入详情页')

    doc.add_heading('游戏详情页', level=3)
    doc.add_paragraph('在详情页可以看到游戏介绍和玩法说明、评分和玩家评价、所需时长（如30分钟/60分钟）、价格信息、支持人数（单人/多人联机）。')

    doc.add_heading('选择体验配置', level=3)
    doc.add_paragraph('1. 选择体验时长')
    doc.add_paragraph('2. 如为多人联机游戏，选择设备数量')
    doc.add_paragraph('3. 点击「开始游戏」按钮进入支付环节')

    # ===== 3. 支付 =====
    doc.add_heading('3. 支付方式', level=2)

    add_image_centered(doc, flow_payment, 5.5)

    doc.add_heading('3.1 微信/支付宝扫码支付（散客/会员均可）', level=3)
    doc.add_paragraph('最常用的支付方式，无需注册会员：')
    doc.add_paragraph('1. 终端屏幕显示支付二维码')
    doc.add_paragraph('2. 打开微信或支付宝「扫一扫」')
    doc.add_paragraph('3. 扫描二维码，手机跳转到支付页面')
    doc.add_paragraph('4. 确认金额，输入密码完成支付')
    doc.add_paragraph('5. 支付成功，终端自动进入下一步')

    doc.add_heading('3.2 会员储值支付（需已注册会员）', level=3)
    doc.add_paragraph('如果你是会员，可使用会员余额和游戏币：')

    p = doc.add_paragraph()
    run = p.add_run('方式A：小程序扫码支付')
    run.font.bold = True
    doc.add_paragraph('1. 终端展示支付二维码')
    doc.add_paragraph('2. 打开头号空间小程序「扫一扫」')
    doc.add_paragraph('3. 小程序自动识别并展示：本次消费金额、预存款余额、可用游戏币、优惠券自动使用')
    doc.add_paragraph('4. 确认支付，系统按顺序自动抵扣：优惠券 → 会员折扣 → 预存款 → 游戏币 → 微信补差')
    doc.add_paragraph('5. 余额不足部分自动计算，用微信支付补差价')

    p = doc.add_paragraph()
    run = p.add_run('方式B：出示会员码')
    run.font.bold = True
    doc.add_paragraph('1. 在微信小程序中打开「会员码」')
    doc.add_paragraph('2. 将二维码对准终端扫码区域')
    doc.add_paragraph('3. 终端识别身份后显示支付确认页')
    doc.add_paragraph('4. 按提示完成支付')

    # ===== 4. 佩戴 =====
    doc.add_heading('4. 佩戴VR设备', level=2)
    doc.add_paragraph('支付成功后，PC终端进入引导界面：')
    doc.add_paragraph('1. 终端显示「请佩戴头显设备」')
    doc.add_paragraph('2. 屏幕展示分配给您的设备编号和佩戴指引图示')
    doc.add_paragraph('3. 按照图示佩戴VR头显：先将头显戴在头上 → 调整顶部绑带松紧 → 调整两侧旋钮至清晰舒适')
    doc.add_paragraph('4. 设备检测到佩戴后自动启动游戏')
    add_note(doc, '如有佩戴问题，举手示意工作人员协助。')

    # ===== 5. 游戏中 =====
    doc.add_heading('5. 游戏中体验', level=2)
    doc.add_paragraph('VR头显内不显示任何菜单和文字，纯粹沉浸式游戏体验：')

    add_styled_table(doc,
        ['功能', '说明'],
        [['自动计时', '系统根据您选择的时长自动计时'],
         ['续费加时', '如需延长时间，可在PC终端选择续费，追加时长继续玩'],
         ['摘盔暂停', '中途取下头显会自动暂停游戏'],
         ['超时结束', '摘下头显超过3分钟未重新佩戴，系统自动结束']],
        [4, 10])

    # ===== 6. 结束 =====
    doc.add_heading('6. 游戏结束', level=2)
    doc.add_paragraph('1. 游戏时长到期或系统自动提醒结束')
    doc.add_paragraph('2. VR头显显示结束画面')
    doc.add_paragraph('3. 取下头显，放回指定位置')
    doc.add_paragraph('4. PC终端自动回到首页，设备恢复为空闲状态')
    doc.add_paragraph('欢迎下次再来体验！')

    # ===== 7. 注册 =====
    doc.add_heading('7. 注册会员', level=2)
    doc.add_paragraph('注册会员可享受折扣、储值优惠、游戏币赠送等福利。')

    p = doc.add_paragraph()
    run = p.add_run('方式一：扫描店铺注册码')
    run.font.bold = True
    doc.add_paragraph('1. 在门店找到店铺注册二维码（通常贴在收银台或终端旁）')
    doc.add_paragraph('2. 微信扫描二维码')
    doc.add_paragraph('3. 填写手机号和基本信息')
    doc.add_paragraph('4. 完成注册，赠送新会员礼包（游戏币+预存款）')

    p = doc.add_paragraph()
    run = p.add_run('方式二：通过小程序注册')
    run.font.bold = True
    doc.add_paragraph('1. 微信搜索「头号空间」小程序')
    doc.add_paragraph('2. 选择最近的门店')
    doc.add_paragraph('3. 填写信息完成注册')

    # ===== 附录 =====
    doc.add_heading('附录：常见问题', level=2)
    faqs = [
        ('Q：不会操作怎么办？', 'A：门店有工作人员全程引导，随时举手示意即可。'),
        ('Q：支付后反悔了能退款吗？', 'A：可以，联系收银员处理退款。'),
        ('Q：近视能玩吗？', 'A：大部分VR头显支持佩戴眼镜使用，也可以调整瞳距。'),
        ('Q：一次能玩多久？', 'A：根据你选择的套餐，通常有30分钟、60分钟等选项。游戏中也可以续费加时。'),
        ('Q：能和朋友一起玩吗？', 'A：部分支持联机的游戏可以多人同时体验，留意「支持人数」说明。'),
        ('Q：会员余额怎么查看？', 'A：打开头号空间小程序，在「我的」页面查看余额和游戏币。'),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.font.bold = True
        doc.add_paragraph(a)

    doc.add_heading('支付优先级说明', level=3)
    doc.add_paragraph('当你是会员时，消费金额按以下顺序自动抵扣：')
    add_styled_table(doc,
        ['优先级', '抵扣项', '说明'],
        [['1', '优惠券', '符合条件的优惠券自动使用'],
         ['2', '会员折扣', '根据会员等级自动打折'],
         ['3', '预存款', '人民币余额优先抵扣'],
         ['4', '游戏币', '1游戏币 = 1元'],
         ['5', '微信补差', '以上抵扣后不足部分，微信支付补足']],
        [2, 4, 7])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('祝您游戏愉快！')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    out_path = os.path.join(OUTPUT_DIR, '头号空间-使用教程-玩家体验篇.docx')
    doc.save(out_path)
    print(f'[OK] {out_path}')
    return out_path

# ============================================================
#  主入口
# ============================================================

if __name__ == '__main__':
    print('正在生成流程图...')
    build_official_docx()
    build_shop_docx()
    build_player_docx()
    print('\n全部完成！生成文件:')
    print(f'  {OUTPUT_DIR}/头号空间-使用教程-官方运营篇.docx')
    print(f'  {OUTPUT_DIR}/头号空间-使用教程-商家门店篇.docx')
    print(f'  {OUTPUT_DIR}/头号空间-使用教程-玩家体验篇.docx')
