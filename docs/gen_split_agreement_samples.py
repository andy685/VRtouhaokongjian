# -*- coding: utf-8 -*-
"""
基于统一分账协议模板，分别以代理商、游戏供应商、商家三种身份
生成已填写的签署样例 docx（占位符替换为示例数据）。
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = os.path.dirname(__file__)

# ===================== 样式工具函数 =====================
def build_doc():
    doc = Document()
    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    normal.font.size = Pt(10.5)
    return doc

def set_cn_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

BLACK = RGBColor(0x00, 0x00, 0x00)

def add_title(doc, text, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, '黑体', size, bold=True)
    p.paragraph_format.space_after = Pt(6)

def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_cn_font(r, '宋体', size, bold=False, color=RGBColor(0x40,0x40,0x40))
    p.paragraph_format.space_after = Pt(12)

def add_clause(doc, text, size=10.5, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, '宋体', size, bold=bold)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(2)

def add_heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_cn_font(r, '黑体', sizes.get(level, 11), bold=True)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_cn_font(r, '黑体', 10, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(val)
            set_cn_font(r, '宋体', 10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()

# ===================== 公共条款（三种身份共用） =====================
def write_common_clauses(doc, partner_type, partner_name, sign_date, contract_no, sign_place):
    add_title(doc, '头号空间平台分账协议')
    add_subtitle(doc, '（拉卡拉支付分账服务配套协议）')

    p = doc.add_paragraph()
    r = p.add_run('合同编号：' + contract_no)
    set_cn_font(r, '宋体', 10.5)
    p = doc.add_paragraph()
    r = p.add_run('签订地点：' + sign_place)
    set_cn_font(r, '宋体', 10.5)
    p = doc.add_paragraph()
    r = p.add_run('签订日期：' + sign_date)
    set_cn_font(r, '宋体', 10.5)

    add_heading(doc, '鉴于')
    add_clause(doc, '1. 头号空间平台（以下简称"平台"）由头号空间（北京）科技有限公司运营，为 VR 线下体验店提供游戏豆采购、内容分发、分润结算等全链路运营管理服务，并通过拉卡拉支付股份有限公司（以下简称"拉卡拉"）的分账服务实现资金的合规清分。')
    add_clause(doc, '2. 平台已与拉卡拉签署《特约商户支付服务合作协议》并开通分账功能，作为分账发起方，有权依据本协议约定的分账规则，通过拉卡拉分账系统将交易资金分配至各分账接收方。')
    add_clause(doc, '3. ' + partner_name + '（以下简称"合作方"）作为平台的' + partner_type + '，与平台存在真实的业务合作关系，并同意接受拉卡拉分账服务的资金清分安排。')
    add_clause(doc, '4. 各方本着平等、自愿、诚实信用的原则，根据《中华人民共和国民法典》《中华人民共和国电子商务法》《非银行支付机构网络支付业务管理办法》及中国人民银行、卡组织相关规定，就分账事宜达成如下协议。')

    add_heading(doc, '第一条  定义')
    add_clause(doc, '1.1 平台：指头号空间平台，由头号空间（北京）科技有限公司运营。')
    add_clause(doc, '1.2 拉卡拉：指拉卡拉支付股份有限公司，持有中国人民银行颁发的《支付业务许可证》，为平台提供支付及分账服务。')
    add_clause(doc, '1.3 分账：指平台根据预先设定的分账接收方及收入分配规则，通过拉卡拉分账系统，将消费者支付的货款或平台收取的游戏豆采购款项，按指定规则分配给参与分账的各方。')
    add_clause(doc, '1.4 分账发起方：指平台，负责通过系统向拉卡拉提交分账指令。')
    add_clause(doc, '1.5 分账接收方：指按本协议约定接收分账资金的各方，包括但不限于代理商、游戏供应商、商家及平台自身。')
    add_clause(doc, '1.6 游戏豆：指平台发行的 B 端运营代币，商家用于启动游戏，平台统一定价 ¥1/豆。')
    add_clause(doc, '1.7 合作方类型：指本协议项下合作方的身份，本协议合作方类型为' + partner_type + '。')
    add_clause(doc, '1.8 分账比例：指分账资金在各方之间的分配比例，以百分比表示，各分账接收方比例之和不超过 100%。')
    add_clause(doc, '1.9 结算周期：指分账资金清算至各分账接收方账户的频率，如 T+1、月结等。')
    add_clause(doc, '1.10 手续费：指拉卡拉就支付及分账服务收取的费用，以及平台就提现/打款收取的费用。')

    add_heading(doc, '第二条  合作内容与分账模式')
    add_clause(doc, '2.1 平台通过拉卡拉分账系统，就以下业务场景实现资金分账：')
    add_clause(doc, '（1）商家向平台采购游戏豆所形成的交易资金；')
    add_clause(doc, '（2）用户游玩消耗游戏豆后，平台向游戏供应商（CP）按次结算的款项；')
    add_clause(doc, '（3）平台向代理商按辖区商家游戏豆消耗量结算的分润款项；')
    add_clause(doc, '（4）其他经双方书面确认的分账场景。')
    add_clause(doc, '2.2 分账模式包括：')
    add_clause(doc, '（1）自动分账：平台系统按预设规则自动向拉卡拉发起分账指令；')
    add_clause(doc, '（2）手动分账：由平台运营人员人工发起分账指令；')
    add_clause(doc, '（3）指定规则分账：按订单维度动态传入分账规则。')
    add_clause(doc, '2.3 本协议采用的分账模式以第三条及附录《分账参数表》约定为准。')

    add_heading(doc, '第四条  拉卡拉分账服务')
    add_clause(doc, '4.1 平台作为拉卡拉特约商户及分账发起方，通过拉卡拉分账系统执行本协议项下的分账指令。拉卡拉依据平台提交的分账规则完成资金清分，不承担本协议项下平台与合作方之间的商业纠纷责任。')
    add_clause(doc, '4.2 分账接收方应在拉卡拉完成入网开户，提供真实、合法、有效的主体资质及银行账户信息。分账接收方账户信息发生变更时，应提前3个工作日书面通知平台及拉卡拉。')
    add_clause(doc, '4.3 拉卡拉分账服务的手续费费率/金额以平台与拉卡拉签署的《特约商户支付服务合作协议》及补充协议为准。')
    add_clause(doc, '4.4 平台保证向拉卡拉提交的分账指令真实、合法、准确，分账资金来源合法，不存在洗钱、套现、欺诈等违法违规情形。')
    add_clause(doc, '4.5 拉卡拉有权根据反洗钱、风控及监管要求，对异常交易进行拦截、延迟结算或冻结，并要求平台及合作方配合提供相关证明材料。')

    add_heading(doc, '第九条  信息安全与保密')
    add_clause(doc, '9.1 各方对因履行本协议获知的他方商业秘密、交易数据、账户信息等负有保密义务，未经书面同意不得向第三方披露。')
    add_clause(doc, '9.2 各方应遵守《中华人民共和国个人信息保护法》《中华人民共和国数据安全法》等规定，采取必要措施保障数据安全。')
    add_clause(doc, '9.3 保密义务在协议终止后3年内继续有效。')

    add_heading(doc, '第十条  违约责任')
    add_clause(doc, '10.1 任何一方违反本协议约定，应承担违约责任，并赔偿守约方因此遭受的直接损失。')
    add_clause(doc, '10.2 合作方提供虚假资质、虚构交易、套取分账资金的，平台有权立即终止分账、冻结相关资金，并要求合作方支付违约金；涉嫌违法的，移送有权机关处理。')
    add_clause(doc, '10.3 因平台未按约定发起分账或分账指令错误导致合作方损失的，平台应予纠正并赔偿合理损失。')
    add_clause(doc, '10.4 因拉卡拉清算系统原因导致的资金延迟或差错，按平台与拉卡拉的协议处理，不视为平台违约。')

    add_heading(doc, '第十一条  协议变更与解除')
    add_clause(doc, '11.1 本协议任何条款的变更或补充，应经各方协商一致并签署书面协议。')
    add_clause(doc, '11.2 发生下列情形之一，一方有权书面通知对方解除本协议：')
    add_clause(doc, '（1）对方严重违约，经书面催告后15日内未纠正；')
    add_clause(doc, '（2）对方丧失继续经营资质或被列入严重违法失信名单；')
    add_clause(doc, '（3）因监管政策变化或拉卡拉分账服务调整，导致本协议无法继续履行；')
    add_clause(doc, '（4）法律法规规定或本协议约定的其他解除情形。')
    add_clause(doc, '11.3 协议解除后，各方应就已发生的交易完成清算与对账，结清款项后互不相欠（另有约定除外）。')

    add_heading(doc, '第十二条  争议解决')
    add_clause(doc, '12.1 本协议适用中华人民共和国法律。')
    add_clause(doc, '12.2 因本协议引起的或与本协议有关的争议，各方应友好协商解决；协商不成的，任何一方均可向平台所在地有管辖权的人民法院提起诉讼。')

    add_heading(doc, '第十三条  其他')
    add_clause(doc, '13.1 本协议自各方法定代表人或授权代表签字（或盖章）并加盖公章（或合同专用章）之日起生效。')
    add_clause(doc, '13.2 本协议一式肆份，各方各执壹份，具有同等法律效力。')
    add_clause(doc, '13.3 本协议与平台及拉卡拉签署的《特约商户支付服务合作协议》及补充协议不一致的，就分账事项以本协议为准；就支付及拉卡拉服务事项以拉卡拉协议为准。')
    add_clause(doc, '13.4 本协议未尽事宜，各方可另行签订补充协议，补充协议与本协议具有同等法律效力。')


def write_sign_page(doc, partner_type, partner_name, account_name, bank_name, bank_account, sign_date, valid_from, valid_to):
    doc.add_page_break()
    add_heading(doc, '签署页')
    add_clause(doc, '（本页为《头号空间平台分账协议》签署页，无正文）')
    add_clause(doc, '')
    add_clause(doc, '甲方（平台）：头号空间（北京）科技有限公司')
    add_clause(doc, '法定代表人/授权代表：张明')
    add_clause(doc, '签字（盖章）：')
    add_clause(doc, '日期：' + sign_date)
    add_clause(doc, '')
    add_clause(doc, '乙方（合作方）：' + partner_name)
    add_clause(doc, '合作方类型：' + partner_type)
    add_clause(doc, '法定代表人/授权代表：__________')
    add_clause(doc, '签字（盖章）：')
    add_clause(doc, '日期：' + sign_date)
    add_clause(doc, '')
    add_clause(doc, '协议有效期限：自 ' + valid_from + ' 至 ' + valid_to)
    add_clause(doc, '结算账户名称：' + account_name)
    add_clause(doc, '开户银行：' + bank_name)
    add_clause(doc, '银行账号：' + bank_account)

# ===================== 样例一：代理商 =====================
def gen_agent_sample():
    doc = build_doc()
    partner_type = '代理商'
    partner_name = '深圳市畅游智能科技有限公司'
    contract_no = 'THKJ-FZ-DL-2026-001'
    sign_date = '2026年7月1日'
    sign_place = '北京市朝阳区'
    valid_from = '2026年7月1日'
    valid_to = '2027年6月30日'
    account_name = '深圳市畅游智能科技有限公司'
    bank_name = '招商银行深圳高新园支行'
    bank_account = '6225 8888 1234 5678'

    write_common_clauses(doc, partner_type, partner_name, sign_date, contract_no, sign_place)

    add_heading(doc, '第三条  分账规则')
    add_clause(doc, '3.1 分账接收方及分账比例')
    add_clause(doc, '各方确认，代理商分账资金按以下规则在各分账接收方之间分配：')
    add_table(doc,
        ['合作方类型', '分账接收方', '分账比例/金额', '结算周期', '备注'],
        [
            ['代理商', '平台', '92%', '月结', '平台收取游戏豆采购收入'],
            ['代理商', '代理商', '7%', '月结', '按辖区商家游戏豆消耗量分润'],
            ['代理商', '拉卡拉(手续费)', '1%', 'T+1', '支付通道费'],
        ],
        widths=[2.5, 2.5, 3, 2, 4]
    )
    add_clause(doc, '3.2 分账比例上限')
    add_clause(doc, '各方确认，分账比例之和为100%，符合拉卡拉分账服务规则。')
    add_clause(doc, '3.3 代理商分润口径')
    add_clause(doc, '代理商分润金额 = 辖区内商家游戏豆消耗量 × 分润单价。其中：')
    add_clause(doc, '（1）分润单价：¥0.07/豆（即游戏豆采购价的7%）；')
    add_clause(doc, '（2）结算依据：以平台系统统计的代理商辖区商家游戏豆消耗记录为准，按月对账结算，次月10日前完成打款；')
    add_clause(doc, '（3）代理商保证金：人民币 50,000元，用于担保履约，协议终止后无争议部分予以退还。')
    add_clause(doc, '3.4 资金冻结与解冻')
    add_clause(doc, '如涉及需管控资金到账时间的场景（如保证金），平台可通过拉卡拉分账的冻结/解冻功能先行冻结资金，待满足解冻条件后释放至对应接收方账户。')

    add_heading(doc, '第五条  各方权利与义务')
    add_heading(doc, '5.1 平台的权利与义务', level=2)
    add_clause(doc, '（1）负责分账系统的对接、运维及分账规则的配置与维护；')
    add_clause(doc, '（2）按本协议约定及结算周期，及时、准确地向拉卡拉发起分账指令；')
    add_clause(doc, '（3）向合作方提供分账明细及对账数据，配合完成对账与差异处理；')
    add_clause(doc, '（4）对合作方的资质真实性、业务合法性进行审核；')
    add_clause(doc, '（5）妥善保管合作方信息及交易数据，不得泄露或滥用。')
    add_heading(doc, '5.2 合作方的权利与义务', level=2)
    add_clause(doc, '（1）保证自身主体资格合法有效，具备签署和履行本协议的资质与能力；')
    add_clause(doc, '（2）按要求在拉卡拉完成入网开户，提供真实、准确的银行账户信息；')
    add_clause(doc, '（3）保证交易背景真实，不得虚构交易、套取分账资金；')
    add_clause(doc, '（4）有权查询自身分账明细，对异常数据及时向平台提出并配合核查；')
    add_clause(doc, '（5）依法承担自身应缴税费，配合平台及拉卡拉的合规要求。')
    add_heading(doc, '5.3 代理商专项义务', level=2)
    add_clause(doc, '（1）按约定缴纳保证金50,000元，维护辖区商家秩序；')
    add_clause(doc, '（2）不得跨区域拓展、不得低价倾销游戏豆；')
    add_clause(doc, '（3）配合平台进行辖区商家的准入审核与日常管理。')

    add_heading(doc, '第六条  资金结算与对账')
    add_clause(doc, '6.1 结算周期：本协议项下分账资金结算周期为月结，次月10日前完成打款，以拉卡拉实际清算时间为准。')
    add_clause(doc, '6.2 对账机制：平台应于每月5日前向代理商提供上月分账对账单。代理商应在收到对账单后3个工作日内完成核对；如有异议，应在异议期内书面提出，逾期未提出视为确认。')
    add_clause(doc, '6.3 差异处理：对账差异经双方核实后，在下一结算周期内多退少补。')

    add_heading(doc, '第七条  手续费与税费')
    add_clause(doc, '7.1 拉卡拉支付及分账服务手续费由平台承担。')
    add_clause(doc, '7.2 平台提现/打款手续费：按打款金额的0.3%收取，由代理商承担，打款时由平台代扣。')
    add_clause(doc, '7.3 各方因履行本协议产生的税费，由各方依法各自承担。')

    add_heading(doc, '第八条  退款与冲正')
    add_clause(doc, '8.1 因消费者退款、交易撤销等导致原分账资金需冲正的，平台有权通过拉卡拉分账系统发起冲正指令，代理商应予配合。')
    add_clause(doc, '8.2 若代理商账户余额不足以冲正，代理商应在接到平台通知后5个工作日内返还差额款项。')

    write_sign_page(doc, partner_type, partner_name, account_name, bank_name, bank_account, sign_date, valid_from, valid_to)

    # 附录
    doc.add_page_break()
    add_heading(doc, '附录  代理商分账参数表')
    add_table(doc,
        ['参数项', '内容'],
        [
            ['代理商名称', partner_name],
            ['代理商级别', '市级代理商'],
            ['辖区范围', '广东省深圳市'],
            ['保证金金额', '¥50,000'],
            ['分润口径', '按辖区商家游戏豆消耗量'],
            ['分润单价/比例', '¥0.07/豆（游戏豆采购价的7%）'],
            ['结算周期', '月结，次月10日前'],
            ['结算账户名', account_name],
            ['开户行', bank_name],
            ['银行账号', bank_account],
            ['提现手续费承担方', '代理商'],
            ['提现手续费率', '0.3%（打款时代扣）'],
        ],
        widths=[4, 12]
    )

    out = os.path.join(OUT_DIR, '分账协议样例-代理商-深圳市畅游智能科技有限公司.docx')
    doc.save(out)
    print('已生成：', out)


# ===================== 样例二：游戏供应商 =====================
def gen_cp_sample():
    doc = build_doc()
    partner_type = '游戏供应商'
    partner_name = '上海幻境数字娱乐有限公司'
    contract_no = 'THKJ-FZ-CP-2026-001'
    sign_date = '2026年7月1日'
    sign_place = '北京市朝阳区'
    valid_from = '2026年7月1日'
    valid_to = '2027年6月30日'
    account_name = '上海幻境数字娱乐有限公司'
    bank_name = '中国工商银行上海市徐汇支行'
    bank_account = '1001 2888 0902 3456'

    write_common_clauses(doc, partner_type, partner_name, sign_date, contract_no, sign_place)

    add_heading(doc, '第三条  分账规则')
    add_clause(doc, '3.1 分账接收方及分账比例')
    add_clause(doc, '各方确认，游戏供应商分账资金按以下规则在各分账接收方之间分配：')
    add_table(doc,
        ['合作方类型', '分账接收方', '分账比例/金额', '结算周期', '备注'],
        [
            ['游戏供应商', '平台', '80%', '月结', '平台毛利'],
            ['游戏供应商', '游戏供应商', '19%', '月结', '体验次数×单次成本价'],
            ['游戏供应商', '拉卡拉(手续费)', '1%', 'T+1', '支付通道费'],
        ],
        widths=[2.5, 2.5, 3, 2, 4]
    )
    add_clause(doc, '3.2 分账比例上限')
    add_clause(doc, '各方确认，分账比例之和为100%，符合拉卡拉分账服务规则。')
    add_clause(doc, '3.3 游戏供应商（CP）结算口径')
    add_clause(doc, '游戏供应商的结算金额 = 有效体验次数 × 单次结算成本价。其中：')
    add_clause(doc, '（1）有效体验次数：以平台系统记录的、经平台与游戏供应商共同确认的有效游玩次数为准；')
    add_clause(doc, '（2）单次结算成本价：由游戏供应商设定并经平台审核确认，单价为 ¥2.00/次；')
    add_clause(doc, '（3）提现/打款手续费由游戏供应商承担，费率为提现金额的0.3%。')
    add_clause(doc, '3.4 资金冻结与解冻')
    add_clause(doc, '如涉及需管控资金到账时间的场景，平台可通过拉卡拉分账的冻结/解冻功能先行冻结资金，待满足解冻条件后释放至对应接收方账户。')

    add_heading(doc, '第五条  各方权利与义务')
    add_heading(doc, '5.1 平台的权利与义务', level=2)
    add_clause(doc, '（1）负责分账系统的对接、运维及分账规则的配置与维护；')
    add_clause(doc, '（2）按本协议约定及结算周期，及时、准确地向拉卡拉发起分账指令；')
    add_clause(doc, '（3）向合作方提供分账明细及对账数据，配合完成对账与差异处理；')
    add_clause(doc, '（4）对合作方的资质真实性、业务合法性进行审核；')
    add_clause(doc, '（5）妥善保管合作方信息及交易数据，不得泄露或滥用。')
    add_heading(doc, '5.2 合作方的权利与义务', level=2)
    add_clause(doc, '（1）保证自身主体资格合法有效，具备签署和履行本协议的资质与能力；')
    add_clause(doc, '（2）按要求在拉卡拉完成入网开户，提供真实、准确的银行账户信息；')
    add_clause(doc, '（3）保证交易背景真实，不得虚构交易、套取分账资金；')
    add_clause(doc, '（4）有权查询自身分账明细，对异常数据及时向平台提出并配合核查；')
    add_clause(doc, '（5）依法承担自身应缴税费，配合平台及拉卡拉的合规要求。')
    add_heading(doc, '5.3 游戏供应商专项义务', level=2)
    add_clause(doc, '（1）保证所供游戏内容合法合规，拥有完整知识产权或授权；')
    add_clause(doc, '（2）保证游戏稳定运行，及时处理故障与客诉；')
    add_clause(doc, '（3）单次结算成本价一经确认不得单方随意调整，调整须提前15日书面通知并经平台同意。')

    add_heading(doc, '第六条  资金结算与对账')
    add_clause(doc, '6.1 结算周期：本协议项下分账资金结算周期为月结，次月15日前完成打款，以拉卡拉实际清算时间为准。')
    add_clause(doc, '6.2 对账机制：平台应于每月10日前向游戏供应商提供上月分账对账单。游戏供应商应在收到对账单后3个工作日内完成核对；如有异议，应在异议期内书面提出，逾期未提出视为确认。')
    add_clause(doc, '6.3 差异处理：对账差异经双方核实后，在下一结算周期内多退少补。')

    add_heading(doc, '第七条  手续费与税费')
    add_clause(doc, '7.1 拉卡拉支付及分账服务手续费由平台承担。')
    add_clause(doc, '7.2 平台提现/打款手续费：按提现金额的0.3%收取，由游戏供应商承担，打款时由平台代扣。')
    add_clause(doc, '7.3 各方因履行本协议产生的税费，由各方依法各自承担。')

    add_heading(doc, '第八条  退款与冲正')
    add_clause(doc, '8.1 游戏供应商结算款项如因有效体验次数核减、客诉退款等原因需调整的，在下一结算周期内据实扣减。')
    add_clause(doc, '8.2 若游戏供应商账户余额不足以冲正，游戏供应商应在接到平台通知后5个工作日内返还差额款项。')

    write_sign_page(doc, partner_type, partner_name, account_name, bank_name, bank_account, sign_date, valid_from, valid_to)

    # 附录
    doc.add_page_break()
    add_heading(doc, '附录  游戏供应商分账参数表')
    add_table(doc,
        ['参数项', '内容'],
        [
            ['游戏供应商名称', partner_name],
            ['游戏名称/编号', '《星际穿越VR》(THKJ-CP-001)'],
            ['单次结算成本价', '¥2.00/次'],
            ['有效体验次数口径', '平台系统记录并经双方确认'],
            ['结算周期', '月结，次月15日前'],
            ['结算账户名', account_name],
            ['开户行', bank_name],
            ['银行账号', bank_account],
            ['提现手续费承担方', '游戏供应商'],
            ['提现手续费率', '0.3%（打款时代扣）'],
            ['知识产权承诺', '游戏内容合法合规，权属清晰'],
        ],
        widths=[4, 12]
    )

    out = os.path.join(OUT_DIR, '分账协议样例-游戏供应商-上海幻境数字娱乐有限公司.docx')
    doc.save(out)
    print('已生成：', out)


# ===================== 样例三：商家 =====================
def gen_merchant_sample():
    doc = build_doc()
    partner_type = '商家'
    partner_name = '成都沉浸世界VR体验馆'
    contract_no = 'THKJ-FZ-MT-2026-001'
    sign_date = '2026年7月1日'
    sign_place = '北京市朝阳区'
    valid_from = '2026年7月1日'
    valid_to = '2027年6月30日'
    account_name = '成都沉浸世界VR体验馆'
    bank_name = '中国建设银行成都春熙路支行'
    bank_account = '5100 1234 5678 9012'

    write_common_clauses(doc, partner_type, partner_name, sign_date, contract_no, sign_place)

    add_heading(doc, '第三条  分账规则')
    add_clause(doc, '3.1 分账接收方及分账比例')
    add_clause(doc, '各方确认，商家采购游戏豆交易资金按以下规则在各分账接收方之间分配：')
    add_table(doc,
        ['合作方类型', '分账接收方', '分账比例/金额', '结算周期', '备注'],
        [
            ['商家', '平台', '99%', 'T+1', '游戏豆采购收入'],
            ['商家', '拉卡拉(手续费)', '1%', 'T+1', '支付通道费'],
        ],
        widths=[2.5, 2.5, 3, 2, 4]
    )
    add_clause(doc, '3.2 分账比例上限')
    add_clause(doc, '各方确认，分账比例之和为100%，符合拉卡拉分账服务规则。')
    add_clause(doc, '3.3 商家采购口径')
    add_clause(doc, '商家向平台采购游戏豆，单价 ¥1/豆，采购款项通过拉卡拉支付并按分账规则清分。商家为游戏豆的购买方与使用方，不直接参与C端消费资金的分账。')
    add_clause(doc, '3.4 资金冻结与解冻')
    add_clause(doc, '如涉及需管控资金到账时间的场景（如预付款），平台可通过拉卡拉分账的冻结/解冻功能先行冻结资金，待满足解冻条件后释放至对应接收方账户。')

    add_heading(doc, '第五条  各方权利与义务')
    add_heading(doc, '5.1 平台的权利与义务', level=2)
    add_clause(doc, '（1）负责分账系统的对接、运维及分账规则的配置与维护；')
    add_clause(doc, '（2）按本协议约定及结算周期，及时、准确地向拉卡拉发起分账指令；')
    add_clause(doc, '（3）向合作方提供分账明细及对账数据，配合完成对账与差异处理；')
    add_clause(doc, '（4）对合作方的资质真实性、业务合法性进行审核；')
    add_clause(doc, '（5）妥善保管合作方信息及交易数据，不得泄露或滥用。')
    add_heading(doc, '5.2 合作方的权利与义务', level=2)
    add_clause(doc, '（1）保证自身主体资格合法有效，具备签署和履行本协议的资质与能力；')
    add_clause(doc, '（2）按要求在拉卡拉完成入网开户，提供真实、准确的银行账户信息；')
    add_clause(doc, '（3）保证交易背景真实，不得虚构交易、套取分账资金；')
    add_clause(doc, '（4）有权查询自身分账明细，对异常数据及时向平台提出并配合核查；')
    add_clause(doc, '（5）依法承担自身应缴税费，配合平台及拉卡拉的合规要求。')
    add_heading(doc, '5.3 商家专项义务', level=2)
    add_clause(doc, '（1）合法经营，按平台定价采购游戏豆，不得私下倒卖；')
    add_clause(doc, '（2）保证C端消费交易真实，妥善处理消费者退款与投诉；')
    add_clause(doc, '（3）配合平台及拉卡拉的资质审核与风控检查。')

    add_heading(doc, '第六条  资金结算与对账')
    add_clause(doc, '6.1 结算周期：本协议项下分账资金结算周期为T+1，以拉卡拉实际清算时间为准。')
    add_clause(doc, '6.2 对账机制：平台应于每日向商家提供前一交易日的分账对账单。商家应在收到对账单后3个工作日内完成核对；如有异议，应在异议期内书面提出，逾期未提出视为确认。')
    add_clause(doc, '6.3 差异处理：对账差异经双方核实后，在下一结算周期内多退少补。')

    add_heading(doc, '第七条  手续费与税费')
    add_clause(doc, '7.1 拉卡拉支付及分账服务手续费由平台承担。')
    add_clause(doc, '7.2 平台无额外提现手续费（商家采购为实时支付）。')
    add_clause(doc, '7.3 各方因履行本协议产生的税费，由各方依法各自承担。')

    add_heading(doc, '第八条  退款与冲正')
    add_clause(doc, '8.1 因消费者退款、交易撤销等导致原分账资金需冲正的，平台有权通过拉卡拉分账系统发起冲正指令，商家应予配合。')
    add_clause(doc, '8.2 若商家账户余额不足以冲正，商家应在接到平台通知后5个工作日内返还差额款项。')

    write_sign_page(doc, partner_type, partner_name, account_name, bank_name, bank_account, sign_date, valid_from, valid_to)

    # 附录
    doc.add_page_break()
    add_heading(doc, '附录  商家分账参数表')
    add_table(doc,
        ['参数项', '内容'],
        [
            ['商家名称', partner_name],
            ['店铺/门店编号', 'THKJ-MT-CD-001'],
            ['所属代理商', '深圳市畅游智能科技有限公司'],
            ['游戏豆采购单价', '¥1/豆（平台统一定价）'],
            ['结算周期', 'T+1'],
            ['结算账户名', account_name],
            ['开户行', bank_name],
            ['银行账号', bank_account],
            ['支付手续费承担方', '平台'],
            ['备注', '无'],
        ],
        widths=[4, 12]
    )

    out = os.path.join(OUT_DIR, '分账协议样例-商家-成都沉浸世界VR体验馆.docx')
    doc.save(out)
    print('已生成：', out)


# ===================== 主入口 =====================
if __name__ == '__main__':
    gen_agent_sample()
    gen_cp_sample()
    gen_merchant_sample()
    print('全部样例生成完成。')
