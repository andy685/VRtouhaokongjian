from __future__ import annotations

from datetime import date
from pathlib import Path
import math
import textwrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "generated"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "头号空间-反扫核销PRD与异常处理流程-独立提炼版.docx"
FLOW_PNG = ASSET_DIR / "反扫核销主流程.png"
STATE_PNG = ASSET_DIR / "异常处理状态机.png"


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def pick_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size, index=0)
        except Exception:
            continue
    return ImageFont.load_default()


def draw_wrapped_text(
    draw: ImageDraw.ImageDraw,
    text: str,
    box: tuple[int, int, int, int],
    font: ImageFont.ImageFont,
    fill: str,
    align: str = "left",
    line_spacing: int = 8,
) -> None:
    x1, y1, x2, y2 = box
    width = x2 - x1
    avg_char_width = max(font.size, 12)
    max_chars = max(6, width // avg_char_width)
    lines = []
    for raw in text.split("\n"):
        lines.extend(textwrap.wrap(raw, width=max_chars) or [""])
    y = y1
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_w = bbox[2] - bbox[0]
        line_h = bbox[3] - bbox[1]
        if align == "center":
            x = x1 + (width - line_w) / 2
        else:
            x = x1
        draw.text((x, y), line, font=font, fill=fill)
        y += line_h + line_spacing
        if y > y2:
            break


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    fill: str,
    outline: str = "#214276",
    title_fill: str = "#12233D",
) -> None:
    draw.rounded_rectangle(box, radius=24, fill=fill, outline=outline, width=3)
    title_font = pick_font(30, bold=True)
    body_font = pick_font(22)
    x1, y1, x2, y2 = box
    draw_wrapped_text(draw, title, (x1 + 24, y1 + 18, x2 - 24, y1 + 70), title_font, title_fill)
    draw.line((x1 + 20, y1 + 78, x2 - 20, y1 + 78), fill=outline, width=2)
    draw_wrapped_text(draw, body, (x1 + 24, y1 + 94, x2 - 24, y2 - 18), body_font, "#1E293B")


def arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    label: str | None = None,
    color: str = "#335C9A",
) -> None:
    draw.line((start, end), fill=color, width=6)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    head_len = 18
    for delta in (math.pi / 7, -math.pi / 7):
        x = end[0] - head_len * math.cos(angle + delta)
        y = end[1] - head_len * math.sin(angle + delta)
        draw.line((end, (x, y)), fill=color, width=6)
    if label:
        mid_x = (start[0] + end[0]) / 2
        mid_y = (start[1] + end[1]) / 2
        font = pick_font(20)
        bbox = draw.textbbox((0, 0), label, font=font)
        pad = 8
        label_box = (
            mid_x - (bbox[2] - bbox[0]) / 2 - pad,
            mid_y - (bbox[3] - bbox[1]) / 2 - pad,
            mid_x + (bbox[2] - bbox[0]) / 2 + pad,
            mid_y + (bbox[3] - bbox[1]) / 2 + pad,
        )
        draw.rounded_rectangle(label_box, radius=10, fill="#F8FAFC", outline="#CBD5E1")
        draw.text((label_box[0] + pad, label_box[1] + pad - 2), label, font=font, fill="#334155")


def generate_flow_diagram() -> None:
    img = Image.new("RGB", (1800, 980), "#F6F8FC")
    draw = ImageDraw.Draw(img)
    title_font = pick_font(42, bold=True)
    sub_font = pick_font(22)
    draw.text((70, 48), "反扫核销主流程（提炼版）", font=title_font, fill="#102A43")
    draw.text(
        (70, 108),
        "依据 PRD §5.8、§4.3.4 与交易参考文档整理；强调小程序承载确认、终端承载识别、后端承载订单与 Session。",
        font=sub_font,
        fill="#486581",
    )

    boxes = [
        ((70, 200, 410, 420), "1. 用户出示会员码", "小程序生成含会员ID的二维码\n每5分钟自动刷新\n用户到店出示给终端或扫码枪", "#E9F2FF"),
        ((490, 200, 830, 420), "2. 终端识别会员", "扫码输入抽象层接收扫码事件\n支持真实扫码与模拟扫码\n解析二维码并拉取会员实时信息", "#EAFBF3"),
        ((910, 200, 1250, 420), "3. 小程序确认支付", "终端通知小程序展示支付确认页\n展示订单明细、会员资产、优惠与补差\n用户在手机端确认支付", "#FFF4D8"),
        ((1330, 200, 1670, 420), "4. 创单与开局", "后端创建订单、扣款并启动会话\n终端提示支付成功与佩戴指引\n体验结束后调用结算/结束接口", "#FFE7E7"),
    ]
    for box, title, body, fill in boxes:
        rounded_box(draw, box, title, body, fill)

    arrow(draw, (410, 310), (490, 310), "扫码事件")
    arrow(draw, (830, 310), (910, 310), "通知确认")
    arrow(draw, (1250, 310), (1330, 310), "支付结果")

    footer_title = pick_font(28, bold=True)
    footer_body = pick_font(22)
    footer_boxes = [
        ((110, 520, 540, 850), "关键规则", "• 不做离线降级，断网即暂停服务\n• 扫码硬件通过抽象层接入，不绑定具体型号\n• 同一会员短时重复扫码必须防重\n• 关键动作须支持幂等与重试", "#FFFFFF"),
        ((685, 520, 1115, 850), "异常收口", "• 覆盖成功、失败、取消、超时四类状态\n• 关注扫码异常、账户异常、扣款异常、创单异常\n• 已扣款未开局必须自动回滚或进入人工处理", "#FFFFFF"),
        ((1260, 520, 1690, 850), "当前可确认的接口口径", "• 终端下单：/api/terminal/trade/consumption-orders\n• 提交支付：/api/terminal/trade/payments\n• 会员查询/创单/开局/结束：PRD已给出流程口径\n• 结构开孔尺寸图暂不在现有资料范围内", "#FFFFFF"),
    ]
    for box, title, body, fill in footer_boxes:
        draw.rounded_rectangle(box, radius=26, fill=fill, outline="#CFD8E3", width=3)
        x1, y1, x2, y2 = box
        draw.text((x1 + 24, y1 + 20), title, font=footer_title, fill="#102A43")
        draw_wrapped_text(draw, body, (x1 + 24, y1 + 76, x2 - 24, y2 - 20), footer_body, "#243B53")

    img.save(FLOW_PNG)


def generate_state_diagram() -> None:
    img = Image.new("RGB", (1800, 980), "#FAFBFD")
    draw = ImageDraw.Draw(img)
    title_font = pick_font(42, bold=True)
    sub_font = pick_font(22)
    draw.text((70, 48), "异常处理状态机（提炼版）", font=title_font, fill="#102A43")
    draw.text(
        (70, 108),
        "依据 PRD §4.6.3、§4.8.4.2、异常订单技术说明与退款流程文档整理。",
        font=sub_font,
        fill="#486581",
    )

    nodes = {
        "pending": ((760, 180, 1060, 320), "待处理", "订单已被标记为异常\n可由商家处理或提交平台", "#FFF5C2"),
        "escalated": ((760, 400, 1060, 540), "提交平台", "商家提交平台审核\n或驳回后重新提交", "#DCEEFF"),
        "rejected": ((1180, 400, 1480, 540), "已驳回", "平台驳回并要求填写驳回原因\n商家可修正重提或自行解决", "#FFE0E0"),
        "resolved": ((760, 660, 1060, 800), "已处理", "终态\n已解决、退款处理完成或商家自行解决", "#DCF7E7"),
        "ignore": ((340, 400, 640, 540), "列表移除", "仅清理异常列表\n不再进入后续审核流转", "#F1F5F9"),
    }
    for box, title, body, fill in nodes.values():
        rounded_box(draw, box, title, body, fill)

    arrow(draw, (910, 320), (910, 400), "提交平台")
    arrow(draw, (760, 250), (640, 430), "忽略")
    arrow(draw, (1060, 250), (1210, 430), "平台驳回")
    arrow(draw, (910, 540), (910, 660), "已解决/退款")
    arrow(draw, (1180, 470), (1060, 720), "自行解决")
    arrow(draw, (1180, 470), (1060, 470), "修正后重提")

    note_box = (120, 690, 620, 880)
    draw.rounded_rectangle(note_box, radius=26, fill="#FFFFFF", outline="#CFD8E3", width=3)
    note_title = pick_font(28, bold=True)
    note_body = pick_font(22)
    draw.text((148, 718), "处理口径补充", font=note_title, fill="#102A43")
    draw_wrapped_text(
        draw,
        "• 商家端：标记入口覆盖6类订单列表页\n• 平台端：审核选项为“已解决 / 驳回”，不再提供退款入口\n• 已结算退款须上传线下退款凭证\n• 点播体验是否退款，要结合体验是否有效与人工判责",
        (148, 774, 592, 860),
        note_body,
        "#243B53",
    )

    img.save(STATE_PNG)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_chinese_font(run, font_name: str = "PingFang SC") -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    r_fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(document: Document, text: str, *, style: str | None = None, bold: bool = False, color: str | None = None):
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_chinese_font(run)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def build_document() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("头号空间-反扫核销PRD与异常处理流程\n独立提炼版")
    set_chinese_font(run)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x10, 0x2A, 0x43)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f"生成日期：{date.today().isoformat()}    |    产出方式：依据现有资料提炼，不修改原PRD")
    set_chinese_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x48, 0x65, 0x81)

    doc.add_paragraph("")

    add_paragraph(doc, "一、结论", style="Heading 1")
    add_paragraph(
        doc,
        "可以。在不调整现有 PRD 文档的前提下，结合仓库内已有 PRD、交易参考、异常订单技术说明和退款流程文档，已经能够提炼出截图中框红的两项交付内容：",
    )
    add_paragraph(doc, "1. 《反扫核销PRD》：可以整理出独立的业务说明、流程、规则、系统边界和当前可确认的接口口径。")
    add_paragraph(doc, "2. 《异常处理流程》：可以整理出异常标记、商家处理、平台审核、退款协同和状态流转的完整闭环。")
    add_paragraph(
        doc,
        "当前边界：现有资料足以形成产品/研发联调文档，但不足以直接产出“机台开孔尺寸图、硬件结构图、最终扫码枪型号接口标准”这类硬件量产资料，本次文档会明确这一边界。",
        color="9F1239",
    )

    add_paragraph(doc, "二、资料来源与提炼范围", style="Heading 1")
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "来源文件"
    hdr[1].text = "核心章节"
    hdr[2].text = "用于提炼的内容"
    for cell in hdr:
        set_cell_shading(cell, "DCE6F7")
        for p in cell.paragraphs:
            for r in p.runs:
                set_chinese_font(r)
                r.bold = True
    rows = [
        ("docs/头号空间-产品需求文档-PRD-v1.3.md", "§4.3.4 / §5.8 / §10.10", "反扫核销流程、点播终端规则、退款与异常处理唯一口径"),
        ("docs/头号空间-PC点播系统-设计说明文档.md", "§4.1.4 / §5.3", "支付与身份识别、反扫/主动扫码两类链路的设计确认"),
        ("docs/异常订单与对账系统-技术说明.md", "§4 / §5 / §6", "异常标记入口、处理弹窗、共享状态与状态机"),
        ("docs/退款交互流程文档.md", "§2 / §4", "已结算退款、线下退款凭证、退款按钮与交互要求"),
        ("docs/2026-06-18-trade-prd-reference.md", "业务入口 / 支付规则", "终端交易、支付接口与 authCode 等字段口径"),
    ]
    for source, section_text, scope in rows:
        cells = table.add_row().cells
        cells[0].text = source
        cells[1].text = section_text
        cells[2].text = scope
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r)

    add_paragraph(doc, "三、《反扫核销PRD》提炼稿", style="Heading 1")
    add_paragraph(doc, "3.1 文档目标", style="Heading 2")
    add_paragraph(
        doc,
        "定义 VR 门店场景下，会员通过小程序会员码被设备端反扫识别后，如何完成会员确认、支付确认、创单、开局、计时、结束与异常收口的完整链路。",
    )
    add_paragraph(doc, "3.2 核心结论", style="Heading 2")
    bullets = [
        "会员码内容：二维码含会员ID，由后端查询实时余额和会员信息；二维码每5分钟自动刷新。",
        "确认端口径：玩家主动扫码支付走小程序；会员码被设备反扫后，支付明细和确认仍在小程序；店员点播走手机端 H5；收银员开单在收银台界面确认。",
        "终端边界：点播终端负责识别扫码结果、触发查询、展示等待与结果；不在大屏展示会员资产扣减明细。",
        "网络依赖：不做离线降级，断网即暂停服务。",
        "研发约束：扫码硬件通过抽象层接入；无硬件阶段必须支持模拟扫码；关键动作需要防重、幂等和异常收口。",
    ]
    for item in bullets:
        add_paragraph(doc, f"• {item}")

    doc.add_picture(str(FLOW_PNG), width=Inches(6.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "图 1  反扫核销主流程图", color="486581").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "3.3 参与系统与职责", style="Heading 2")
    role_table = doc.add_table(rows=1, cols=2)
    role_table.style = "Table Grid"
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    role_hdr = role_table.rows[0].cells
    role_hdr[0].text = "系统"
    role_hdr[1].text = "职责"
    for cell in role_hdr:
        set_cell_shading(cell, "E8EEF7")
        for p in cell.paragraphs:
            for r in p.runs:
                set_chinese_font(r)
                r.bold = True
    role_rows = [
        ("小程序（miniapp-payment）", "生成会员码；承载主动扫码与反扫后的支付确认页；展示会员资产、优惠和补差；确认支付。"),
        ("点播系统", "接收扫码结果；查询会员；通知小程序展示确认页；展示支付结果与设备引导；触发 Session 开始与结束。"),
        ("收银台（cashier-ui）", "扫描 tab 接入同一扫码输入抽象层；收银员开单时在收银台界面完成明细确认。"),
        ("手机端 H5", "承载员工点播身份校验与确认。"),
        ("商家后台（admin-dashboard）", "消费订单、Session、设备状态与异常数据的查询、处理和经营分析。"),
    ]
    for left, right in role_rows:
        cells = role_table.add_row().cells
        cells[0].text = left
        cells[1].text = right
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r)

    add_paragraph(doc, "3.4 当前可确认的接口口径", style="Heading 2")
    add_paragraph(doc, "说明：以下内容来自现有交易参考与 PRD 流程口径，可用于研发联调文档；不等同于最终后端接口定稿。")
    api_table = doc.add_table(rows=1, cols=3)
    api_table.style = "Table Grid"
    api_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    api_hdr = api_table.rows[0].cells
    api_hdr[0].text = "接口/字段"
    api_hdr[1].text = "当前口径"
    api_hdr[2].text = "来源"
    for cell in api_hdr:
        set_cell_shading(cell, "E8EEF7")
        for p in cell.paragraphs:
            for r in p.runs:
                set_chinese_font(r)
                r.bold = True
    api_rows = [
        ("POST /api/terminal/trade/consumption-orders", "终端/PC 消费下单，来源固定为终端。", "trade-prd-reference"),
        ("POST /api/terminal/trade/payments", "终端提交支付；拉卡拉被扫支付通常要求携带 authCode。", "trade-prd-reference"),
        ("authCode", "扫码枪扫到的客户付款码；拉卡拉被扫支付必填。", "trade-prd-reference"),
        ("POST /api/member/{id}/query", "PRD 已明确存在会员实时查询动作，用于获取会员名/等级/余额/折扣。", "PRD §5.8"),
        ("POST /api/order/create", "PRD 已明确创单+扣余额步骤。", "PRD §5.8"),
        ("POST /api/session/start / end", "PRD 已明确 Session 启动、心跳、结束与结算口径。", "PRD §5.8"),
    ]
    for a, b, c in api_rows:
        cells = api_table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_chinese_font(r)

    add_paragraph(doc, "3.5 必须写入独立文档的研发约束", style="Heading 2")
    constraints = [
        "扫码硬件适配：统一走“扫码输入抽象层”，系统消费标准化扫码结果事件，不直接绑定具体扫码枪厂商实现。",
        "无硬件阶段开发：在真实扫码枪接入前，必须支持模拟扫码输入，用于完成识别、扣款、开局、结束和异常验证。",
        "防重与并发：同一会员码短时间重复扫码必须防重；同一会员同一时刻不得在多个点播终端同时成功开局。",
        "异常覆盖：至少覆盖扫码异常、会员账户异常、扣款异常、订单创建异常、Session 启动异常、设备离线/占用异常、网络异常。",
        "异常收口：已扣款但未成功开局时，必须自动回滚或进入人工处理，不允许形成长期悬挂订单。",
    ]
    for item in constraints:
        add_paragraph(doc, f"• {item}")

    add_paragraph(doc, "四、《异常处理流程》提炼稿", style="Heading 1")
    add_paragraph(doc, "4.1 异常来源", style="Heading 2")
    for item in [
        "订单页人工标记：商家后台 6 类订单列表页可直接标记异常。",
        "自动对账：系统模拟或真实比对后发现差异，写入异常列表。",
        "上传对账单：导入 xlsx/csv 后发现差异，写入异常列表。",
        "手动录入核对：按订单号逐笔核对渠道金额，发现差异后写入异常列表。",
        "平台巡检：平台端订单流水页也可标记异常，来源记为“平台巡检”。",
    ]:
        add_paragraph(doc, f"• {item}")

    add_paragraph(doc, "4.2 商家端处理口径", style="Heading 2")
    merchant_table = doc.add_table(rows=1, cols=3)
    merchant_table.style = "Table Grid"
    merchant_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    merchant_hdr = merchant_table.rows[0].cells
    merchant_hdr[0].text = "状态"
    merchant_hdr[1].text = "弹窗标题"
    merchant_hdr[2].text = "可选操作"
    for cell in merchant_hdr:
        set_cell_shading(cell, "E8EEF7")
        for p in cell.paragraphs:
            for r in p.runs:
                set_chinese_font(r)
                r.bold = True
    merchant_rows = [
        ("pending", "处理异常订单", "已解决 / 退款 / 提交平台 / 忽略"),
        ("escalated", "处理异常订单", "已解决 / 退款 / 提交平台 / 忽略"),
        ("rejected", "处理驳回订单", "修正后重新提交平台 / 自行解决"),
        ("resolved", "查看异常订单", "只读查看"),
    ]
    for row in merchant_rows:
        cells = merchant_table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = val
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_chinese_font(r)

    doc.add_picture(str(STATE_PNG), width=Inches(6.8))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "图 2  异常处理状态机", color="486581").alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, "4.3 平台端审核口径", style="Heading 2")
    add_paragraph(doc, "平台对账中心汇总商家提交与平台巡检的异常订单。平台端审核动作收口为两项：")
    add_paragraph(doc, "• 已解决：pending / escalated → resolved。")
    add_paragraph(doc, "• 驳回：pending / escalated → rejected，且驳回原因必填并写入 handleRemark。")
    add_paragraph(doc, "平台端不再提供退款入口；退款由商家在商家后台处理。")

    add_paragraph(doc, "4.4 退款协同口径", style="Heading 2")
    refund_items = [
        "已结算订单退款：店铺线下退款，系统记录退款凭证、操作人、原因和金额。",
        "线下退款凭证：已结算退款必须上传图片凭证，未上传时确认按钮禁用。",
        "预存款/游戏币退款：系统退回会员预存款或游戏币，冲正商家收入确认，不产生拉卡拉退款。",
        "点播体验退款例外：用户主动提前结束、摘盔超时、个人原因中途退出默认不自动退款；设备故障、游戏崩溃、心跳异常等场景进入异常审核。",
    ]
    for item in refund_items:
        add_paragraph(doc, f"• {item}")

    add_paragraph(doc, "五、建议作为独立交付物的文档结构", style="Heading 1")
    for item in [
        "封面与结论：明确“基于现有资料提炼，不修改原 PRD”。",
        "第一部分《反扫核销PRD》：背景、适用范围、参与系统、主流程图、关键规则、研发约束、接口口径、边界说明。",
        "第二部分《异常处理流程》：异常来源、标记入口、状态机图、商家端处理、平台端审核、退款协同、体验异常判责。",
        "附录：来源章节索引，便于回溯到原始 PRD 与技术说明。",
    ]:
        add_paragraph(doc, f"• {item}")

    add_paragraph(doc, "六、本次提炼的边界说明", style="Heading 1")
    add_paragraph(
        doc,
        "本稿已经覆盖截图中框红的两项内容，但不主动新增原资料中不存在的规则。特别是“机台开孔尺寸、结构件定位、扫码枪最终通信协议、硬件量产图纸”未在当前仓库资料中形成可直接出图的依据，因此本次仅能输出产品/研发说明文档，不输出硬件结构图纸。",
        color="9F1239",
    )

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    generate_flow_diagram()
    generate_state_diagram()
    build_document()
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
